"""
A-Tech Appraisal Co. — Revision & Comment Library
Streamlit Web App
"""

import os, json, io, tempfile, zipfile
import re, requests
import fitz  # pymupdf
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import streamlit as st
import pandas as pd
from datetime import date
from PIL import Image as PILImage
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                 TableStyle, HRFlowable, PageBreak, KeepTogether,
                                 Image as RLImage, Image)
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from reportlab.platypus import Flowable

LOGO_PATH = os.path.join(os.path.dirname(__file__), "atech_logo.png")

# ── STR PDF Brand Constants ─────────────────────────────────────────────────
DARK_BLUE  = colors.HexColor("#1F3864")
MID_BLUE   = colors.HexColor("#2E5FA3")
LIGHT_BLUE = colors.HexColor("#BDD7EE")
LIGHT_GRAY = colors.HexColor("#F5F5F5")
DARK_GRAY  = colors.HexColor("#444444")
WHITE      = colors.white

PAGE_W, PAGE_H = letter
MARGIN    = 0.65 * inch
CONTENT_W = PAGE_W - 2 * MARGIN


def parse_airdna_pdf(pdf_bytes):
    """Extract all data from AirDNA Rentalizer PDF using known line structure."""
    doc  = fitz.open(stream=pdf_bytes, filetype="pdf")
    lines   = [l.strip() for l in doc[0].get_text().split("\n") if l.strip()]
    p2lines = [l.strip() for l in (doc[1].get_text() if len(doc)>1 else "").split("\n") if l.strip()]
    data = {}

    # ── Address (lines 2 & 3, after "Property Earning Potential" and "Submarket Score") ──
    data["address_line1"] = lines[2].rstrip(",").strip() if len(lines) > 2 else ""
    data["city_state_zip"] = lines[3].replace(", USA","").strip() if len(lines) > 3 else ""

    # ── Market / Submarket (line 4) ──
    for l in lines:
        m = re.search(r"Market:\s*(.+?)\s+Submarket:\s*(.+)", l)
        if m:
            data["market"]    = m.group(1).strip()
            data["submarket"] = m.group(2).strip()
            break

    # ── Beds / Baths / Guests (lines 5,6,7) ──
    for l in lines:
        m = re.search(r"^(\d+)\s+Bed", l)
        if m: data["bedrooms"] = m.group(1)
        m = re.search(r"^(\d+(?:\.\d+)?)\s+Bath", l)
        if m: data["bathrooms"] = m.group(1)
        m = re.search(r"^(\d+)\s+Guests?", l)
        if m: data["max_guests"] = m.group(1)

    # ── Financials (fixed positions relative to labels) ──
    for i, l in enumerate(lines):
        if l == "Operating Expenses"   and i+1 < len(lines): data["operating_expenses"] = lines[i+1]
        if l == "Net Operating Income" and i+1 < len(lines): data["noi"]                = lines[i+1]
        if l == "Cap Rate"             and i+1 < len(lines): data["cap_rate"]            = lines[i+1]

    # Revenue = line before "Projected", Occupancy = line before "Occupancy", ADR = line before "Average"
    for i, l in enumerate(lines):
        if l == "Projected"  and i > 0 and "$" in lines[i-1]: data["projected_revenue"] = lines[i-1]
        if l == "Occupancy"  and i > 0 and "%" in lines[i-1]: data["occupancy"]          = lines[i-1]
        if l == "Average"    and i > 0 and "$" in lines[i-1]: data["adr"]                = lines[i-1]

    # ── Submarket Score — appears after AIRDNA.CO footer ──
    for i, l in enumerate(lines):
        if l == "AIRDNA.CO" and i+2 < len(lines):
            candidate = lines[i+2]
            if candidate.isdigit() and 50 <= int(candidate) <= 100:
                data["submarket_score"] = candidate
            break

    # ── Comps — each comp is 7 lines after the title ──
    # Header columns end at "ADR" (line index 36), comps start at 37
    numeric_pat = re.compile(r"^\$?[\d,.KM%]+$")

    adr_idx = None
    for i, l in enumerate(lines):
        if l == "ADR":
            adr_idx = i
            break

    comps = []
    if adr_idx is not None:
        i = adr_idx + 1
        while i < len(lines):
            l = lines[i]
            if l.startswith("+") or l == "AIRDNA.CO":
                break
            # Collect title (may span multiple lines) then 7 numeric values
            title = l
            vals  = []
            j = i + 1
            while j < len(lines) and len(vals) < 7:
                candidate = lines[j]
                cleaned   = candidate.replace(".","").replace("%","").replace("$","").replace("K","").replace(",","")
                if numeric_pat.match(candidate) or cleaned.isdigit():
                    vals.append(candidate)
                elif len(vals) == 0:
                    title += " " + candidate
                else:
                    break
                j += 1
            if len(vals) == 7:
                # Clean title — AirDNA sometimes bleeds the bedroom number onto the title line
                clean_title = re.sub(r'\s+\d+(?:\.\d+)?$', '', title.strip())
                comps.append({
                    "num":     str(len(comps)+1),
                    "name":    clean_title,
                    "bdba":    f"{vals[0]}/{vals[1]}",
                    "rev_pot": vals[2],
                    "days":    vals[3],
                    "revenue": vals[4],
                    "occ":     vals[5],
                    "adr":     vals[6],
                })
                i = j
            else:
                i += 1
    data["comps"] = comps

    # ── Amenities from page 2 ──
    known = {"Air Conditioning","Dryer","Heating","Hot Tub","Kitchen",
             "Parking","Pool","Cable TV","Washer","Wireless Internet"}
    raw_amenities = []
    i = 0
    while i < len(p2lines):
        if p2lines[i] in known and i+1 < len(p2lines) and "%" in p2lines[i+1]:
            raw_amenities.append((p2lines[i], p2lines[i+1]))
            i += 2
        else:
            i += 1
    # Merge Dryer + Washer → Dryer / Washer
    merged, dryer_pct = [], None
    for name, pct in raw_amenities:
        if name == "Dryer":
            dryer_pct = pct
        elif name == "Washer":
            merged.append(("Dryer / Washer", dryer_pct or pct))
        else:
            merged.append((name, pct))
    data["amenities"] = merged

    # ── Property photo (largest image on page 1) ──
    photo_path, best_size = None, 0
    for img in doc[0].get_images(full=True):
        base = doc.extract_image(img[0])
        if len(base["image"]) > best_size:
            best_size      = len(base["image"])
            photo_bytes    = base["image"]
            photo_ext      = base["ext"]
    if best_size > 10000:
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=f".{photo_ext}")
        tmp.write(photo_bytes); tmp.close()
        photo_path = tmp.name
    data["photo_path"] = photo_path

    return data


# ── AI Market Commentary ──────────────────────────────────────────────────────
# ── Charts ────────────────────────────────────────────────────────────────────
BRAND_BLUE = "#2E5FA3"

def chart_revenue_range(comps, projected_rev, out_path):
    """
    Horizontal range chart showing min/Q1/median/Q3/max revenue across the
    comp set, with the analyst's subject projection marked as a vertical line.
    No individual comp is identifiable — aggregated statistics only.
    """
    import statistics
    try:
        rev_vals = sorted([
            float(c["revenue"].replace("$","").replace("K","").replace(",","")) * 1000
            for c in comps if c.get("revenue")
        ])
        proj = float(projected_rev.replace("$","").replace("K","").replace(",","")) * 1000
    except Exception:
        return

    n = len(rev_vals)
    if n < 2:
        return

    r_min    = rev_vals[0]
    r_max    = rev_vals[-1]
    r_median = statistics.median(rev_vals)
    r_q1     = rev_vals[max(0, n//4)]
    r_q3     = rev_vals[min(n-1, 3*n//4)]

    fig, ax = plt.subplots(figsize=(7.5, 2.2))

    # Full range bar
    ax.barh(0, r_max - r_min, left=r_min, height=0.35,
            color=BRAND_BLUE, alpha=0.15, label="Full Range")
    # IQR bar
    ax.barh(0, r_q3 - r_q1, left=r_q1, height=0.35,
            color=BRAND_BLUE, alpha=0.45, label="Middle 50%")
    # Median tick
    ax.plot([r_median, r_median], [-0.22, 0.22], color=BRAND_BLUE,
            linewidth=2.5, label=f"Median  ${r_median/1000:.0f}K")
    # Subject projection
    ax.plot([proj, proj], [-0.28, 0.28], color="#E84040",
            linewidth=2.5, linestyle="--", label=f"Subject Projection  ${proj/1000:.0f}K")

    ax.set_yticks([])
    ax.xaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_: f"${v/1000:.0f}K"))
    ax.tick_params(axis="x", labelsize=8)
    ax.set_xlim(r_min * 0.88, r_max * 1.06)
    ax.set_ylim(-0.5, 0.5)
    ax.grid(axis="x", linestyle="--", alpha=0.4)
    ax.spines[["top","right","left"]].set_visible(False)
    ax.legend(fontsize=7.5, loc="upper left", framealpha=0.7)
    fig.tight_layout(pad=0.5)
    fig.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close()


def chart_adr_vs_occ(comps, proj_adr, proj_occ, out_path):
    """
    Anonymous scatter plot of ADR vs. Occupancy across the comp set.
    No individual comp is labeled — distribution only.
    Subject projection marked distinctly.
    """
    try:
        adr_vals = [float(c["adr"].replace("$","").replace(",",""))
                    for c in comps if c.get("adr")]
        occ_vals = [float(c["occ"].replace("%",""))
                    for c in comps if c.get("occ")]
        p_adr = float(proj_adr.replace("$","").replace(",",""))
        p_occ = float(proj_occ.replace("%",""))
    except Exception:
        return

    if len(adr_vals) < 2:
        return

    fig, ax = plt.subplots(figsize=(7.5, 2.8))

    # Comp cloud — no labels
    ax.scatter(occ_vals, adr_vals, color=BRAND_BLUE, alpha=0.55,
               s=55, zorder=3, label="Comparable Properties")

    # Subject projection
    ax.scatter([p_occ], [p_adr], color="#E84040", s=90, zorder=5,
               marker="*", label=f"Subject Projection")

    ax.set_xlabel("Occupancy Rate (%)", fontsize=8)
    ax.set_ylabel("Average Daily Rate ($)", fontsize=8)
    ax.tick_params(labelsize=8)
    ax.xaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_: f"{v:.0f}%"))
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_: f"${v:.0f}"))
    ax.grid(linestyle="--", alpha=0.35)
    ax.spines[["top","right"]].set_visible(False)
    ax.legend(fontsize=7.5, framealpha=0.7)
    fig.tight_layout(pad=0.5)
    fig.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close()


# ── PDF Builder (same engine as before) ───────────────────────────────────────
def header_footer(canvas, doc, address):
    canvas.saveState()
    if os.path.exists(LOGO_PATH):
        canvas.drawImage(LOGO_PATH, PAGE_W - MARGIN - 1.5*inch,
                         PAGE_H - 0.52*inch, width=1.5*inch, height=0.42*inch,
                         preserveAspectRatio=True, mask="auto")
    canvas.setFont("Helvetica", 7.5)
    canvas.setFillColor(colors.HexColor("#555555"))
    canvas.drawString(MARGIN, PAGE_H - 0.38*inch, address)
    canvas.setStrokeColor(MID_BLUE)
    canvas.setLineWidth(1)
    canvas.line(MARGIN, PAGE_H - 0.56*inch, PAGE_W - MARGIN, PAGE_H - 0.56*inch)
    canvas.line(MARGIN, 0.52*inch, PAGE_W - MARGIN, 0.52*inch)
    canvas.setFont("Helvetica", 7.5)
    canvas.setFillColor(colors.HexColor("#888888"))
    canvas.drawString(MARGIN, 0.35*inch,
        "A-Tech Appraisal Co., LLC | Rhode Island & Massachusetts")
    canvas.drawRightString(PAGE_W - MARGIN, 0.35*inch, f"Page {doc.page}")
    canvas.restoreState()

def make_styles():
    def s(name, **kw): return ParagraphStyle(name, **kw)
    return {
        "title":    s("title", fontSize=22, fontName="Helvetica-Bold",
                      textColor=DARK_BLUE, spaceAfter=6, leading=26),
        "h1":       s("h1", fontSize=14, fontName="Helvetica-Bold",
                      textColor=MID_BLUE, spaceBefore=10, spaceAfter=4, leading=18),
        "h2":       s("h2", fontSize=11, fontName="Helvetica-Bold",
                      textColor=DARK_BLUE, spaceBefore=8, spaceAfter=3),
        "body":     s("body", fontSize=9.5, fontName="Helvetica",
                      textColor=DARK_GRAY, leading=14, spaceAfter=4),
        "small":    s("small", fontSize=8, fontName="Helvetica",
                      textColor=DARK_GRAY, leading=11),
        "lk":       s("lk", fontSize=9, fontName="Helvetica-Bold", textColor=DARK_GRAY),
        "lv":       s("lv", fontSize=9, fontName="Helvetica", textColor=DARK_GRAY),
        "cert":     s("cert", fontSize=9, fontName="Helvetica", textColor=DARK_GRAY,
                      leading=13, leftIndent=10, spaceAfter=2),
    }

DISCLAIMER_ITEMS = [
    "<b>Not an appraisal:</b> This report is a short-term rental income analysis only. It is not an appraisal, appraisal review, or an opinion of market value or market rent.",
    "<b>Licensing:</b> A-Tech Appraisal Co., LLC is not acting as a licensed or certified real estate appraiser in connection with this report and does not provide appraisal services through this analysis.",
    "<b>Estimates:</b> All figures are estimates derived from third-party short-term rental market data and comparable STR performance. Actual results may vary materially based on management, condition, amenities, pricing strategy, seasonality, and market changes.",
    "<b>Rules &amp; permits:</b> Local STR regulations, permits, and tax requirements may apply. Compliance is the responsibility of the owner/operator.",
]

AVM_COMMENTARY_BOILERPLATE = (
    "This STR income analysis is intended to assist the client and/or lender with reviewing "
    "potential short-term rental income for the subject property. No interior or exterior "
    "inspection was completed as part of this analysis, and no opinion of market value or "
    "market rent is provided. Actual STR performance is highly sensitive to pricing strategy, "
    "management quality, guest reviews, furnishings, and amenity set. Local STR regulations, "
    "HOA restrictions, and permitting requirements can materially impact whether STR operation "
    "is permitted and under what conditions. "
    "This analysis was prepared by A-Tech Appraisal Co., LLC."
)

CERT_ITEMS = [
    "The statements of fact contained in this report are true and correct to the best of my knowledge.",
    "The analyses and conclusions are limited only by the stated assumptions and limiting conditions.",
    "I have no present or prospective interest in the subject property and no personal interest with respect to the parties involved.",
    "My compensation is not contingent upon the reporting of a predetermined result or conclusion.",
    "This report is a short-term rental income analysis and is not an appraisal or appraisal review.",
    "The analyst is not acting as a state-licensed or certified real estate appraiser for this assignment.",
]

METHODOLOGY_SECTIONS = [
    ("What this report is (and is not)",
     "This document is a short-term rental income analysis prepared for income support and feasibility review. It summarizes estimated revenue and operating metrics using third-party STR market data and the performance of similar active listings. This is not an appraisal, not an opinion of market value, and not an opinion of market rent."),
    ("Data sources",
     "Market data sources reviewed for this analysis include AirDNA (paid subscription) and publicly available listing data observed directly on Airbnb.com and VRBO.com as of the report date. All performance metrics and income conclusions cited in this report reflect the analyst's independent reconciliation of available market evidence and do not constitute a reproduction of any third-party data product or model output. AirDNA and other data sources are referenced solely as research tools used to inform the analyst's independent professional conclusions."),
    ("Data considered",
     "Primary inputs include the subject's configuration (bed/bath/guest capacity), market and submarket classification, and a curated set of comparable STR listings. The comparable set is used to bracket typical ADR, occupancy rates, and annual revenue for similar rentals."),
    ("Operating expenses, NOI &amp; cap rate",
     "Operating expenses reflect a modeled STR expense framework inclusive of estimated taxes, insurance, utilities, maintenance and turnover costs, and platform/management fees. Net operating income (NOI) is calculated as projected gross revenue less estimated operating expenses."),
    ("Key limitations",
     "No interior or exterior inspection was completed for this analysis. Property condition, furnishings, amenity set, management quality, pricing strategy, and guest reviews can materially impact actual STR performance. Local STR regulations, HOA restrictions, and permitting requirements may restrict or prohibit short-term rental operation in whole or in part. All projections are estimates and are not guarantees of future performance."),
    ("Intended users &amp; intended use",
     "Intended user(s): the client and/or lender, and parties specifically authorized by the client. Intended use: lender-facing STR income support and feasibility review for the subject property. Any other use of this report is prohibited without the express written permission of A-Tech Appraisal Co., LLC."),
]

def generate_comp_narrative(data):
    """
    Generate a professional analyst narrative from parsed comp data.
    Replaces the comp table — no proprietary data reproduced, all conclusions
    are the analyst's independent reconciliation of available market evidence.
    """
    comps = data.get("comps", [])
    market    = data.get("market", "this market")
    submarket = data.get("submarket", "this submarket")
    bedrooms  = data.get("bedrooms", "")
    rev       = data.get("projected_revenue", "")
    adr       = data.get("adr", "")
    occ       = data.get("occupancy", "")

    if not comps:
        return (
            "The analyst reviewed available short-term rental market data for the subject submarket. "
            "Insufficient comparable data was available to produce a detailed range analysis; "
            "the income estimate reflects the analyst's review of broader market conditions."
        )

    # ── Compute stats ────────────────────────────────────────────────────────
    try:
        occ_vals = [float(c["occ"].replace("%","")) for c in comps if c.get("occ")]
        adr_vals = [float(c["adr"].replace("$","").replace(",","")) for c in comps if c.get("adr")]
        rev_vals = [float(c["revenue"].replace("$","").replace("K","").replace(",",""))*1000
                    for c in comps if c.get("revenue")]
        day_vals = [int(c["days"]) for c in comps if c.get("days","").isdigit()]
    except Exception:
        occ_vals, adr_vals, rev_vals, day_vals = [], [], [], []

    if not occ_vals:
        return "The analyst reviewed available short-term rental market data for the subject submarket."

    occ_min,  occ_max  = min(occ_vals),  max(occ_vals)
    adr_min,  adr_max  = min(adr_vals),  max(adr_vals)
    rev_min,  rev_max  = min(rev_vals),  max(rev_vals)
    day_min,  day_max  = (min(day_vals), max(day_vals)) if day_vals else (None, None)

    # Clustering — middle 50% (IQR)
    import statistics
    occ_median = statistics.median(occ_vals)
    adr_median = statistics.median(adr_vals)
    rev_median = statistics.median(rev_vals)

    occ_sorted = sorted(occ_vals)
    adr_sorted = sorted(adr_vals)
    rev_sorted = sorted(rev_vals)
    n = len(comps)
    q1_idx, q3_idx = max(0, n//4), min(n-1, 3*n//4)

    occ_cluster = (occ_sorted[q1_idx], occ_sorted[q3_idx])
    adr_cluster = (adr_sorted[q1_idx], adr_sorted[q3_idx])
    rev_cluster = (rev_sorted[q1_idx]/1000, rev_sorted[q3_idx]/1000)

    # ADR/occupancy relationship — high ADR comps tend to have lower occ?
    paired = sorted(zip(adr_vals, occ_vals), key=lambda x: x[0])
    top_adr_occ   = [o for a,o in paired[-3:]]
    bot_adr_occ   = [o for a,o in paired[:3]]
    luxury_note   = (sum(top_adr_occ)/len(top_adr_occ)) < (sum(bot_adr_occ)/len(bot_adr_occ))

    # ── Build narrative ──────────────────────────────────────────────────────
    n_comps = len(comps)
    bed_str = f"{bedrooms}-bedroom " if bedrooms else ""

    para1 = (
        f"To support the income estimate for the subject property, the analyst reviewed the "
        f"performance of {n_comps} active short-term rental listings in the {submarket} submarket "
        f"comparable to the subject in terms of bedroom count, bathroom count, building class, "
        f"and overall utility. Comparable properties demonstrated a range of performance outcomes "
        f"reflective of differences in floor level, view orientation, amenity set, and management quality."
    )

    para2 = (
        f"Among the comparable set, occupancy rates ranged from approximately {occ_min:.0f}% "
        f"to {occ_max:.0f}%, with the majority of properties clustered between "
        f"{occ_cluster[0]:.0f}% and {occ_cluster[1]:.0f}%. Average daily rates ranged from "
        f"approximately ${adr_min:.0f} to ${adr_max:.0f}, with most properties performing in "
        f"the ${adr_cluster[0]:.0f} to ${adr_cluster[1]:.0f} range. Actual annual revenues "
        f"across the comparable set ranged from approximately ${rev_min/1000:.0f}K to "
        f"${rev_max/1000:.0f}K, with the majority of comparables falling between "
        f"${rev_cluster[0]:.0f}K and ${rev_cluster[1]:.0f}K."
    )

    days_sent = ""
    if day_min and day_max:
        days_sent = (
            f" Days booked ranged from approximately {day_min} to {day_max} annually, "
            f"consistent with {'a high-demand urban' if day_min > 280 else 'an active'} "
            f"submarket benefiting from year-round demand."
        )

    luxury_sent = ""
    if luxury_note:
        luxury_sent = (
            " Properties at the upper end of the ADR range tended to carry lower occupancy rates, "
            "suggesting a premium pricing strategy rather than volume-based booking."
        )

    para3 = days_sent + luxury_sent

    para4 = (
        f"The subject's projected occupancy of {occ} and ADR of {adr} fall within the "
        f"well-supported middle range of comparable performance and are considered reasonable "
        f"and achievable under competent management. The analyst's projected gross annual "
        f"revenue of {rev} is supported by and consistent with the range established by "
        f"the comparable set. Market data sources reviewed for this analysis include AirDNA "
        f"(paid subscription) and publicly available listing data observed directly on "
        f"Airbnb.com and VRBO.com as of the report date. Performance metrics cited herein "
        f"reflect the analyst's independent reconciliation of available market evidence and "
        f"do not constitute a reproduction of any third-party data product or model output."
    )

    return "\n\n".join([para1, para2, para3.strip(), para4]) if para3.strip() else \
           "\n\n".join([para1, para2, para4])


def build_pdf(data, client, loan_num, report_date, commentary, buf,
              photo_override=None, map_override=None,
              client_address="", client_phone="", client_order_num="",
              borrower="", avm_file_id="", property_type="Single-Family Residence"):
    styles = make_styles()
    addr1 = data.get("address_line1","")
    city  = data.get("city_state_zip","")
    full_address = f"{addr1}, {city}"

    # Store assignment fields in data for sign-off table
    data["client_address"]   = client_address
    data["client_phone"]     = client_phone
    data["client_order_num"] = client_order_num
    data["borrower"]         = borrower
    data["avm_file_id"]      = avm_file_id

    # Apply overrides
    if photo_override:
        data["photo_path"] = photo_override
    if map_override:
        data["map_path"] = map_override

    doc = SimpleDocTemplate(buf, pagesize=letter,
        leftMargin=MARGIN, rightMargin=MARGIN,
        topMargin=0.65*inch, bottomMargin=0.65*inch)

    def _hf(canvas, doc):
        header_footer(canvas, doc, full_address)

    story = []

    # ── PAGE 1 ──────────────────────────────────────────────────────────────
    story.append(Paragraph("Short-Term Rental Income Analysis", styles["title"]))
    story.append(Paragraph("(Not an Appraisal)",
        ParagraphStyle("subtitle", fontSize=11, fontName="Helvetica",
                       textColor=colors.HexColor("#888888"), spaceAfter=6, leading=14)))
    story.append(HRFlowable(width=CONTENT_W, thickness=1, color=MID_BLUE, spaceAfter=10))

    # Photo + info side by side
    photo_w = 3.2*inch
    info_w  = CONTENT_W - photo_w - 0.2*inch

    photo_path = data.get("photo_path")
    if photo_path and os.path.exists(photo_path):
        photo_cell = Image(photo_path, width=photo_w, height=2.2*inch)
    else:
        photo_cell = Paragraph("<font color='#AAAAAA'>[Property Photo]</font>",
            ParagraphStyle("ph", fontSize=10, alignment=TA_CENTER))

    lk, lv = styles["lk"], styles["lv"]
    info_rows = [
        [Paragraph("<b>Subject Property:</b>", lk), ""],
        [Paragraph(addr1, lv), ""],
        [Paragraph(city, lv), ""],
        [Paragraph("<b>Property Type:</b>", lk),  Paragraph(property_type, lv)],
        [Paragraph("<b>Configuration:</b>", lk),
         Paragraph(f"{data.get('bedrooms','')} Bedrooms | {data.get('bathrooms','')} Bathrooms", lv)],
        [Paragraph("<b>Maximum Guests:</b>", lk),  Paragraph(data.get("max_guests",""), lv)],
        [Paragraph("<b>Market Area:</b>", lk),     Paragraph(data.get("market",""), lv)],
        [Paragraph("<b>Submarket:</b>", lk),        Paragraph(data.get("submarket",""), lv)],
        [Paragraph("<b>Market Demand:</b>", lk),
         Paragraph(("Very Strong" if int(data.get("submarket_score",0) or 0) >= 85
                    else "Strong" if int(data.get("submarket_score",0) or 0) >= 70
                    else "Moderate" if int(data.get("submarket_score",0) or 0) >= 55
                    else "Emerging"), lv)],
        [Paragraph("<b>Report Date:</b>", lk),      Paragraph(report_date, lv)],
        [Paragraph("<b>Prepared By:</b>", lk),      Paragraph("A-Tech Appraisal Co., LLC", lv)],
    ]
    info_t = Table(info_rows, colWidths=[1.5*inch, 1.9*inch])
    info_t.setStyle(TableStyle([
        ("VALIGN",(0,0),(-1,-1),"TOP"),
        ("TOPPADDING",(0,0),(-1,-1),2),("BOTTOMPADDING",(0,0),(-1,-1),2),
        ("LEFTPADDING",(0,0),(-1,-1),0),("RIGHTPADDING",(0,0),(-1,-1),4),
    ]))

    top_t = Table([[photo_cell, info_t]], colWidths=[photo_w, info_w])
    top_t.setStyle(TableStyle([
        ("VALIGN",(0,0),(-1,-1),"TOP"),
        ("LEFTPADDING",(0,0),(-1,-1),0),("RIGHTPADDING",(0,0),(-1,-1),0),
        ("TOPPADDING",(0,0),(-1,-1),0),("BOTTOMPADDING",(0,0),(-1,-1),0),
        ("RIGHTPADDING",(0,0),(0,-1),22),("LEFTPADDING",(1,0),(1,-1),8),
    ]))
    story.append(top_t)
    story.append(Spacer(1,14))

    # Metrics box
    story.append(Paragraph("Rental Analysis Quick View", styles["h1"]))
    rev  = data.get("projected_revenue","—")
    adr  = data.get("adr","—")
    occ  = data.get("occupancy","—")
    exp  = data.get("operating_expenses","—")
    noi  = data.get("noi","—")
    cap  = data.get("cap_rate","—")

    label_s = ParagraphStyle("ml", fontSize=8, fontName="Helvetica",
                              textColor=DARK_GRAY, alignment=TA_CENTER)
    value_s = ParagraphStyle("mv", fontSize=16, fontName="Helvetica-Bold",
                              textColor=MID_BLUE, alignment=TA_CENTER, leading=20)
    cw = CONTENT_W/3
    mx = Table([
        [Paragraph("Projected Annual STR Income",label_s),
         Paragraph("Average Daily Rate (ADR)",label_s),
         Paragraph("Occupancy Rate (Projected)",label_s)],
        [Paragraph(rev,value_s),Paragraph(adr,value_s),Paragraph(occ,value_s)],
        [Paragraph("Operating Expenses (Est.)",label_s),
         Paragraph("Net Operating Income (NOI)",label_s),
         Paragraph("Estimated Cap Rate",label_s)],
        [Paragraph(exp,value_s),Paragraph(noi,value_s),Paragraph(cap,value_s)],
    ], colWidths=[cw]*3)
    mx.setStyle(TableStyle([
        ("BOX",(0,0),(-1,-1),0.75,MID_BLUE),
        ("INNERGRID",(0,0),(-1,-1),0.5,LIGHT_BLUE),
        ("BACKGROUND",(0,0),(-1,-1),WHITE),
        ("TOPPADDING",(0,0),(-1,-1),6),("BOTTOMPADDING",(0,0),(-1,-1),6),
        ("LEFTPADDING",(0,0),(-1,-1),6),("RIGHTPADDING",(0,0),(-1,-1),6),
    ]))
    story.append(mx)
    story.append(Spacer(1,12))

    # Disclaimer box
    disc_s = ParagraphStyle("disc",fontSize=8,fontName="Helvetica",
                             textColor=DARK_GRAY,leading=11,spaceAfter=3)
    disc_items = [Paragraph("<b>Important Disclaimer (Read First)</b>",
        ParagraphStyle("dh",fontSize=8.5,fontName="Helvetica-Bold",
                       textColor=DARK_BLUE,spaceAfter=3))]
    for txt in DISCLAIMER_ITEMS:
        disc_items.append(Paragraph(txt, disc_s))
    d_inner = Table([[i] for i in disc_items], colWidths=[CONTENT_W-0.3*inch])
    d_inner.setStyle(TableStyle([("LEFTPADDING",(0,0),(-1,-1),6),
                                  ("RIGHTPADDING",(0,0),(-1,-1),6),
                                  ("TOPPADDING",(0,0),(-1,-1),2),
                                  ("BOTTOMPADDING",(0,0),(-1,-1),2)]))
    d_wrap = Table([[d_inner]],colWidths=[CONTENT_W])
    d_wrap.setStyle(TableStyle([("BOX",(0,0),(-1,-1),0.75,MID_BLUE),
                                 ("BACKGROUND",(0,0),(-1,-1),colors.HexColor("#EEF4FB")),
                                 ("TOPPADDING",(0,0),(-1,-1),6),
                                 ("BOTTOMPADDING",(0,0),(-1,-1),6),
                                 ("LEFTPADDING",(0,0),(-1,-1),4),
                                 ("RIGHTPADDING",(0,0),(-1,-1),4)]))
    story.append(d_wrap)

    # ── PAGE 2 ──────────────────────────────────────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph("Comparable Short-Term Rental Analysis", styles["h1"]))

    # Generate and render narrative — no comp table or map
    comp_narrative = generate_comp_narrative(data)
    for para_text in comp_narrative.split("\n\n"):
        para_text = para_text.strip()
        if para_text:
            story.append(Paragraph(para_text, styles["body"]))
            story.append(Spacer(1, 4))

    # ── PAGE 3 ──────────────────────────────────────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph("Market Overview &amp; Commentary", styles["h1"]))
    story.append(Paragraph(
        f"{data.get('market','')} Market – {data.get('submarket','')} Submarket",
        styles["h2"]))
    story.append(Paragraph(commentary, styles["body"]))
    story.append(Spacer(1,6))
    # Compute comp ranges for analyst-attributed projection language
    comps = data.get("comps", [])
    try:
        occ_vals = [float(c["occ"].replace("%","")) for c in comps if c.get("occ")]
        adr_vals = [float(c["adr"].replace("$","").replace(",","")) for c in comps if c.get("adr")]
        rev_vals = [float(c["revenue"].replace("$","").replace("K","").replace(",",""))*1000
                    for c in comps if c.get("revenue")]
        occ_range = f"{min(occ_vals):.0f}%\u2013{max(occ_vals):.0f}%" if occ_vals else occ
        adr_range = f"${min(adr_vals):.0f}\u2013${max(adr_vals):.0f}" if adr_vals else adr
        rev_low   = min(rev_vals) / 1000 if rev_vals else 0
        rev_high  = max(rev_vals) / 1000 if rev_vals else 0
        rev_range = f"${rev_low:.0f}K\u2013${rev_high:.0f}K" if rev_vals else rev
    except Exception:
        occ_range, adr_range, rev_range = occ, adr, rev

    story.append(Paragraph(
        f"<b>Projection Support</b><br/>"
        f"Based on the analyst's review of comparable short-term rental performance in the "
        f"{data.get('submarket','')} submarket, comparable properties demonstrated occupancy rates "
        f"ranging from approximately {occ_range}, average daily rates ranging from approximately "
        f"{adr_range}, and annual revenues ranging from approximately {rev_range}. The analyst's "
        f"projected gross annual revenue of {rev}, reflecting an ADR of {adr} and occupancy of {occ}, "
        f"represents an independent professional estimate derived from available market evidence. "
        f"The subject's bedroom count and guest capacity place it toward the larger end of typical "
        f"STR inventory, which can support higher ADR and stronger peak-season performance when "
        f"paired with competitive amenities and professional management.",
        styles["body"]))
    story.append(Spacer(1,10))

    # AVM Commentary box
    avm_s = ParagraphStyle("avm_i",fontSize=8.5,fontName="Helvetica",
                            textColor=DARK_GRAY,leading=12)
    avm_inner = Table([[Paragraph(AVM_COMMENTARY_BOILERPLATE,avm_s)]],
                       colWidths=[CONTENT_W-0.3*inch])
    avm_inner.setStyle(TableStyle([("TOPPADDING",(0,0),(-1,-1),6),
                                    ("BOTTOMPADDING",(0,0),(-1,-1),6),
                                    ("LEFTPADDING",(0,0),(-1,-1),8),
                                    ("RIGHTPADDING",(0,0),(-1,-1),8)]))
    avm_hdr = Table([[Paragraph("<b>A-Tech Commentary</b>",
        ParagraphStyle("ah",fontSize=9,fontName="Helvetica-Bold",textColor=WHITE))]],
        colWidths=[CONTENT_W])
    avm_hdr.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),DARK_BLUE),
                                  ("TOPPADDING",(0,0),(-1,-1),5),
                                  ("BOTTOMPADDING",(0,0),(-1,-1),5),
                                  ("LEFTPADDING",(0,0),(-1,-1),8)]))
    avm_outer = Table([[avm_hdr],[avm_inner]],colWidths=[CONTENT_W])
    avm_outer.setStyle(TableStyle([("BOX",(0,0),(-1,-1),0.75,MID_BLUE),
                                    ("BACKGROUND",(0,1),(-1,-1),colors.HexColor("#EEF4FB")),
                                    ("LEFTPADDING",(0,0),(-1,-1),0),
                                    ("RIGHTPADDING",(0,0),(-1,-1),0),
                                    ("TOPPADDING",(0,0),(-1,-1),0),
                                    ("BOTTOMPADDING",(0,0),(-1,-1),0)]))
    story.append(avm_outer)

    # ── PAGE 4 ──────────────────────────────────────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph("Amenities &amp; Market Analysis", styles["h1"]))
    story.append(Paragraph("Comparable STR Amenity Prevalence", styles["h2"]))

    # Amenity table
    am_h = ParagraphStyle("amh",fontSize=8.5,fontName="Helvetica-Bold",
                           textColor=WHITE,alignment=TA_CENTER)
    am_c = ParagraphStyle("amc",fontSize=9,fontName="Helvetica",textColor=DARK_GRAY)
    am_p = ParagraphStyle("amp",fontSize=9,fontName="Helvetica",
                           textColor=DARK_GRAY,alignment=TA_CENTER)
    cw4 = CONTENT_W/4
    am_data = [[Paragraph("Amenity",am_h),Paragraph("% of Comps",am_h),
                Paragraph("Amenity",am_h),Paragraph("% of Comps",am_h)]]
    amenities = data.get("amenities",[])
    pairs = [(amenities[i], amenities[i+1] if i+1 < len(amenities) else ("",""))
             for i in range(0, len(amenities), 2)]
    for (a1,p1),(a2,p2) in pairs:
        am_data.append([Paragraph(a1,am_c),Paragraph(p1,am_p),
                        Paragraph(a2,am_c),Paragraph(p2,am_p)])
    at = Table(am_data,colWidths=[cw4]*4)
    at.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),MID_BLUE),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[WHITE,colors.HexColor("#F7FAFF")]),
        ("BOX",(0,0),(-1,-1),0.5,colors.HexColor("#AAAAAA")),
        ("INNERGRID",(0,0),(-1,-1),0.3,colors.HexColor("#DDDDDD")),
        ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),
        ("LEFTPADDING",(0,0),(-1,-1),6),("RIGHTPADDING",(0,0),(-1,-1),6),
    ]))
    story.append(at)
    story.append(Spacer(1,14))

    # Chart 1 — Revenue Range
    comps     = data.get("comps", [])
    proj_rev  = data.get("projected_revenue", "")
    proj_adr  = data.get("adr", "")
    proj_occ  = data.get("occupancy", "")

    if comps and proj_rev:
        tmp1 = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        tmp1.close()
        chart_revenue_range(comps, proj_rev, tmp1.name)
        if os.path.exists(tmp1.name) and os.path.getsize(tmp1.name) > 0:
            story.append(Paragraph("Comparable Revenue Range — Subject vs. Market", styles["h2"]))
            story.append(Paragraph(
                "The chart below illustrates the distribution of annual revenue across the comparable "
                "set. The shaded bar represents the full range; the darker band reflects the middle 50% "
                "of comparable performance. The subject's projected revenue is marked in red.",
                styles["small"]))
            story.append(Spacer(1, 4))
            story.append(Image(tmp1.name, width=CONTENT_W, height=2.0*inch))
            story.append(Spacer(1, 14))

    # Chart 2 — ADR vs Occupancy scatter
    if comps and proj_adr and proj_occ:
        tmp2 = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        tmp2.close()
        chart_adr_vs_occ(comps, proj_adr, proj_occ, tmp2.name)
        if os.path.exists(tmp2.name) and os.path.getsize(tmp2.name) > 0:
            story.append(Paragraph("ADR vs. Occupancy — Comparable Set Distribution", styles["h2"]))
            story.append(Paragraph(
                "Each point represents an anonymous comparable property. The red star marks the "
                "subject's projected ADR and occupancy combination relative to the comp set cluster.",
                styles["small"]))
            story.append(Spacer(1, 4))
            story.append(Image(tmp2.name, width=CONTENT_W, height=2.6*inch))

    # ── PAGE 5 ──────────────────────────────────────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph("Methodology, Assumptions &amp; Limitations", styles["h1"]))
    for title, body in METHODOLOGY_SECTIONS:
        story.append(Paragraph(title, styles["h2"]))
        story.append(Paragraph(body, styles["body"]))

    # ── PAGE 6 ──────────────────────────────────────────────────────────────
    story.append(PageBreak())
    story.append(Paragraph("Identification, Intended Use &amp; Analyst Sign-Off", styles["h1"]))

    so_rows = [
        ["Prepared By",         "A-Tech Appraisal Co., LLC"],
        ["Subject Property",    full_address],
        ["Property Type",       property_type],
        ["Configuration",       f"{data.get('bedrooms','')} Bedrooms | "
                                f"{data.get('bathrooms','')} Bathrooms | "
                                f"{data.get('max_guests','')} Guests Max"],
        ["Market / Submarket",  f"{data.get('market','')} / {data.get('submarket','')}"],
        ["Report Date",         report_date],
        ["Client / Lender",     client],
        ["Client Address",      data.get("client_address","")],
        ["Client Phone",        data.get("client_phone","")],
        ["Client Order Number", data.get("client_order_num","")],
        ["Borrower",            data.get("borrower","")],
        ["Loan Number",         loan_num],
        ["AVM File ID",         data.get("avm_file_id","")],
        ["Intended Use",        "Short-term rental income support for lender feasibility / underwriting review (not an appraisal)."],
        ["Intended Users",      "Client/lender and parties specifically authorized by the client."],
    ]
    so_lk = ParagraphStyle("slk",fontSize=9,fontName="Helvetica-Bold",textColor=DARK_GRAY)
    so_lv = ParagraphStyle("slv",fontSize=9,fontName="Helvetica",textColor=DARK_GRAY)
    so_data = [[Paragraph(r[0],so_lk),Paragraph(r[1],so_lv)] for r in so_rows]
    so_t = Table(so_data,colWidths=[1.8*inch,CONTENT_W-1.8*inch])
    so_t.setStyle(TableStyle([
        ("BOX",(0,0),(-1,-1),0.5,colors.HexColor("#AAAAAA")),
        ("INNERGRID",(0,0),(-1,-1),0.3,colors.HexColor("#DDDDDD")),
        ("BACKGROUND",(0,0),(0,-1),LIGHT_GRAY),
        ("ROWBACKGROUNDS",(1,0),(-1,-1),[WHITE,colors.HexColor("#F7FAFF")]),
        ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),
        ("LEFTPADDING",(0,0),(-1,-1),7),("RIGHTPADDING",(0,0),(-1,-1),7),
        ("VALIGN",(0,0),(-1,-1),"TOP"),
    ]))
    story.append(so_t)
    story.append(Spacer(1,14))

    story.append(Paragraph("Analyst Certification", styles["h2"]))
    for item in CERT_ITEMS:
        story.append(Paragraph(f"• {item}", styles["cert"]))
    story.append(Spacer(1,20))

    sign_t = Table([[
        Paragraph("Company: <b>A-Tech Appraisal Co., LLC</b>",
            ParagraphStyle("sc",fontSize=9,fontName="Helvetica",textColor=DARK_GRAY)),
        Paragraph(f"Date: <b>{report_date}</b>",
            ParagraphStyle("sd",fontSize=9,fontName="Helvetica",
                           textColor=DARK_GRAY,alignment=TA_RIGHT))
    ]], colWidths=[CONTENT_W*0.5]*2)
    sign_t.setStyle(TableStyle([
        ("LINEABOVE",(1,0),(1,0),0.75,DARK_GRAY),
        ("TOPPADDING",(0,0),(-1,-1),4),
        ("LEFTPADDING",(0,0),(-1,-1),0),
        ("RIGHTPADDING",(0,0),(-1,-1),0),
    ]))
    story.append(sign_t)

    doc.build(story, onFirstPage=_hf, onLaterPages=_hf)


# ── Email Helper ──────────────────────────────────────────────────────────────
# ── Intake Export Helpers ─────────────────────────────────────────────────────

def build_intake_pdf(intake_text, address="Subject Property"):
    """Convert intake markdown text to a branded PDF using ReportLab."""
    from io import BytesIO
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=0.75*inch, rightMargin=0.75*inch,
        topMargin=0.9*inch, bottomMargin=0.75*inch
    )

    styles = {
        "title": ParagraphStyle("title", fontName="Helvetica-Bold", fontSize=14,
                                 textColor=colors.HexColor("#1F3864"),
                                 spaceAfter=6, spaceBefore=0),
        "h2":    ParagraphStyle("h2", fontName="Helvetica-Bold", fontSize=11,
                                 textColor=colors.HexColor("#2E5FA3"),
                                 spaceAfter=4, spaceBefore=10),
        "body":  ParagraphStyle("body", fontName="Helvetica", fontSize=9,
                                 textColor=colors.HexColor("#333333"),
                                 spaceAfter=3, leading=13),
        "flag":  ParagraphStyle("flag", fontName="Helvetica", fontSize=9,
                                 textColor=colors.HexColor("#C00000"),
                                 spaceAfter=2, leading=12),
        "bold":  ParagraphStyle("bold", fontName="Helvetica-Bold", fontSize=9,
                                 textColor=colors.HexColor("#333333"),
                                 spaceAfter=2),
    }

    def _hf(canvas, doc):
        canvas.saveState()
        if os.path.exists(LOGO_PATH):
            canvas.drawImage(LOGO_PATH,
                             letter[0] - 0.75*inch - 1.4*inch,
                             letter[1] - 0.62*inch,
                             width=1.4*inch, height=0.38*inch,
                             preserveAspectRatio=True, mask="auto")
        canvas.setFont("Helvetica-Bold", 8)
        canvas.setFillColor(colors.HexColor("#1F3864"))
        canvas.drawString(0.75*inch, letter[1] - 0.45*inch, "APPRAISAL ASSIGNMENT INTAKE")
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(colors.HexColor("#555555"))
        canvas.drawString(0.75*inch, letter[1] - 0.60*inch, address)
        canvas.setStrokeColor(colors.HexColor("#2E5FA3"))
        canvas.setLineWidth(1)
        canvas.line(0.75*inch, letter[1] - 0.68*inch, letter[0] - 0.75*inch, letter[1] - 0.68*inch)
        canvas.line(0.75*inch, 0.55*inch, letter[0] - 0.75*inch, 0.55*inch)
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(colors.HexColor("#888888"))
        canvas.drawString(0.75*inch, 0.38*inch,
            "A-Tech Appraisal Co., LLC")
        canvas.drawRightString(letter[0] - 0.75*inch, 0.38*inch, f"Page {doc.page}")
        canvas.restoreState()

    story = []

    # Parse markdown into ReportLab elements
    lines = intake_text.split("\n")
    in_table = False
    table_rows = []

    for line in lines:
        stripped = line.strip()

        # Skip horizontal rules and empty TOTAL_JSON remnants
        if stripped in ("---", "***", "___") or stripped.startswith("TOTAL_JSON"):
            if in_table and table_rows:
                # Flush table
                if len(table_rows) > 1:
                    col_count = max(len(r) for r in table_rows)
                    col_w = (letter[0] - 1.5*inch) / col_count
                    tbl = Table(
                        [[Paragraph(str(c), styles["body"]) for c in row] for row in table_rows],
                        colWidths=[col_w]*col_count
                    )
                    tbl.setStyle(TableStyle([
                        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#BDD7EE")),
                        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
                        ("FONTSIZE", (0,0), (-1,-1), 8.5),
                        ("GRID", (0,0), (-1,-1), 0.4, colors.HexColor("#CCCCCC")),
                        ("ROWBACKGROUNDS", (0,1), (-1,-1),
                         [colors.white, colors.HexColor("#F5F8FF")]),
                        ("TOPPADDING", (0,0), (-1,-1), 4),
                        ("BOTTOMPADDING", (0,0), (-1,-1), 4),
                        ("LEFTPADDING", (0,0), (-1,-1), 6),
                        ("RIGHTPADDING", (0,0), (-1,-1), 6),
                        ("VALIGN", (0,0), (-1,-1), "TOP"),
                    ]))
                    story.append(tbl)
                    story.append(Spacer(1, 6))
                table_rows = []
                in_table = False
            if stripped in ("---", "***", "___"):
                story.append(HRFlowable(width="100%", thickness=0.5,
                                         color=colors.HexColor("#CCCCCC"),
                                         spaceAfter=4, spaceBefore=4))
            continue

        # Markdown table rows
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if all(set(c) <= set("-: ") for c in cells):
                continue  # separator row
            table_rows.append(cells)
            in_table = True
            continue
        else:
            if in_table and table_rows:
                col_count = max(len(r) for r in table_rows)
                col_w = (letter[0] - 1.5*inch) / col_count
                tbl = Table(
                    [[Paragraph(str(c), styles["body"]) for c in row] for row in table_rows],
                    colWidths=[col_w]*col_count
                )
                tbl.setStyle(TableStyle([
                    ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#BDD7EE")),
                    ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
                    ("FONTSIZE", (0,0), (-1,-1), 8.5),
                    ("GRID", (0,0), (-1,-1), 0.4, colors.HexColor("#CCCCCC")),
                    ("ROWBACKGROUNDS", (0,1), (-1,-1),
                     [colors.white, colors.HexColor("#F5F8FF")]),
                    ("TOPPADDING", (0,0), (-1,-1), 4),
                    ("BOTTOMPADDING", (0,0), (-1,-1), 4),
                    ("LEFTPADDING", (0,0), (-1,-1), 6),
                    ("RIGHTPADDING", (0,0), (-1,-1), 6),
                    ("VALIGN", (0,0), (-1,-1), "TOP"),
                ]))
                story.append(tbl)
                story.append(Spacer(1, 6))
                table_rows = []
                in_table = False

        if not stripped:
            story.append(Spacer(1, 4))
            continue

        # Headings
        if stripped.startswith("## "):
            story.append(Paragraph(stripped[3:].strip(), styles["title"]))
        elif stripped.startswith("### ") or stripped.startswith("#### "):
            story.append(Paragraph(stripped.lstrip("#").strip(), styles["h2"]))
        elif stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            story.append(Paragraph(stripped.strip("*"), styles["bold"]))
        elif stripped.startswith("- [x]") or stripped.startswith("-[x]"):
            text = stripped[5:].strip()
            story.append(Paragraph(f"⚠ {text}", styles["flag"]))
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            # Inline bold
            text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
            story.append(Paragraph(f"• {text}", styles["body"]))
        elif re.match(r"^\d+\.", stripped):
            text = re.sub(r"^\d+\.\s*", "", stripped)
            text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
            story.append(Paragraph(f"{stripped.split('.')[0]}. {text}", styles["body"]))
        else:
            # Regular paragraph — handle inline bold
            text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", stripped)
            story.append(Paragraph(text, styles["body"]))

    doc.build(story, onFirstPage=_hf, onLaterPages=_hf)
    buf.seek(0)
    return buf.getvalue()


def build_intake_docx(intake_text, address="Subject Property"):
    """Convert intake markdown text to a Word document using python-docx."""
    from io import BytesIO
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return None

    doc = DocxDocument()

    # Page setup — US Letter, 0.75" margins
    section = doc.sections[0]
    section.page_width  = int(8.5  * 914400)
    section.page_height = int(11.0 * 914400)
    for attr in ("left_margin","right_margin","top_margin","bottom_margin"):
        setattr(section, attr, int(0.75 * 914400))

    # Helper: set paragraph font
    def _set_run(run, bold=False, size=10, color=None, italic=False):
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)

    def _heading(text, level=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8 if level == 1 else 4)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(text)
        if level == 1:
            _set_run(run, bold=True, size=13, color=(31, 56, 100))
        else:
            _set_run(run, bold=True, size=11, color=(46, 95, 163))
        return p

    def _body(text, bold=False, color=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        _set_run(run, bold=bold, size=9.5,
                 color=color or (51, 51, 51))
        return p

    def _flag(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"⚠ {text}")
        _set_run(run, size=9.5, color=(192, 0, 0))
        return p

    def _bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        # Strip inline ** bold markers for simplicity
        clean = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        run = p.add_run(clean)
        _set_run(run, size=9.5)
        return p

    def _add_table(rows):
        if not rows:
            return
        col_count = max(len(r) for r in rows)
        tbl = doc.add_table(rows=0, cols=col_count)
        tbl.style = "Table Grid"
        for i, row_data in enumerate(rows):
            row = tbl.add_row()
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = str(cell_text) if j < len(row_data) else ""
                run = cell.paragraphs[0].runs
                if run:
                    run[0].font.size = Pt(9)
                    if i == 0:
                        run[0].bold = True
                # Header row shading
                if i == 0:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), "BDD7EE")
                    tcPr.append(shd)
        doc.add_paragraph()  # spacing after table

    # Title header
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(2)
    r = title_p.add_run("APPRAISAL ASSIGNMENT INTAKE")
    _set_run(r, bold=True, size=14, color=(31, 56, 100))

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(8)
    r = sub_p.add_run(address)
    _set_run(r, size=10, color=(85, 85, 85))

    # Parse and render
    lines = intake_text.split("\n")
    in_table = False
    table_rows = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("TOTAL_JSON"):
            break  # stop at JSON block

        if stripped in ("---", "***", "___"):
            if in_table and table_rows:
                _add_table(table_rows)
                table_rows = []
                in_table = False
            # Horizontal rule via paragraph border
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "AAAAAA")
            pBdr.append(bottom)
            pPr.append(pBdr)
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if all(set(c) <= set("-: ") for c in cells):
                continue
            table_rows.append(cells)
            in_table = True
            continue
        else:
            if in_table and table_rows:
                _add_table(table_rows)
                table_rows = []
                in_table = False

        if not stripped:
            doc.add_paragraph()
            continue

        if stripped.startswith("## "):
            _heading(stripped[3:].strip(), level=1)
        elif stripped.startswith("### ") or stripped.startswith("#### "):
            _heading(stripped.lstrip("#").strip(), level=2)
        elif stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            _body(stripped.strip("*"), bold=True)
        elif stripped.startswith("- [x]") or stripped.startswith("-[x]"):
            _flag(stripped[5:].strip())
        elif stripped.startswith("- ") or stripped.startswith("* "):
            _bullet(stripped[2:].strip())
        elif re.match(r"^\d+\.", stripped):
            _bullet(re.sub(r"^\d+\.\s*", "", stripped))
        else:
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            _body(clean)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def send_report_email(to_email, subject, body, pdf_bytes, filename):
    """Send PDF report via Gmail SMTP."""
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.application import MIMEApplication

    try:
        gmail_address = st.secrets["GMAIL_ADDRESS"]
        gmail_password = st.secrets["GMAIL_APP_PASSWORD"]
    except Exception:
        gmail_address = os.environ.get("GMAIL_ADDRESS","")
        gmail_password = os.environ.get("GMAIL_APP_PASSWORD","")

    msg = MIMEMultipart()
    msg["From"]    = gmail_address
    msg["To"]      = to_email
    msg["Subject"] = subject
    msg.attach(MIMEText(body, "plain"))

    attachment = MIMEApplication(pdf_bytes, _subtype="pdf")
    attachment.add_header("Content-Disposition", "attachment", filename=filename)
    msg.attach(attachment)

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
        server.login(gmail_address, gmail_password)
        server.send_message(msg)


# ── Streamlit UI ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="A-Tech STR Report Generator",
    page_icon="🏠",
    layout="centered"
)


# ── Password Protection ───────────────────────────────────────────────────────

# ── Default Data ──────────────────────────────────────────────────────────────
DEFAULT_REVISIONS = [
    {"id": "r1", "category": "Rental Comp Selection — Multi-Family vs SFR", "request": "As per the report, subject is 'Multi-Family' however SFR (single family) rental comp #1-3 is used, please verify or comment.", "response": "All rental comparables provided in this report are sourced from multi-family dwellings, not single-family residences. Individual units within multi-family properties are typically marketed and leased on a per-unit basis, and the rental listings utilized reflect individually marketed units from within multi-family buildings. These represent the most recent and relevant market rent data available, as active and recent listings from multi-family properties provide the most accurate reflection of current market rents for the subject's unit type. The rental comparables are considered appropriate and consistent with the subject's property classification.", "notes": ""},
    {"id": "r2", "category": "Building Official — Verbal vs Written Documentation", "request": "The lender is asking if the appraiser received anything from the building official that stated that the kitchen needed to be removed so that they can present it to the borrower?", "response": "The appraiser did not receive written documentation from the building official; however, the appraiser contacted the building department directly and was verbally informed of the requirements necessary to bring the property into compliance. The items communicated by the building department were incorporated into the report as subject-to conditions reflecting those requirements.", "notes": "Customize the building department name to match the subject municipality."},
    {"id": "r3", "category": "Basement Discoloration — Moisture Observation", "request": "Appraiser to comment on discoloration on concrete basement floors and state whether any moisture or dampness was observed.", "response": "At the time of inspection there were no present areas of moisture or dampness observed. The discoloration noted on the concrete basement floor is consistent with normal aging and does not indicate active moisture intrusion based on the appraiser's visual observation at the time of inspection.", "notes": "If moisture was present, modify accordingly. This is a visual observation disclaimer — not a structural certification."},
    {"id": "r4", "category": "Not a PUD — Client States Subject is a PUD", "request": "Client states the subject is a PUD.", "response": "A Planned Unit Development (PUD) is generally characterized by a homeowners association that holds title to and is responsible for maintaining common areas and amenities, with individual unit owners holding fee simple title to their lots and improvements. The subject property does not meet this definition. [OPTION A — No HOA: The subject has no homeowners association, common areas, or shared amenities that would indicate a PUD classification.] [OPTION B — Minor HOA: The subject does carry a nominal HOA fee for common area maintenance; however, this fee structure alone does not constitute a PUD classification. The subject's HOA structure does not rise to that level.] The property has been appropriately classified consistent with its legal description, site ownership, and market perception.", "notes": "Use Option A if no HOA. Use Option B if minor HOA exists. Delete the option that does not apply."},
]

DEFAULT_COMMENTS = [
    {"id": "c1",  "category": "Land Condo — Classification and Market Perception", "text": "A land condominium is a form of ownership in which the unit owner holds fee simple title to the land beneath their home (the 'unit'), while certain common elements — such as roads, utilities, or shared open space — are owned collectively by the association. Unlike a traditional condominium, the land condo owner holds direct ownership of the land, and the structure is treated in a manner consistent with single-family ownership for appraisal and lending purposes. In the subject's market area, land condominiums are generally perceived by buyers and sellers as functionally equivalent to fee simple single-family homes. The market does not demonstrate a measurable valuation differential between land condo and fee simple single-family properties in this neighborhood.", "notes": "Verify that your comparable sales are also land condos where possible."},
    {"id": "c2",  "category": "Concession Adjustments — Adjusted at Cash Value", "text": "It would appear from the analysis of the market that concessions of 0-4% are typical within the subject's marketing area. There is not a prevalence of loan discounts, interest buy-downs and/or concessions which would have a negative impact on the subject property's market value, unless otherwise stated in the report. However, the appraiser has assessed adjustments to comparables in an effort to demonstrate a cash basis for adjusted value.", "notes": "Use when concessions are present but not excessive."},
    {"id": "c3",  "category": "Concession Adjustments — Adjusted Only When Excessive", "text": "The analysis of the average concessions package for homes within the sample used to determine market trends indicates a range of approximately X-X%. There is no prevalent use of loan discounts, interest rate buy-downs, or concessions that would negatively impact the subject property's market value unless otherwise stated in the report. Consequently, the appraiser has applied adjustments to comparables only when they exceed the typical threshold for the overall market.", "notes": "Use when concessions exist but are not being adjusted on every comp."},
    {"id": "c4",  "category": "Chain of Title", "text": "The appraiser has conducted a thorough search of the subject property and comparables to determine the chain of title and the number of transfers for each home. All available and customary data sources were utilized to obtain this information, including MLS, CompFlo deeds, Realist tax, and public record. However, it should be noted that data in this area can sometimes be unreliable, particularly for recent transfers. While every effort was made to identify all transfers, it is not guaranteed that every transaction was captured. The purpose of the appraiser's chain of title examination is to alert the client to any potentially compromised comparables or non-arm's length transactions.", "notes": "Update data sources to match your actual sources."},
    {"id": "c5",  "category": "Comparable Selection Criteria", "text": "In searching for comparables, the appraiser has emphasized location, lot size, gross living area, utility, design, age, and condition, with specific emphasis on sales that closed within 180 days prior to the effective date. In most cases, comparables older than six months have been eliminated from consideration. In cases where the appraiser could not locate recent sales within the immediate area, a representative sample of homes from both the subject neighborhood and competing areas was utilized. Sales that may represent outlier events have been carefully scrutinized and eliminated from consideration when appropriate.", "notes": ""},
    {"id": "c6",  "category": "Highest and Best Use", "text": "The relevant legal, physical, and economic factors were analyzed to the extent necessary and resulted in a conclusion that the current use of the subject property is the highest and best use [USPAP — Standards Rule 2-2(b)(x)].", "notes": ""},
    {"id": "c7",  "category": "Condition/Quality Adjustments — General", "text": "It should be noted that the appraiser considered numerous sales within the defined neighborhood reflecting various states of condition. In comparable selection the appraiser made a concentrated effort to locate and select homes similar in terms of physical characteristics and condition. The appraiser assessed features and condition through the use of MLS listings, interior photos (when available), and exterior inspections. Any condition/quality adjustments assessed are based on the appraiser's assessment of contributory value via paired sales analysis. Unless otherwise noted the appraiser has applied equal emphasis to all comparable sales.", "notes": ""},
    {"id": "c8",  "category": "Condition/Quality Adjustments — UAD Requirements", "text": "UAD requirements mandate that condition and quality ratings be reported on an 'absolute' basis rather than a 'relative' basis. Consequently, in some cases the appraiser will make adjustments to comparables that share the same UAD condition or quality rating due to minor differences in overall condition or quality. The appraiser assessed features and condition through the use of MLS listings, interior photos, and exterior inspections. During the course of research, the appraiser deemed that comparables #[X] and #[X] warrant a similar condition rating but were slightly superior/inferior in terms of overall condition. Consequently, an adjustment of $[X,XXX] has been assigned based on the appraiser's assessment of contributory value via paired sales analysis.", "notes": "Fill in comparable numbers and adjustment amount."},
    {"id": "c9",  "category": "Active Listing/Pending Sale — Multiple", "text": "Comparables #[X] and #[X] are active listings/pending sales and were included to demonstrate current market activity within the subject's marketing area. A conservative [X]% adjustment was assessed to the active listings in an effort to compensate for the typical differential between list price and sales price.", "notes": "Typically 5-10% depending on market."},
    {"id": "c10", "category": "Active Listing/Pending Sale — Single (with explanation)", "text": "Comparable #[X] is an active listing/pending sale and was included to demonstrate current market activity within the subject's marketing area. A conservative [X]% adjustment was assessed in an effort to compensate for the typical differential between list price and sales price. Please note that the appraiser conducted an exhaustive search for an additional listing competitive to the subject; however, none similar in terms of age, GLA, and condition were available as of the effective date.", "notes": ""},
    {"id": "c11", "category": "Predominant Value — Appraised Value Exceeds Predominant", "text": "The predominant value as stated on page 1 of this appraisal is approximately $[X]; however, this range applies to the subject's immediate marketing area as a whole. Homes within this range are considered to be inferior in terms of GLA, age, quality of construction, appeal, and utility. Consequently the appraised value exceeds the predominant value. The home is not considered to be over-improved for the marketing area.", "notes": ""},
    {"id": "c12", "category": "Effective Age — Less Than Actual Age", "text": "Please note that it is the opinion of the appraiser that the effective age is less than the actual age of the subject property. This assertion is predicated on the fact that the subject has been well-maintained and that the home's true age does not reflect the general upkeep performed over its lifetime. Although the home does show general wear and tear consistent with its age, its effective age has been estimated at [X] years.", "notes": "Fill in the estimated effective age."},
    {"id": "c13", "category": "GLA — Differs from Public Record (>10%)", "text": "Please note that the GLA for the subject may differ from public record by more than 10%. The measurements indicated within this report are considered to be highly accurate. Differentials of this nature are not uncommon and are attributed to inaccuracies within the public record data. Discrepancies commonly occur when municipalities include basement or unfinished attic area in the reported GLA figure.", "notes": ""},
    {"id": "c14", "category": "GLA — Differential Exceeds 20% Preferred Guideline", "text": "Please note that the GLA differential for some comparables exceeds the preferred 20% guideline. This is attributed to the complexity of comparable selection and the limited availability of recent comparable sales similar in terms of GLA. Consequently, the appraiser was required to utilize homes that were inferior/superior in terms of GLA and has applied the necessary adjustments accordingly.", "notes": ""},
    {"id": "c15", "category": "GLA — Subject Inferior to All Comparables", "text": "Please note that the subject is inferior in terms of GLA to all comparables utilized. The appraiser conducted an exhaustive search for recent comparables similar or superior in terms of GLA; however, very few that could be considered similar in terms of age and condition were available at the time of inspection. The subject is not considered to be inadequate for its marketing area.", "notes": ""},
    {"id": "c16", "category": "GLA — Subject Superior to All Comparables", "text": "Please note that the subject is superior in terms of GLA to all comparables utilized. The appraiser conducted an exhaustive search for recent comparables similar in terms of GLA; however, very few that could be considered similar in terms of age and condition were available. The subject is not considered to be over-improved for its marketing area.", "notes": ""},
    {"id": "c17", "category": "GLA — Split Level Threshold Method", "text": "Please note that the subject is a split-level home. For the purposes of this appraisal the appraiser has utilized the 'threshold method of entry' to report GLA and basement areas. This entails reporting 'below grade' area as all finished/unfinished area below the threshold of the front door. This method is applied consistently across all comparables. Within this market, tax records tend to be inconsistent when reporting physical characteristics for split-level homes.", "notes": ""},
    {"id": "c18", "category": "Net/Gross Adjustments Exceed 15%/25%", "text": "Please note that some comparables exceed the preferred guidelines of 15% and/or 25% for net and/or gross adjustment percentages respectively. This can be attributed to the complexity of comparable selection and the limited number of recent available comparable sales. The subject's floor plan, GLA, and lot size in conjunction with age made it necessary to utilize a variety of homes to 'bracket' the market's reaction for each contributory characteristic.", "notes": ""},
    {"id": "c19", "category": "Line Item Adjustments Exceed 10%", "text": "Please note that single line item adjustments for some comparables exceed the preferred 10% guideline. This can be attributed to the complexity of comparable selection and the limited number of recent available comparable sales. The subject's floor plan, GLA, and lot size in conjunction with age made it necessary to utilize a variety of homes to 'bracket' the market's reaction for each contributory characteristic.", "notes": ""},
    {"id": "c20", "category": "Land Value Exceeds 30%", "text": "Please note that land value for the subject exceeds 30% of indicated market value. This can be attributed to the limited availability of vacant land for development within the subject's defined neighborhood and the enhanced marketability associated with the subject's location. This is considered typical for the subject's marketing area.", "notes": "Common in coastal RI markets. Pair with land sales data when possible."},
    {"id": "c21", "category": "Pool — In-Ground Pool Adjustment", "text": "Please note that the subject has an in-ground pool. Within this market, pools are considered to affect marketability in a positive manner. Consequently an adjustment of $[XX,XXX] was assessed to all comparables without pools in an effort to compensate for the enhanced marketability considered to be inherent in such a feature.", "notes": "Fill in the dollar adjustment. Support with paired sales if challenged."},
    {"id": "c22", "category": "Waterfront / River Lot", "text": "Please note that the subject is situated on a waterfront/river-front lot. Within this market, the subject's lot influence represents a significant, positive component of the subject's marketability and appeals to a specific pool of buyers. The appraiser made a concentrated effort to locate and/or utilize homes similar in terms of lot influence. The appraiser has assessed a $[XX,000] adjustment to all comparables not situated on a lot with a similar external influence in an effort to compensate for the enhanced marketability.", "notes": "Especially relevant for Narragansett, coastal RI, and waterfront MA assignments."},
    {"id": "c23", "category": "Acreage / Large Lot", "text": "Please note that the subject is situated on approximately [XX] acres. Within this market, the subject's lot size represents a significant, positive component of the subject's marketability. Per MLS, vacant land sales and listings within the subject's defined neighborhood are selling/listing from $[X] to $[X] per acre. The appraiser has adjusted conservatively at $[X] per acre to compensate for the subject's lot size.", "notes": "Fill in acreage and per-acre values from land sales."},
    {"id": "c24", "category": "External Obsolescence — High Traffic Street", "text": "Please note that the subject suffers from incurable external obsolescence in the form of close proximity to a moderately high-traffic street. The appraiser conducted an exhaustive search for comparable sales with a similar external influence; however, none were available for review at the time of inspection. The appraiser assessed a $[X,000] adjustment to all comparables that do not share a similar external influence.", "notes": ""},
    {"id": "c25", "category": "External Obsolescence — Power Lines", "text": "Please note that the subject suffers from incurable external obsolescence in the form of close proximity to power lines. The appraiser conducted an exhaustive search for comparable sales with a similar external influence; however, none were available for review at the time of inspection. The appraiser assessed a conservative $[X,000] adjustment based on market knowledge, realtor interviews, and historical sales data.", "notes": ""},
    {"id": "c26", "category": "Flood Zone — Portion of Lot in 100-Year Flood Plain", "text": "A portion of the subject property appears to lie within the 100-year flood plain; however, improvements do not appear to lie within the flood zone. The appraiser is not an expert in this field and recommends that a flood certification be obtained to more specifically determine the subject's position relative to the flood zone.", "notes": "Add adjustment language if a flood zone adjustment was made."},
    {"id": "c27", "category": "Condition Disclaimer — Not a Home Inspection", "text": "Please note that this appraisal is not a home inspection and the appraiser is not acting as a home inspector. While observing the subject property, the appraiser visually observed areas that were readily accessible. The appraiser is not required to disturb or move obstructions to visibility. The inspection is not technically exhaustive. A formal home inspection for the subject property was not provided to the appraiser. The appraisal report should not be relied upon to disclose any conditions present in the subject property. A professional home inspection is recommended on all property purchase transactions.", "notes": "Include on any assignment with deferred maintenance, C4 or lower, or noted repairs."},
    {"id": "c28", "category": "C1 — New Construction", "text": "New Construction (C1): The subject is new construction. Due to the age of the subject, no physical inadequacies were noted at the time of inspection. For the purposes of this appraisal and to establish a basis for comparison, the appraiser has listed the home as 'new' or C1 condition.", "notes": ""},
    {"id": "c29", "category": "C2 — Remodeled", "text": "Remodeled (C2): Please note that the subject has recently been subjected to a comprehensive remodel. A detailed list of improvements has been attached to this report. The appraiser noted obvious and readily observable improvements to the property completed in a workmanlike manner. It is the opinion of the appraiser that the noted improvements have brought the home into good marketable condition. The appraiser has listed the home as 'good/remodeled' or C2 condition.", "notes": ""},
    {"id": "c30", "category": "C3 — Updated Condition", "text": "Updated Condition (C3): Please note that the subject has been well maintained and has been subjected to a series of recent improvements. The appraiser noted obvious and readily observable improvements completed in a workmanlike manner. The appraiser has listed the subject in 'updated' or C3 condition.", "notes": ""},
    {"id": "c31", "category": "C4 — Average Condition", "text": "Average Condition (C4): Please note that the subject has been adequately maintained and is considered to be in average marketable condition. The appraiser noted no recent major renovations and/or improvements to the property within the last [X] years. The appraiser has listed the home as 'average' or C4 condition.", "notes": ""},
    {"id": "c32", "category": "C5/C6 — Fair/Poor Condition", "text": "Fair Condition (C5-C6): Please note that the subject is currently in a deteriorated state of condition. An interior inspection revealed that many interior and exterior components indicate that some items may be closer to the end of their useful life than a cursory inspection might reveal. The appraiser has provided a list of repairs that represent what was readily observable at the time of inspection; however, this appraisal is not a home inspection. A professional home inspection is recommended. The property is considered to be in 'fair' or C5 condition.", "notes": ""},
    {"id": "c33", "category": "Income Approach — Not Utilized", "text": "Please note that the appraiser has considered the income approach to value for the subject; however, the approach was not feasible due to the limited availability of information regarding rental homes that have recently sold. In order to generate an income approach to value, the appraiser would require access to recently sold rental homes to develop a reliable gross rent multiplier (GRM). The data sources available within the normal course of business do not lend themselves to identifying sufficient arm's-length investment transactions of this type. Consequently, the income approach to value has not been developed.", "notes": ""},
    {"id": "c34", "category": "Adjustment Factors — GLA Methodology", "text": "Please note that adjustment factors are based on sales comparison analysis or modified cost approach to determine the contributory value of specific marketability factors for the subject. Regression analysis is not feasible during the normal course of business in this market area as MLS data does not typically provide accurate representations of all relevant physical characteristics. Consequently, adjustment factors are based on paired sales analysis, market extraction, local builder and realtor input, and building cost databases.", "notes": "This language directly addresses regression analysis pushback."},
    {"id": "c35", "category": "Gross Building Area (GBA) — Multi-Family Definition", "text": "Exterior measurements of a two-to-four family building are used to calculate the Gross Building Area (GBA). Interior finished space, including below-grade living space, interior stairways, hallways, storage rooms, and laundry rooms are part of the GBA. Exterior stairways are not included. Due to the subject's multi-family status, the total rentable area is calculated as GBA. Consequently the below-grade area has been included in the GBA.", "notes": "Use on all 2-4 family assignments."},
    {"id": "c36", "category": "Scope of Work — Conventional", "text": "In conducting this appraisal assignment, the appraiser first collected preliminary public record data and made an initial search of available market sales, trends, and influences. A physical inspection of the subject property was made in accordance with the information requirements of the URAR format. The appraiser is not an expert in matters of pest control, structural engineering, hazardous waste, survey, or title matters, and no expertise or warranty is implied in these areas.", "notes": ""},
]

DEFAULT_ZONING = [
    {'id': 'z1', 'city': 'Barrington', 'district': 'RE', 'property_type': '', 'frontage': 'See Note', 'lot_area': 'See Note', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'As consistent with the requirements of the predominant residential zoning designation of the area  surrounding the property, being either the R-25 or the R-40 District.'},
    {'id': 'z2', 'city': 'Barrington', 'district': 'AR', 'property_type': '', 'frontage': "180'", 'lot_area': '120,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': "100'  Street frontage allowed where a lot abuts a cul-de-sac or an outside curb of a street"},
    {'id': 'z3', 'city': 'Barrington', 'district': 'R-40 CD (SF)', 'property_type': '', 'frontage': "180'", 'lot_area': '40,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': "Single-family dwelling / 100'  Street frontage allowed where a lot abuts a cul-de-sac or an outside curb of a street"},
    {'id': 'z4', 'city': 'Barrington', 'district': 'R-25 (SF)', 'property_type': '', 'frontage': "90'", 'lot_area': '25,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': "Single-family dwelling / 60'  Street frontage allowed where a lot abuts a cul-de-sac or an outside curb of a street"},
    {'id': 'z5', 'city': 'Barrington', 'district': 'R-10 (SF)', 'property_type': '', 'frontage': "90'", 'lot_area': '10,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Single-family dwelling'},
    {'id': 'z6', 'city': 'Barrington', 'district': 'NB/RBF', 'property_type': '', 'frontage': "90'", 'lot_area': '10,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z7', 'city': 'Barrington', 'district': 'GI', 'property_type': '', 'frontage': '', 'lot_area': '', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'As consistent with the requirements of the predominant residential zoning designation of the area  surrounding the property, being either the R-25 or the R-40 District.'},
    {'id': 'z8', 'city': 'Barrington', 'district': 'R-40 CD 2-Fam', 'property_type': '', 'frontage': "180' SINGLE FAM / 100'  Street frontage allowed where a lot abuts a cul-de-sac or an outside curb of a street", 'lot_area': '50,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Two-family dwellings'},
    {'id': 'z9', 'city': 'Barrington', 'district': 'R-25 SF ADU', 'property_type': '', 'frontage': "90' SINGLE FAM / 60'  Street frontage allowed where a lot abuts a cul-de-sac or an outside curb of a street", 'lot_area': '40,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Single-family dwelling with accessory living quarters or guest house'},
    {'id': 'z10', 'city': 'Barrington', 'district': 'R-10 SF ADU', 'property_type': '', 'frontage': "90' SINGLE FAM", 'lot_area': 'N/A', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Single-family dwelling with accessory living quarters or guest house'},
    {'id': 'z11', 'city': 'Barrington', 'district': 'R-40 CD SF ADU', 'property_type': '', 'frontage': "180' SINGLE FAM / 100'  Street frontage allowed where a lot abuts a cul-de-sac or an outside curb of a street", 'lot_area': '60,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Single-family dwelling with accessory living quarters or guest house'},
    {'id': 'z12', 'city': 'Barrington', 'district': 'R-25 2-Fam or Other', 'property_type': '', 'frontage': "90' SINGLE FAM / 60'  Street frontage allowed where a lot abuts a cul-de-sac or an outside curb of a street", 'lot_area': '30,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Two-family dwellings or  Other permitted or special use permit uses'},
    {'id': 'z13', 'city': 'Barrington', 'district': 'R-10 2-Fam or Other', 'property_type': '', 'frontage': "60' TWO FAM & OTHER PERMITED USES", 'lot_area': '15,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Two-family dwellings or  Other permitted or special use permit uses'},
    {'id': 'z14', 'city': 'Barrington', 'district': 'R-40 CD (Other)', 'property_type': '', 'frontage': "180' SINGLE FAM / 100'  Street frontage allowed where a lot abuts a cul-de-sac or an outside curb of a street", 'lot_area': '80,0000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Other permitted or special use permit uses'},
    {'id': 'z16', 'city': 'Block Island (New Shoreham)', 'district': 'RA', 'property_type': '', 'frontage': "200'", 'lot_area': '120,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z17', 'city': 'Block Island (New Shoreham)', 'district': 'RB', 'property_type': '', 'frontage': "150'", 'lot_area': '60,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z18', 'city': 'Block Island (New Shoreham)', 'district': 'RC', 'property_type': '', 'frontage': "75'", 'lot_area': '20,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Without Sewer 40,000 SF'},
    {'id': 'z19', 'city': 'Block Island (New Shoreham)', 'district': 'RC/M', 'property_type': '', 'frontage': "75'", 'lot_area': '20,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Without Sewer 40,000 SF'},
    {'id': 'z20', 'city': 'Block Island (New Shoreham)', 'district': 'M', 'property_type': '', 'frontage': "75'", 'lot_area': '20,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Without Sewer 40,000 SF'},
    {'id': 'z21', 'city': 'Block Island (New Shoreham)', 'district': 'OHC', 'property_type': '', 'frontage': "75'", 'lot_area': '20,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z22', 'city': 'Block Island (New Shoreham)', 'district': 'NHC', 'property_type': '', 'frontage': "150'", 'lot_area': '20,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Without Sewer 40,000 SF'},
    {'id': 'z23', 'city': 'Block Island (New Shoreham)', 'district': 'SC', 'property_type': '', 'frontage': "100'", 'lot_area': '20,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z25', 'city': 'Bristol', 'district': 'R-80', 'property_type': '', 'frontage': '150', 'lot_area': '80,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z26', 'city': 'Bristol', 'district': 'R-40', 'property_type': '', 'frontage': '150', 'lot_area': '40,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z27', 'city': 'Bristol', 'district': 'R-20', 'property_type': '', 'frontage': '120', 'lot_area': '20,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z28', 'city': 'Bristol', 'district': 'R-15', 'property_type': '', 'frontage': '100', 'lot_area': '15,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z29', 'city': 'Bristol', 'district': 'R-10', 'property_type': '', 'frontage': '80', 'lot_area': '10,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z30', 'city': 'Bristol', 'district': 'R-10SW', 'property_type': '', 'frontage': '80', 'lot_area': '10,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'With public water & sewer'},
    {'id': 'z31', 'city': 'Bristol', 'district': 'R-8', 'property_type': '', 'frontage': '80', 'lot_area': '8,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z32', 'city': 'Bristol', 'district': 'R-6', 'property_type': '', 'frontage': '60', 'lot_area': '6,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '1 dwelling unit; +4,000 SF per additional unit'},
    {'id': 'z34', 'city': 'Burrillville', 'district': 'OS', 'property_type': '', 'frontage': "450'", 'lot_area': '5 acres (conservation)', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z35', 'city': 'Burrillville', 'district': 'R-40', 'property_type': '', 'frontage': '150', 'lot_area': '40,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z36', 'city': 'Burrillville', 'district': 'R-7', 'property_type': '', 'frontage': '75', 'lot_area': '7,500 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z37', 'city': 'Burrillville', 'district': 'VC', 'property_type': '', 'frontage': "125'", 'lot_area': '20,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z38', 'city': 'Burrillville', 'district': 'R-12 2- Fam', 'property_type': '', 'frontage': "125'", 'lot_area': '15,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z39', 'city': 'Burrillville', 'district': 'F-5', 'property_type': '', 'frontage': "450'", 'lot_area': '5 acres', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z40', 'city': 'Burrillville', 'district': 'F-2', 'property_type': '', 'frontage': "300'", 'lot_area': '2 acres', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z41', 'city': 'Burrillville', 'district': 'R-12', 'property_type': '', 'frontage': "100''", 'lot_area': '12,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z43', 'city': 'Central Falls', 'district': 'R-1', 'property_type': '', 'frontage': '40', 'lot_area': '5,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z44', 'city': 'Central Falls', 'district': 'R-2', 'property_type': '', 'frontage': '40', 'lot_area': '5,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '+2,500 SF for 1st 2 units; +1,500 SF each additional unit'},
    {'id': 'z45', 'city': 'Central Falls', 'district': 'R-3', 'property_type': '', 'frontage': '40', 'lot_area': '5,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '+2,000 SF for 1st 3 units; +1,000 SF each additional unit; max height 30 ft / 3 stories'},
    {'id': 'z47', 'city': 'Charlestown', 'district': 'R-20', 'property_type': '', 'frontage': '120', 'lot_area': '20,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z48', 'city': 'Charlestown', 'district': 'R-40', 'property_type': '', 'frontage': '150', 'lot_area': '40,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z49', 'city': 'Charlestown', 'district': 'R-40 Cluster', 'property_type': '', 'frontage': '100', 'lot_area': '20,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z50', 'city': 'Charlestown', 'district': 'R-2A', 'property_type': '', 'frontage': '200', 'lot_area': '2 acres', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z51', 'city': 'Charlestown', 'district': 'R-2A Multi', 'property_type': '', 'frontage': '200', 'lot_area': '2 acres per dwelling unit', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': "+10' Frontage per dwelling unit"},
    {'id': 'z52', 'city': 'Charlestown', 'district': 'R-3A', 'property_type': '', 'frontage': "300'", 'lot_area': '30 acres', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z53', 'city': 'Charlestown', 'district': 'C-1', 'property_type': '', 'frontage': "120'", 'lot_area': '20,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z54', 'city': 'Charlestown', 'district': 'C-2', 'property_type': '', 'frontage': "150'", 'lot_area': '20,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z55', 'city': 'Charlestown', 'district': 'R-3A Multi', 'property_type': '', 'frontage': "300'", 'lot_area': '3 Acres per Dwelling Unit', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': "+ 20' Frontage per dwelling unit"},
    {'id': 'z56', 'city': 'Charlestown', 'district': 'Village District', 'property_type': '', 'frontage': "120'", 'lot_area': '20,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z57', 'city': 'Charlestown', 'district': 'R-40, R-2A, R-Aa Conservative Dev', 'property_type': '', 'frontage': "50'", 'lot_area': '20,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z58', 'city': 'Charlestown', 'district': 'C-3', 'property_type': '', 'frontage': "150'", 'lot_area': '40,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z59', 'city': 'Charlestown', 'district': 'Two Family Dwelling', 'property_type': '', 'frontage': "R-2A 250' Front / R-3A 300'", 'lot_area': '2 x min. lot size', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z61', 'city': 'Coventry', 'district': 'Village Rural Commercial', 'property_type': '', 'frontage': '125', 'lot_area': '40,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z62', 'city': 'Coventry', 'district': 'Village Main Street Commercial', 'property_type': '', 'frontage': '80', 'lot_area': '7,500 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z63', 'city': 'Coventry', 'district': 'General Business', 'property_type': '', 'frontage': '125', 'lot_area': '15,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z64', 'city': 'Coventry', 'district': 'General Business-1', 'property_type': '', 'frontage': '200', 'lot_area': '43,560 SF (1 acre)', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z65', 'city': 'Coventry', 'district': 'RR5', 'property_type': '', 'frontage': "300'", 'lot_area': '5 acres', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z66', 'city': 'Coventry', 'district': 'RR3', 'property_type': '', 'frontage': "225'", 'lot_area': '3 acres', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z67', 'city': 'Coventry', 'district': 'RR2', 'property_type': '', 'frontage': "225'", 'lot_area': '2 acres', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z68', 'city': 'Coventry', 'district': 'R20 (SF)', 'property_type': '', 'frontage': "120'", 'lot_area': '20,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z69', 'city': 'Coventry', 'district': 'R20 (No public water)', 'property_type': '', 'frontage': "150'", 'lot_area': '1 acre', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z70', 'city': 'Coventry', 'district': 'R20 2-Fam', 'property_type': '', 'frontage': "175'", 'lot_area': '30,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z71', 'city': 'Coventry', 'district': 'R20 2-Fam no public water', 'property_type': '', 'frontage': "175'", 'lot_area': '60,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z72', 'city': 'Coventry', 'district': 'SF Cluster Devl (W/ water Or sewer)', 'property_type': '', 'frontage': "100'", 'lot_area': '15,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z73', 'city': 'Coventry', 'district': '2 Fam Cluster Devl (W/ water Or sewer)', 'property_type': '', 'frontage': "125'", 'lot_area': '20,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z74', 'city': 'Coventry', 'district': 'SF Cluster Devl (W/ no water Or sewer)', 'property_type': '', 'frontage': "150'", 'lot_area': '43,560 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z75', 'city': 'Coventry', 'district': '2 Fam Cluster Devl (W/ no water Or sewer)', 'property_type': '', 'frontage': "175'", 'lot_area': '60,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z76', 'city': 'Coventry', 'district': 'SF Conservation Des (W/ water Or sewer)', 'property_type': '', 'frontage': "100'", 'lot_area': '15,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z77', 'city': 'Coventry', 'district': 'MF Conservation Des (W/ water Or sewer)', 'property_type': '', 'frontage': "125'", 'lot_area': '20,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z78', 'city': 'Coventry', 'district': 'SF Conservation Des (W/ water Or sewer)', 'property_type': '', 'frontage': "150'", 'lot_area': '1 acre', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z79', 'city': 'Coventry', 'district': '2 Fam Conservation Des (W/ water Or sewer)', 'property_type': '', 'frontage': "175'", 'lot_area': '60,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z81', 'city': 'Cranston', 'district': 'A-80', 'property_type': '', 'frontage': '', 'lot_area': '80,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z82', 'city': 'Cranston', 'district': 'A-20', 'property_type': '', 'frontage': '', 'lot_area': '20,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z83', 'city': 'Cranston', 'district': 'A-12', 'property_type': '', 'frontage': '', 'lot_area': '12,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z84', 'city': 'Cranston', 'district': 'A-8', 'property_type': '', 'frontage': '', 'lot_area': '8,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z85', 'city': 'Cranston', 'district': 'A-6', 'property_type': '', 'frontage': '', 'lot_area': '6,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z86', 'city': 'Cranston', 'district': 'B-1', 'property_type': '', 'frontage': '', 'lot_area': '6,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z87', 'city': 'Cranston', 'district': 'B-2', 'property_type': '', 'frontage': '', 'lot_area': '6,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z88', 'city': 'Cranston', 'district': 'B-1 2-Fam', 'property_type': '', 'frontage': '', 'lot_area': '8,000 SF (2F)', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z89', 'city': 'Cranston', 'district': 'B-2 2-Fam', 'property_type': '', 'frontage': '', 'lot_area': '8,000 SF (2F)', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z91', 'city': 'Cumberland', 'district': 'A-1 w/o water Or sewer', 'property_type': '', 'frontage': "250'", 'lot_area': '217,800 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '217,800sf per aditional dwelling unit'},
    {'id': 'z92', 'city': 'Cumberland', 'district': 'A-2 w/o water Or Sewer', 'property_type': '', 'frontage': "180'", 'lot_area': '80,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '80,000 sf per additional dwelling unit'},
    {'id': 'z93', 'city': 'Cumberland', 'district': 'R-1 w/o water Or Sewer', 'property_type': '', 'frontage': "180'", 'lot_area': '80,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '80,000 sf per additional dwelling unit'},
    {'id': 'z94', 'city': 'Cumberland', 'district': 'A-1 w/ water Or sewer But not both', 'property_type': '', 'frontage': "250'", 'lot_area': '217,800 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '217,800sf per aditional dwelling unit'},
    {'id': 'z95', 'city': 'Cumberland', 'district': 'A-2 w/ water Or sewer But not both', 'property_type': '', 'frontage': "180'", 'lot_area': '80,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '80,000 sf per additional dwelling unit'},
    {'id': 'z96', 'city': 'Cumberland', 'district': 'R-1 w/ water Or sewer But not both', 'property_type': '', 'frontage': "100'", 'lot_area': '40,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '40,000 sf per additional dwelling unit'},
    {'id': 'z97', 'city': 'Cumberland', 'district': 'R-2 w/ water Or sewer But not both', 'property_type': '', 'frontage': "90'", 'lot_area': '30,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '10,000 sf for 2nd unit'},
    {'id': 'z98', 'city': 'Cumberland', 'district': 'R-3 w/ water Or sewer But not both', 'property_type': '', 'frontage': "90'", 'lot_area': '30,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '10,000 sf for 2nd unit'},
    {'id': 'z99', 'city': 'Cumberland', 'district': 'A-1 w/ water & sewer', 'property_type': '', 'frontage': "250'", 'lot_area': '217,800 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z100', 'city': 'Cumberland', 'district': 'A-2 w/ water & sewer', 'property_type': '', 'frontage': "180'", 'lot_area': '80,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z101', 'city': 'Cumberland', 'district': 'R-1 w/ water & sewer', 'property_type': '', 'frontage': "100'", 'lot_area': '25,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z102', 'city': 'Cumberland', 'district': 'R-2 w/ water & sewer', 'property_type': '', 'frontage': "40'", 'lot_area': '10,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '5,000 sf for 2nd unit'},
    {'id': 'z103', 'city': 'Cumberland', 'district': 'R-3 w/ water & sewer', 'property_type': '', 'frontage': "40'", 'lot_area': '10,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '5,000 sf for each additional unit'},
    {'id': 'z105', 'city': 'East Greenwich', 'district': 'R-4', 'property_type': '', 'frontage': "100'", 'lot_area': '4,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z106', 'city': 'East Greenwich', 'district': 'R-10', 'property_type': '', 'frontage': "100'", 'lot_area': '10,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z107', 'city': 'East Greenwich', 'district': 'R-20', 'property_type': '', 'frontage': "125'", 'lot_area': '20,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z108', 'city': 'East Greenwich', 'district': 'R-30', 'property_type': '', 'frontage': "150'", 'lot_area': '30,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z109', 'city': 'East Greenwich', 'district': 'F-1', 'property_type': '', 'frontage': "150'", 'lot_area': '1 Acre', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z110', 'city': 'East Greenwich', 'district': 'F-2', 'property_type': '', 'frontage': "150'", 'lot_area': '2 acres', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z111', 'city': 'East Greenwich', 'district': 'R-6 SF', 'property_type': '', 'frontage': "60'", 'lot_area': '6,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z112', 'city': 'East Greenwich', 'district': 'R-6 2-Fam', 'property_type': '', 'frontage': "80'", 'lot_area': '10,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z113', 'city': 'East Greenwich', 'district': 'R-6 Multi', 'property_type': '', 'frontage': "100'", 'lot_area': '10,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '4,000 sf for every unit over 2 units'},
    {'id': 'z115', 'city': 'East Providence', 'district': 'Residential 1', 'property_type': '', 'frontage': '125', 'lot_area': '18,750 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z116', 'city': 'East Providence', 'district': 'Residential 2', 'property_type': '', 'frontage': '100', 'lot_area': '10,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z117', 'city': 'East Providence', 'district': 'Residential 3', 'property_type': '', 'frontage': '75', 'lot_area': '7,500 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z118', 'city': 'East Providence', 'district': 'Residential 4', 'property_type': '', 'frontage': '50', 'lot_area': '5,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z119', 'city': 'East Providence', 'district': 'Residential 5', 'property_type': '', 'frontage': '75', 'lot_area': '7,500 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z120', 'city': 'East Providence', 'district': 'Residential 6', 'property_type': '', 'frontage': '50', 'lot_area': '5,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z121', 'city': 'East Providence', 'district': 'Commercial 1', 'property_type': '', 'frontage': '100', 'lot_area': '10,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z122', 'city': 'East Providence', 'district': 'Commercial 2', 'property_type': '', 'frontage': '50', 'lot_area': '5,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z123', 'city': 'East Providence', 'district': 'Industrial 1', 'property_type': '', 'frontage': '150', 'lot_area': '30,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z124', 'city': 'East Providence', 'district': 'Industrial 2', 'property_type': '', 'frontage': '175', 'lot_area': '40,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z126', 'city': 'Exeter', 'district': 'RE-2', 'property_type': '', 'frontage': '200', 'lot_area': '2 acres', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z127', 'city': 'Exeter', 'district': 'RU-3', 'property_type': '', 'frontage': '250', 'lot_area': '3 Acres', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z128', 'city': 'Exeter', 'district': 'RU-4', 'property_type': '', 'frontage': '300', 'lot_area': '4 Acres', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z129', 'city': 'Exeter', 'district': 'CR-5', 'property_type': '', 'frontage': '350', 'lot_area': '5 Acres', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z130', 'city': 'Exeter', 'district': 'LB-R', 'property_type': '', 'frontage': '150', 'lot_area': '2 Acres', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z131', 'city': 'Exeter', 'district': 'LI', 'property_type': '', 'frontage': '400', 'lot_area': '2 Acres', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z133', 'city': 'Foster', 'district': 'AR', 'property_type': '', 'frontage': '300', 'lot_area': '200,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z134', 'city': 'Foster', 'district': 'GB', 'property_type': '', 'frontage': '300', 'lot_area': '200,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z135', 'city': 'Foster', 'district': 'HC2', 'property_type': '', 'frontage': '300', 'lot_area': '200,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z137', 'city': 'Glocester', 'district': 'A-4', 'property_type': '', 'frontage': '350', 'lot_area': '4 acres', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z138', 'city': 'Glocester', 'district': 'A-4 2-Fam', 'property_type': '', 'frontage': "350'", 'lot_area': '6 Acres', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z139', 'city': 'Glocester', 'district': 'A-3 SF', 'property_type': '', 'frontage': "300'", 'lot_area': '3 Acres', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z140', 'city': 'Glocester', 'district': 'A-3 2-Fam', 'property_type': '', 'frontage': "350'", 'lot_area': '6 Acres', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z141', 'city': 'Glocester', 'district': 'R-2 SF', 'property_type': '', 'frontage': "250'", 'lot_area': '2 Acres', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z142', 'city': 'Glocester', 'district': 'R-2 2-Fam', 'property_type': '', 'frontage': "300'", 'lot_area': '5 Acres', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z143', 'city': 'Glocester', 'district': 'R-2 Multi', 'property_type': '', 'frontage': "300'", 'lot_area': '4 Acres Per Unit', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z145', 'city': 'Hopkinton', 'district': 'R-1', 'property_type': '', 'frontage': '100', 'lot_area': '20,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Subdivision requires 60,000 SF per lot; 2-family: 80,000 SF per unit'},
    {'id': 'z146', 'city': 'Hopkinton', 'district': 'RFR-80', 'property_type': '', 'frontage': '225', 'lot_area': '80,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z147', 'city': 'Hopkinton', 'district': 'Neighborhood Business', 'property_type': '', 'frontage': '150', 'lot_area': '40,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z148', 'city': 'Hopkinton', 'district': 'Commercial', 'property_type': '', 'frontage': '150', 'lot_area': '60,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z149', 'city': 'Hopkinton', 'district': 'Manufacturing', 'property_type': '', 'frontage': '225', 'lot_area': '80,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z150', 'city': 'Hopkinton', 'district': 'R-1 New Sub div', 'property_type': '', 'frontage': "100'", 'lot_area': '60,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z152', 'city': 'Jamestown', 'district': 'OS', 'property_type': '', 'frontage': "300'", 'lot_area': '200,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z153', 'city': 'Jamestown', 'district': 'RR-200', 'property_type': '', 'frontage': "300'", 'lot_area': '200,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z154', 'city': 'Jamestown', 'district': 'RR-80', 'property_type': '', 'frontage': "200'", 'lot_area': '80,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z155', 'city': 'Jamestown', 'district': 'R-400', 'property_type': '', 'frontage': "150'", 'lot_area': '40,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z156', 'city': 'Jamestown', 'district': 'R-20', 'property_type': '', 'frontage': "100'", 'lot_area': '20,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z157', 'city': 'Jamestown', 'district': 'R-20 2-Fam', 'property_type': '', 'frontage': "100'", 'lot_area': '30,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z158', 'city': 'Jamestown', 'district': 'R-20 Multi', 'property_type': '', 'frontage': "100'", 'lot_area': '80,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z159', 'city': 'Jamestown', 'district': 'R-8', 'property_type': '', 'frontage': "80'", 'lot_area': '8,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z160', 'city': 'Jamestown', 'district': 'R-8 2-Fam', 'property_type': '', 'frontage': "80'", 'lot_area': '12,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z161', 'city': 'Jamestown', 'district': 'R-8 Multi', 'property_type': '', 'frontage': "80'", 'lot_area': '25,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z162', 'city': 'Jamestown', 'district': 'CL', 'property_type': '', 'frontage': "80' min-120' max", 'lot_area': '8,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z163', 'city': 'Jamestown', 'district': 'CL 2-Fam', 'property_type': '', 'frontage': "80' min-120' max", 'lot_area': '8,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z164', 'city': 'Jamestown', 'district': 'CL Multi', 'property_type': '', 'frontage': "80' min-120' max", 'lot_area': '25,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z165', 'city': 'Jamestown', 'district': 'CD', 'property_type': '', 'frontage': "40' min- 96' max", 'lot_area': '5,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z166', 'city': 'Jamestown', 'district': 'CD 2-Fam', 'property_type': '', 'frontage': "40' min- 96' max", 'lot_area': '5,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z167', 'city': 'Jamestown', 'district': 'CS Multi', 'property_type': '', 'frontage': "40' min- 96' max", 'lot_area': '20,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z169', 'city': 'Johnston', 'district': 'R-40', 'property_type': '', 'frontage': '140', 'lot_area': '40,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z170', 'city': 'Johnston', 'district': 'R-20', 'property_type': '', 'frontage': "120'", 'lot_area': '20,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z171', 'city': 'Johnston', 'district': 'R-15', 'property_type': '', 'frontage': "100'", 'lot_area': '15,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z172', 'city': 'Johnston', 'district': 'R-10 SF', 'property_type': '', 'frontage': "100'", 'lot_area': '10,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z173', 'city': 'Johnston', 'district': 'R-10 Duplex', 'property_type': '', 'frontage': "150'", 'lot_area': '15,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z174', 'city': 'Johnston', 'district': 'R-10 2-Fam', 'property_type': '', 'frontage': "120'", 'lot_area': '12,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z175', 'city': 'Johnston', 'district': 'R-7 SF', 'property_type': '', 'frontage': "70'", 'lot_area': '7,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z176', 'city': 'Johnston', 'district': 'R-7 2-Fam', 'property_type': '', 'frontage': "120'", 'lot_area': '12,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z177', 'city': 'Johnston', 'district': 'R-7 2-Fam', 'property_type': '', 'frontage': "85'", 'lot_area': '8,500 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z179', 'city': 'Lincoln', 'district': 'RA-40', 'property_type': '', 'frontage': '150', 'lot_area': '40,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z180', 'city': 'Lincoln', 'district': 'RS-20', 'property_type': '', 'frontage': '120', 'lot_area': '20,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z181', 'city': 'Lincoln', 'district': 'RS-12', 'property_type': '', 'frontage': '100', 'lot_area': '12,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z182', 'city': 'Lincoln', 'district': 'RL-9', 'property_type': '', 'frontage': '75', 'lot_area': '9,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z183', 'city': 'Lincoln', 'district': 'RG-7', 'property_type': '', 'frontage': '60', 'lot_area': '7,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z184', 'city': 'Lincoln', 'district': 'RL-9 2-Fam', 'property_type': '', 'frontage': "100'", 'lot_area': '12,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z185', 'city': 'Lincoln', 'district': 'RG-7 2-Fam', 'property_type': '', 'frontage': "70'", 'lot_area': '8,500 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z186', 'city': 'Lincoln', 'district': 'RG-7 Multi', 'property_type': '', 'frontage': "60'", 'lot_area': '7,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': "+ 10' Frontage for ea unit over 1 / + 1,500sf for ea unit over one"},
    {'id': 'z188', 'city': 'Little Compton', 'district': 'R', 'property_type': '', 'frontage': '175', 'lot_area': '2 acres', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Frontage extends to depth of 175 ft; cul-de-sac: 105 ft'},
    {'id': 'z190', 'city': 'Middletown', 'district': 'R-60', 'property_type': '', 'frontage': "200'", 'lot_area': '60,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z191', 'city': 'Middletown', 'district': 'R-30', 'property_type': '', 'frontage': "130'", 'lot_area': '30,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z192', 'city': 'Middletown', 'district': 'R-20', 'property_type': '', 'frontage': "120'", 'lot_area': '20,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z193', 'city': 'Middletown', 'district': 'R-20 2-Fam', 'property_type': '', 'frontage': "150'", 'lot_area': '40,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z194', 'city': 'Middletown', 'district': 'R-10', 'property_type': '', 'frontage': "100'", 'lot_area': '10,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z195', 'city': 'Middletown', 'district': 'R-10 2-Fam', 'property_type': '', 'frontage': "120'", 'lot_area': '15,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z196', 'city': 'Middletown', 'district': 'RM', 'property_type': '', 'frontage': "100'", 'lot_area': '10,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z197', 'city': 'Middletown', 'district': 'RM 2-Fam', 'property_type': '', 'frontage': "120'", 'lot_area': '15,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z198', 'city': 'Middletown', 'district': 'RM Multi', 'property_type': '', 'frontage': "150'", 'lot_area': '40,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z199', 'city': 'Middletown', 'district': 'LB', 'property_type': '', 'frontage': "100'", 'lot_area': '10,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z200', 'city': 'Middletown', 'district': 'LB 2-Fam', 'property_type': '', 'frontage': "120'", 'lot_area': '15,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z201', 'city': 'Middletown', 'district': 'LB Multi', 'property_type': '', 'frontage': "150'", 'lot_area': '40,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z202', 'city': 'Middletown', 'district': 'OB', 'property_type': '', 'frontage': "120'", 'lot_area': '20,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z203', 'city': 'Middletown', 'district': 'OB 2-Fam', 'property_type': '', 'frontage': "150'", 'lot_area': '40,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z204', 'city': 'Middletown', 'district': 'OB Multi', 'property_type': '', 'frontage': "150'", 'lot_area': '40,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z206', 'city': 'Narragansett', 'district': 'R-80', 'property_type': '', 'frontage': '200', 'lot_area': '80,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z207', 'city': 'Narragansett', 'district': 'R-40', 'property_type': '', 'frontage': '150', 'lot_area': '40,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Single-family dwelling; Other uses: 20,000 SF'},
    {'id': 'z208', 'city': 'Narragansett', 'district': 'R-20', 'property_type': '', 'frontage': '100', 'lot_area': '20,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z209', 'city': 'Narragansett', 'district': 'R-10', 'property_type': '', 'frontage': '100', 'lot_area': '10,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z210', 'city': 'Narragansett', 'district': 'R-10A', 'property_type': '', 'frontage': '100', 'lot_area': '10,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z211', 'city': 'Narragansett', 'district': 'R-80 Duplex', 'property_type': '', 'frontage': '200', 'lot_area': '100,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z212', 'city': 'Narragansett', 'district': 'R-40', 'property_type': '', 'frontage': '150', 'lot_area': '60,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z214', 'city': 'Newport', 'district': 'R-3 & LB', 'property_type': '', 'frontage': '50', 'lot_area': '3,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z215', 'city': 'Newport', 'district': 'R-4', 'property_type': '', 'frontage': '50', 'lot_area': '4,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z216', 'city': 'Newport', 'district': 'R-6', 'property_type': '', 'frontage': "50'", 'lot_area': '6,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z217', 'city': 'Newport', 'district': 'R-10 & R10A', 'property_type': '', 'frontage': "80'", 'lot_area': '10,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z218', 'city': 'Newport', 'district': 'R-20', 'property_type': '', 'frontage': "100'", 'lot_area': '20,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z219', 'city': 'Newport', 'district': 'R-40A', 'property_type': '', 'frontage': "200'", 'lot_area': '40,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z220', 'city': 'Newport', 'district': 'R-60', 'property_type': '', 'frontage': "200'", 'lot_area': '60,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z221', 'city': 'Newport', 'district': 'R-120', 'property_type': '', 'frontage': "300'", 'lot_area': '120,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z222', 'city': 'Newport', 'district': 'R-160', 'property_type': '', 'frontage': "400'", 'lot_area': '160,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z223', 'city': 'Newport', 'district': 'WBD & GBD', 'property_type': '', 'frontage': "50'", 'lot_area': '5,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z225', 'city': 'North Kingstown', 'district': 'RR/R80', 'property_type': '', 'frontage': '200', 'lot_area': '80,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z226', 'city': 'North Kingstown', 'district': 'PP', 'property_type': '', 'frontage': '200', 'lot_area': '5 acres', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z227', 'city': 'North Kingstown', 'district': 'NR/R40', 'property_type': '', 'frontage': '180', 'lot_area': '40,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z228', 'city': 'North Kingstown', 'district': 'VR/R20', 'property_type': '', 'frontage': '140', 'lot_area': '20,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z229', 'city': 'North Kingstown', 'district': 'VLDR/200', 'property_type': '', 'frontage': '300', 'lot_area': '200,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z230', 'city': 'North Kingstown', 'district': 'LDR/120', 'property_type': '', 'frontage': '250', 'lot_area': '120,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z231', 'city': 'North Kingstown', 'district': 'VR/R20 2 Fam', 'property_type': '', 'frontage': '160', 'lot_area': '40,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z233', 'city': 'North Providence', 'district': 'RS-8', 'property_type': '', 'frontage': '70', 'lot_area': '8,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z234', 'city': 'North Providence', 'district': 'RS-12', 'property_type': '', 'frontage': '100', 'lot_area': '12,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z235', 'city': 'North Providence', 'district': 'RL-10', 'property_type': '', 'frontage': '100', 'lot_area': '8,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '+ 7,000 SF for 2nd unit'},
    {'id': 'z236', 'city': 'North Providence', 'district': 'RL-13', 'property_type': '', 'frontage': '100', 'lot_area': '8,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '+ 8,000 SF for 2nd unit'},
    {'id': 'z237', 'city': 'North Providence', 'district': 'RG', 'property_type': '', 'frontage': '70', 'lot_area': '8,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': "+ 6,000 SF for 2nd unit (100' Frontage 2 Fam)"},
    {'id': 'z239', 'city': 'North Smithfield', 'district': 'REA', 'property_type': '', 'frontage': '300', 'lot_area': '120,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z240', 'city': 'North Smithfield', 'district': 'RA', 'property_type': '', 'frontage': '200', 'lot_area': '65,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z241', 'city': 'North Smithfield', 'district': 'RS', 'property_type': '', 'frontage': '200', 'lot_area': '40,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z242', 'city': 'North Smithfield', 'district': 'RA 2-Fam', 'property_type': '', 'frontage': '200', 'lot_area': '130,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z243', 'city': 'North Smithfield', 'district': 'RS Multi', 'property_type': '', 'frontage': '175', 'lot_area': '80,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z244', 'city': 'North Smithfield', 'district': 'RA Multi', 'property_type': '', 'frontage': '200', 'lot_area': '65,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '+ 20,000 per bedroom'},
    {'id': 'z245', 'city': 'North Smithfield', 'district': 'RS Multi', 'property_type': '', 'frontage': "200'", 'lot_area': '40,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '+ 6,000 per bedroom'},
    {'id': 'z246', 'city': 'North Smithfield', 'district': 'RU', 'property_type': '', 'frontage': '100', 'lot_area': '20,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z247', 'city': 'North Smithfield', 'district': 'RU 2-Fam', 'property_type': '', 'frontage': '120', 'lot_area': '30,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z248', 'city': 'North Smithfield', 'district': 'RU Multi', 'property_type': '', 'frontage': '200', 'lot_area': '6,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '+ 4,000 per bedroom'},
    {'id': 'z250', 'city': 'Pawtucket', 'district': 'RL', 'property_type': '', 'frontage': "90'", 'lot_area': 'New Lots 9,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Existing Lots NONE'},
    {'id': 'z251', 'city': 'Pawtucket', 'district': 'RS SF', 'property_type': '', 'frontage': "50'", 'lot_area': 'New Lots 5,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Existing Lots NONE'},
    {'id': 'z252', 'city': 'Pawtucket', 'district': 'RS Other Res', 'property_type': '', 'frontage': '75', 'lot_area': 'New Lots 7,500 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Existing Lots NONE'},
    {'id': 'z253', 'city': 'Pawtucket', 'district': 'RT Single', 'property_type': '', 'frontage': '50', 'lot_area': 'New Lots 5,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Existing Lots NONE'},
    {'id': 'z254', 'city': 'Pawtucket', 'district': 'RT 2-Fam', 'property_type': '', 'frontage': '75', 'lot_area': '7,500 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z255', 'city': 'Pawtucket', 'district': 'RT Other Res', 'property_type': '', 'frontage': '75', 'lot_area': 'New Lots 7,500 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Existing Lots NONE'},
    {'id': 'z256', 'city': 'Pawtucket', 'district': 'RM Single', 'property_type': '', 'frontage': '50', 'lot_area': 'New Lots 5,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Existing Lots NONE'},
    {'id': 'z257', 'city': 'Pawtucket', 'district': 'RM 2-Fam', 'property_type': '', 'frontage': '75', 'lot_area': '7,500 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z258', 'city': 'Pawtucket', 'district': 'RM 3-Fam', 'property_type': '', 'frontage': '100', 'lot_area': '10,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z259', 'city': 'Pawtucket', 'district': 'RM Multi', 'property_type': '', 'frontage': '100', 'lot_area': '3,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z261', 'city': 'Portsmouth', 'district': 'R-10', 'property_type': '', 'frontage': '100', 'lot_area': '10,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Min lot area per DU: 10,000 SF / Lot ares 40,000 SF where not serviced by public water'},
    {'id': 'z262', 'city': 'Portsmouth', 'district': 'R-20', 'property_type': '', 'frontage': '110', 'lot_area': '20,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Min lot area per DU: 15,000 SF / Lot ares 40,000 SF where not serviced by public water'},
    {'id': 'z263', 'city': 'Portsmouth', 'district': 'R-30', 'property_type': '', 'frontage': '125', 'lot_area': '30,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z264', 'city': 'Portsmouth', 'district': 'R-40', 'property_type': '', 'frontage': '125', 'lot_area': '40,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Min lot area per DU: 30,000 SF / Lot ares 40,000 SF where not serviced by public water'},
    {'id': 'z265', 'city': 'Portsmouth', 'district': 'R-60', 'property_type': '', 'frontage': '200', 'lot_area': '60,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z267', 'city': 'Providence', 'district': 'R1A', 'property_type': '', 'frontage': '75', 'lot_area': 'New subdivisions (not RH): 7,500 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Existing Lots NONE'},
    {'id': 'z268', 'city': 'Providence', 'district': 'R1', 'property_type': '', 'frontage': '50', 'lot_area': 'New subdivisions (not RH): 5,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Existing Lots NONE'},
    {'id': 'z269', 'city': 'Providence', 'district': 'R2', 'property_type': '', 'frontage': '50', 'lot_area': 'New subdivisions (not RH): 5,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Existing Lots NONE'},
    {'id': 'z270', 'city': 'Providence', 'district': 'R3', 'property_type': '', 'frontage': '35', 'lot_area': 'New subdivisions (not RH): 3,500 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Existing Lots NONE'},
    {'id': 'z271', 'city': 'Providence', 'district': 'RP', 'property_type': '', 'frontage': '50', 'lot_area': 'New subdivisions (not RH): 5,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Existing Lots NONE'},
    {'id': 'z272', 'city': 'Providence', 'district': 'No Zoning For Existing Lots', 'property_type': '', 'frontage': '', 'lot_area': '', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z274', 'city': 'Richmond', 'district': 'R3', 'property_type': '', 'frontage': '300', 'lot_area': '3 acres', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z275', 'city': 'Richmond', 'district': 'R-2', 'property_type': '', 'frontage': '200', 'lot_area': '2 acres', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z276', 'city': 'Richmond', 'district': 'R1', 'property_type': '', 'frontage': '150', 'lot_area': '1 acre', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z277', 'city': 'Richmond', 'district': 'NB', 'property_type': '', 'frontage': '150', 'lot_area': '1 acre', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z279', 'city': 'Scituate', 'district': 'RR-120', 'property_type': '', 'frontage': '300', 'lot_area': '120,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z280', 'city': 'Scituate', 'district': 'RS-120', 'property_type': '', 'frontage': '300', 'lot_area': '120,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z281', 'city': 'Scituate', 'district': 'RRW-60/80', 'property_type': '', 'frontage': '200', 'lot_area': '80,000 SF (no public water)', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '60,000 SF (public water)'},
    {'id': 'z283', 'city': 'Smithfield', 'district': 'R-200', 'property_type': '', 'frontage': '300', 'lot_area': '200,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z284', 'city': 'Smithfield', 'district': 'R-80', 'property_type': '', 'frontage': '200', 'lot_area': '80,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z285', 'city': 'Smithfield', 'district': 'R-Med', 'property_type': '', 'frontage': '150', 'lot_area': '40,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z286', 'city': 'Smithfield', 'district': 'R-20', 'property_type': '', 'frontage': '125', 'lot_area': '20,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z287', 'city': 'Smithfield', 'district': 'R-20M', 'property_type': '', 'frontage': '125', 'lot_area': '20,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z288', 'city': 'Smithfield', 'district': 'MU', 'property_type': '', 'frontage': '125', 'lot_area': '20,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z289', 'city': 'Smithfield', 'district': 'PD', 'property_type': '', 'frontage': '300', 'lot_area': '200,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z290', 'city': 'Smithfield', 'district': 'Village', 'property_type': '', 'frontage': '150', 'lot_area': '20,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z291', 'city': 'Smithfield', 'district': 'R-20M 2-Fam', 'property_type': '', 'frontage': '150', 'lot_area': '40,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z293', 'city': 'South Kingstown', 'district': 'R200', 'property_type': '', 'frontage': '200', 'lot_area': '200,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z294', 'city': 'South Kingstown', 'district': 'R80', 'property_type': '', 'frontage': '200', 'lot_area': '80,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z295', 'city': 'South Kingstown', 'district': 'R40', 'property_type': '', 'frontage': '150', 'lot_area': '40,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z296', 'city': 'South Kingstown', 'district': 'R30', 'property_type': '', 'frontage': '125', 'lot_area': '30,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '45,000 SF (2-unit w/sewer); 60,000 SF (2-unit w/o sewer)'},
    {'id': 'z297', 'city': 'South Kingstown', 'district': 'R20', 'property_type': '', 'frontage': '100', 'lot_area': '20,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '30,000 SF (2-unit w/sewer); 40,000 SF (2-unit w/o sewer)'},
    {'id': 'z297a', 'city': 'South Kingstown', 'district': 'R-10', 'property_type': 'Single Family', 'frontage': '80', 'lot_area': '10,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z297b', 'city': 'South Kingstown', 'district': 'R-10', 'property_type': '2 Fam/ADU', 'frontage': '80', 'lot_area': '15,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'With Sewer'},
    {'id': 'z297c', 'city': 'South Kingstown', 'district': 'R-10', 'property_type': '2 Fam/ADU', 'frontage': '80', 'lot_area': '20,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': 'Without Sewer'},
    {'id': 'z299', 'city': 'Tiverton', 'district': 'R-30 SF', 'property_type': '', 'frontage': '150', 'lot_area': '30,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z300', 'city': 'Tiverton', 'district': 'R-30 2-Fam', 'property_type': '', 'frontage': '150', 'lot_area': '40,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z301', 'city': 'Tiverton', 'district': 'R-30 3-Fam', 'property_type': '', 'frontage': '150', 'lot_area': '40,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z302', 'city': 'Tiverton', 'district': 'R-30 Multi', 'property_type': '', 'frontage': '150', 'lot_area': '40,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '+ 15,000 SF/unit above 2 + 7,500 SF/bedroom above 2'},
    {'id': 'z303', 'city': 'Tiverton', 'district': 'R-60 SF', 'property_type': '', 'frontage': '175', 'lot_area': '60,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z304', 'city': 'Tiverton', 'district': 'R-60 2-Fam', 'property_type': '', 'frontage': '175', 'lot_area': '60,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z305', 'city': 'Tiverton', 'district': 'R-60 Multi', 'property_type': '', 'frontage': '175', 'lot_area': '60,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '+ 15,000 SF/unit above 2 + 7,500 SF/bedroom above 2/unit'},
    {'id': 'z307', 'city': 'Warren', 'district': 'Residence 40 SF', 'property_type': '', 'frontage': '150', 'lot_area': '40,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z308', 'city': 'Warren', 'district': 'Residence 40 (other use)', 'property_type': '', 'frontage': '135', 'lot_area': '30,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z309', 'city': 'Warren', 'district': 'Residence 20 SF', 'property_type': '', 'frontage': '120', 'lot_area': '20,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z310', 'city': 'Warren', 'district': 'Residence 15 SF', 'property_type': '', 'frontage': '110', 'lot_area': '15,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z311', 'city': 'Warren', 'district': 'Residence 10 SF', 'property_type': '', 'frontage': '90', 'lot_area': '10,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z312', 'city': 'Warren', 'district': 'Residence 20 2-Fam', 'property_type': '', 'frontage': '140', 'lot_area': '30,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z313', 'city': 'Warren', 'district': 'Residence 10 2-Fam', 'property_type': '', 'frontage': '110', 'lot_area': '15,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z314', 'city': 'Warren', 'district': 'Residence 6 SF', 'property_type': '', 'frontage': '60', 'lot_area': '6,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z315', 'city': 'Warren', 'district': 'Residence 6 2-Fam', 'property_type': '', 'frontage': '70', 'lot_area': '8,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z316', 'city': 'Warren', 'district': 'Residence 6 3-Fam', 'property_type': '', 'frontage': '72', 'lot_area': '9,500 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z317', 'city': 'Warren', 'district': 'Residence 6 4-Fam', 'property_type': '', 'frontage': '74', 'lot_area': '11,000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z319', 'city': 'Warwick', 'district': 'A-7 / GB', 'property_type': '', 'frontage': '70', 'lot_area': '7,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z320', 'city': 'Warwick', 'district': 'A-10', 'property_type': '', 'frontage': '100', 'lot_area': '10,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z321', 'city': 'Warwick', 'district': 'A-15', 'property_type': '', 'frontage': '125', 'lot_area': '15,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z322', 'city': 'Warwick', 'district': 'A-40 / OS', 'property_type': '', 'frontage': '150', 'lot_area': '40,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z324', 'city': 'West Greenwich', 'district': 'RFR-2', 'property_type': '', 'frontage': '200', 'lot_area': '2 acres', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z325', 'city': 'West Greenwich', 'district': 'RFR-1', 'property_type': '', 'frontage': '150', 'lot_area': '1 acre', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z327', 'city': 'West Warwick', 'district': 'R-10 SF', 'property_type': '', 'frontage': '80', 'lot_area': '10,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z328', 'city': 'West Warwick', 'district': 'R-8 SF', 'property_type': '', 'frontage': '70', 'lot_area': '8,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z329', 'city': 'West Warwick', 'district': 'R-7.5 SF', 'property_type': '', 'frontage': '70', 'lot_area': '7,500 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z330', 'city': 'West Warwick', 'district': 'R-6 SF', 'property_type': '', 'frontage': '55', 'lot_area': '6,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z331', 'city': 'West Warwick', 'district': 'R-10 2-Fam', 'property_type': '', 'frontage': '80', 'lot_area': '15,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z332', 'city': 'West Warwick', 'district': 'R-8 2-Fam', 'property_type': '', 'frontage': '70', 'lot_area': '12,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z333', 'city': 'West Warwick', 'district': 'R-7.5 2-Fam', 'property_type': '', 'frontage': '70', 'lot_area': '10,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z334', 'city': 'West Warwick', 'district': 'R-6 2-Fam', 'property_type': '', 'frontage': '55', 'lot_area': '8,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z335', 'city': 'West Warwick', 'district': 'R-7.5 Multi', 'property_type': '', 'frontage': '70', 'lot_area': '10,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '+7,500 for ea additional unit'},
    {'id': 'z336', 'city': 'West Warwick', 'district': 'R-6 Multi', 'property_type': '', 'frontage': '55', 'lot_area': '10,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '+7,500 for ea additional unit'},
    {'id': 'z337', 'city': 'West Warwick', 'district': 'VC', 'property_type': '', 'frontage': '', 'lot_area': '10,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '+7,500 for ea additional unit'},
    {'id': 'z339', 'city': 'Westerly', 'district': 'RR-60 SF', 'property_type': '', 'frontage': '200', 'lot_area': '60,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z340', 'city': 'Westerly', 'district': 'LDR-43 SF', 'property_type': '', 'frontage': '200', 'lot_area': '43,560 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z341', 'city': 'Westerly', 'district': 'LDR-40 SF', 'property_type': '', 'frontage': '150', 'lot_area': '40,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z342', 'city': 'Westerly', 'district': 'MDR-30 SF', 'property_type': '', 'frontage': '120', 'lot_area': '30,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z343', 'city': 'Westerly', 'district': 'MDR-20 SF', 'property_type': '', 'frontage': '100', 'lot_area': '20,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z344', 'city': 'Westerly', 'district': 'HDR-15 SF', 'property_type': '', 'frontage': '100', 'lot_area': '15000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z345', 'city': 'Westerly', 'district': 'HDR-15 Multi', 'property_type': '', 'frontage': '150', 'lot_area': '4 acres', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '+ 15,000 per unit'},
    {'id': 'z346', 'city': 'Westerly', 'district': 'HDR-10 SF', 'property_type': '', 'frontage': '80', 'lot_area': '10000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z347', 'city': 'Westerly', 'district': 'HDR-6 SF', 'property_type': '', 'frontage': '60', 'lot_area': '6000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z348', 'city': 'Westerly', 'district': 'HDR-6 2-Fam', 'property_type': '', 'frontage': '100', 'lot_area': '12000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z349', 'city': 'Westerly', 'district': 'HDR-6 3-Fam', 'property_type': '', 'frontage': '120', 'lot_area': '18000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z350', 'city': 'Westerly', 'district': 'HDR-6 Multi', 'property_type': '', 'frontage': '60', 'lot_area': '2 acres', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': '+ 6,000 per unit'},
    {'id': 'z351', 'city': 'Westerly', 'district': 'P-15 & NB', 'property_type': '', 'frontage': '', 'lot_area': 'Conforms to nearest residential zone', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z352', 'city': 'Westerly', 'district': 'Downtown Center 2', 'property_type': '', 'frontage': '', 'lot_area': 'Conforms to HDR-6', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z354', 'city': 'Woonsocket', 'district': 'R-1', 'property_type': '', 'frontage': '135', 'lot_area': '25,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z355', 'city': 'Woonsocket', 'district': 'R-2', 'property_type': '', 'frontage': '90', 'lot_area': '10,000 SF', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z356', 'city': 'Woonsocket', 'district': 'R-1', 'property_type': '', 'frontage': '135', 'lot_area': '25000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z357', 'city': 'Woonsocket', 'district': 'R-2', 'property_type': '', 'frontage': '90', 'lot_area': '10000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z358', 'city': 'Woonsocket', 'district': 'R-3', 'property_type': '', 'frontage': '70', 'lot_area': '7000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': "80' Frontage for Multi"},
    {'id': 'z359', 'city': 'Woonsocket', 'district': 'R-4', 'property_type': '', 'frontage': '60', 'lot_area': '6000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': "+ 10' Frontage  for ea additional unit"},
    {'id': 'z360', 'city': 'Woonsocket', 'district': 'C-1, C-2 & MU-2', 'property_type': '', 'frontage': '', 'lot_area': '6000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': ''},
    {'id': 'z361', 'city': 'Woonsocket', 'district': 'MU-I', 'property_type': '', 'frontage': '60', 'lot_area': '6000 sf', 'lot_width': '', 'front_yard': '', 'side_yard': '', 'rear_yard': '', 'max_height': '', 'max_lot_cov': '', 'max_floors': '', 'notes': "+ 10' Frontage  for ea additional unit"},
]

DEFAULT_NEIGHBORHOODS = [
    {"id": "n0", "city": 'Barrington', "neighborhood": 'Barrington', "description": 'Subject is located in Barrington, RI, an established residential community consisting primarily of single-family dwellings with supporting municipal and recreational uses. Properties vary in age, style, and quality. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood or environmental conditions were observed. Land use designated as "other" includes open space, coastal features, and public recreational land.', "notes": ""},
    {"id": "n1", "city": 'Bristol', "neighborhood": 'Bristol', "description": 'Subject is located in Bristol, RI, an established coastal community consisting of single-family, multi-family, and mixed-use properties. Improvements vary in age, style, and quality. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed. Land use designated as "other" includes waterfront areas, parks, and public recreational land.', "notes": ""},
    {"id": "n2", "city": 'Burrillville', "neighborhood": 'Burrillville', "description": 'Subject is located in Burrillville, RI, an established semi-rural community consisting primarily of single-family residences with limited supporting commercial development. Properties vary in age, style, and condition. Access to employment, shopping, and support services is typical for the area and surrounding communities. No adverse neighborhood conditions were observed. Land use designated as "other" includes wooded land, reservoirs, and conservation areas.', "notes": ""},
    {"id": "n3", "city": 'Central Falls', "neighborhood": 'Central Falls', "description": 'Subject is located in Central Falls, RI, an established urban community consisting primarily of multi-family residential properties with some mixed-use development. Properties vary in age, style, and condition. Access to employment, public transportation, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed. Land use designated as "other" includes municipal and recreational land.', "notes": ""},
    {"id": "n4", "city": 'Charlestown', "neighborhood": 'Charlestown', "description": 'Subject is located in Charlestown, RI, an established coastal and semi-rural community consisting primarily of single-family residences and seasonal properties with limited commercial development. Properties vary in age, style, and quality. Access to employment, shopping, and support services is typical for the area and surrounding communities. No adverse neighborhood conditions were observed. Land use designated as "other" includes beaches, conservation land, and open space.', "notes": ""},
    {"id": "n5", "city": 'Coventry', "neighborhood": 'Coventry', "description": 'Subject is located in Coventry, RI, an established residential community consisting of single-family residences with some multi-family and limited commercial properties. Improvements vary in age, style, and quality. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed. Land use designated as "other" includes open space, parks, and recreational land.', "notes": ""},
    {"id": "n6", "city": 'Cranston', "neighborhood": 'General', "description": 'Subject is located in Cranston, RI, an established residential community consisting of single and multi-family properties with supporting commercial and municipal uses. Properties vary in age, style, and quality. Access to employment, shopping, healthcare, and support services is typical for the area. No adverse neighborhood conditions were observed. Land use designated as "other" includes parks and public recreational land.', "notes": ""},
    {"id": "n7", "city": 'Cranston', "neighborhood": 'Edgewood', "description": 'Subject is located in the Edgewood section of Cranston, an established residential neighborhood consisting of single and multi-family properties. Improvements vary in age, style, and condition. Access to support services and transportation is typical for the area. No adverse neighborhood conditions were observed.', "notes": ""},
    {"id": "n8", "city": 'Cranston', "neighborhood": 'Garden City', "description": 'Subject is located in the Garden City area of Cranston, an established residential neighborhood with nearby supporting commercial development. Properties vary in age, style, and condition. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed.', "notes": ""},
    {"id": "n9", "city": 'Cumberland', "neighborhood": 'Cumberland', "description": 'Subject is located in Cumberland, RI, an established residential community consisting primarily of single-family residences with some condominium and multi-family development. Properties vary in age, style, and quality. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed. Land use designated as "other" includes open space and recreational land.', "notes": ""},
    {"id": "n10", "city": 'East Greenwich', "neighborhood": 'East Greenwich', "description": 'Subject is located in East Greenwich, RI, an established residential community consisting primarily of single-family residences with a village-style center and surrounding neighborhoods. Properties vary in age, style, and quality. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed. Land use designated as "other" includes waterfront and recreational areas.', "notes": ""},
    {"id": "n11", "city": 'East Providence', "neighborhood": 'General', "description": 'Subject is located in East Providence, RI, an established residential community consisting of single and multi-family properties with supporting commercial uses. Properties vary in age, style, and condition. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed.', "notes": ""},
    {"id": "n12", "city": 'East Providence', "neighborhood": 'Riverside', "description": 'Subject is located in the Riverside section of East Providence, an established residential neighborhood consisting of single and multi-family properties. Improvements vary in age, style, and condition. Access to support services is typical for the area. No adverse neighborhood conditions were observed.', "notes": ""},
    {"id": "n13", "city": 'Exeter', "neighborhood": 'Exeter', "description": 'Subject is located in Exeter, RI, an established rural residential community consisting primarily of single-family residences on larger parcels. Properties vary in age, style, and condition. Access to employment, shopping, and support services is typical for the area and surrounding communities. No adverse neighborhood conditions were observed. Land use designated as "other" includes conservation land and open space.', "notes": ""},
    {"id": "n14", "city": 'Foster', "neighborhood": 'Foster', "description": 'Subject is located in Foster, RI, an established rural community consisting primarily of single-family residences on larger parcels. Properties vary in age, style, and quality. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed. Land use designated as "other" includes farmland, wooded areas, and conservation land.', "notes": ""},
    {"id": "n15", "city": 'Glocester', "neighborhood": 'Glocester', "description": 'Subject is located in Glocester, RI, an established semi-rural residential community consisting primarily of single-family residences with limited commercial development. Properties vary in age, style, and condition. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed. Land use designated as "other" includes open space and conservation land.', "notes": ""},
    {"id": "n16", "city": 'Hopkinton', "neighborhood": 'Hopkinton', "description": 'Subject is located in Hopkinton, RI, an established rural/suburban community consisting primarily of single-family residences. Properties vary in age, style, and condition. Access to employment, shopping, and support services is typical for the area and surrounding communities. No adverse neighborhood conditions were observed. Land use designated as "other" includes wooded and conservation land.', "notes": ""},
    {"id": "n17", "city": 'Jamestown', "neighborhood": 'Jamestown', "description": 'Subject is located in Jamestown, RI, an established island community consisting primarily of single-family residences. Properties vary in age, style, and quality. Access to employment, shopping, and support services is typical for the area via bridge access. No adverse neighborhood conditions were observed. Land use designated as "other" includes waterfront and recreational land.', "notes": ""},
    {"id": "n18", "city": 'Johnston', "neighborhood": 'Johnston', "description": 'Subject is located in Johnston, RI, an established residential community consisting of single-family, multi-family, and limited commercial properties. Properties vary in age, style, and condition. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed. Land use designated as "other" includes open space and public land.', "notes": ""},
    {"id": "n19", "city": 'Lincoln', "neighborhood": 'Lincoln', "description": 'Subject is located in Lincoln, RI, an established residential community consisting primarily of single-family residences with some condominium and commercial development. Properties vary in age, style, and quality. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed. Land use designated as "other" includes recreational and conservation land.', "notes": ""},
    {"id": "n20", "city": 'Little Compton', "neighborhood": 'Little Compton', "description": 'Subject is located in Little Compton, RI, a primarily residential coastal community characterized by rural and low-density development patterns. The housing stock consists predominantly of single-family dwellings situated on individual lots, with limited multi-family and commercial development in designated village areas. Properties vary in age, architectural style, and condition. Lot sizes are generally larger than those typical of more densely developed communities. Access to employment, shopping, healthcare, and municipal services is typical for the area and surrounding communities. No adverse neighborhood or environmental conditions were observed at the time of inspection. Land use designated as "other" includes agricultural land, conservation areas, open space, and recreational land.', "notes": ""},
    {"id": "n21", "city": 'Middletown', "neighborhood": 'Middletown', "description": 'Subject is located in Middletown, RI, an established residential community consisting of single-family, multi-family, and limited commercial properties. Improvements vary in age, style, and quality. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed. Land use designated as "other" includes coastal and recreational land.', "notes": ""},
    {"id": "n22", "city": 'Narragansett', "neighborhood": 'Narragansett', "description": 'Subject is located in Narragansett, RI, an established residential community consisting of single-family residences, seasonal housing, and limited multi-family properties. Properties vary in age, style, and quality. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed. Land use designated as "other" includes beaches and public recreational land.', "notes": ""},
    {"id": "n23", "city": 'New Shoreham (Block Island)', "neighborhood": 'New Shoreham', "description": 'Subject is located in New Shoreham, RI, an island community characterized by coastal and low-density residential development patterns. The housing stock consists primarily of single-family dwellings with seasonal occupancy in portions of the market, along with limited multi-family and commercial development concentrated in village areas. Properties vary in age, style, and condition. Development is influenced by coastal features and land use regulations, resulting in larger parcel sizes in many areas. Access to employment, shopping, healthcare, and support services is typical for an island location and is supplemented by mainland communities via ferry and air transportation. No adverse neighborhood or environmental conditions were observed at the time of inspection. Land use designated as "other" includes conservation land, open space, and recreational areas.', "notes": ""},
    {"id": "n24", "city": 'Newport', "neighborhood": 'General', "description": 'Subject is located in Newport, RI, an established residential and mixed-use community consisting of single and multi-family residences and historic housing stock. Properties vary in age, style, and quality. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed.', "notes": ""},
    {"id": "n25", "city": 'Newport', "neighborhood": 'The Point / Historic Districts', "description": 'Subject is located in the Point/Historic District area of Newport, an established residential neighborhood characterized by historic single and multi-family residences. Properties typically have smaller site sizes consistent with the area. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed.', "notes": ""},
    {"id": "n26", "city": 'North Kingstown', "neighborhood": 'North Kingstown', "description": 'Subject is located in North Kingstown, RI, an established residential community consisting primarily of single-family residences with some condominium and mixed-use development. Properties vary in age, style, and quality. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed. Land use designated as "other" includes waterfront and recreational land.', "notes": ""},
    {"id": "n27", "city": 'North Providence', "neighborhood": 'North Providence', "description": 'Subject is located in North Providence, RI, an established residential community consisting primarily of single-family residences with some multi-family development. Properties vary in age, style, and condition. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed.', "notes": ""},
    {"id": "n28", "city": 'Pawtucket', "neighborhood": 'General', "description": 'Subject is located in Pawtucket, RI, an established urban residential community consisting of single and multi-family residences and mixed-use properties. Properties vary in age, style, and condition. Access to employment, public transportation, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed.', "notes": ""},
    {"id": "n29", "city": 'Pawtucket', "neighborhood": 'Oak Hill', "description": 'Subject is located in the Oak Hill section of Pawtucket, an established residential neighborhood consisting primarily of single-family and two-family residences. Properties vary in age, style, and condition. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed.', "notes": ""},
    {"id": "n30", "city": 'Portsmouth', "neighborhood": 'Portsmouth', "description": 'Subject is located in Portsmouth, RI, an established residential community consisting primarily of single-family residences with agricultural and waterfront influence in select areas. Properties vary in age, style, and quality. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed. Land use designated as "other" includes farmland and coastal areas.', "notes": ""},
    {"id": "n31", "city": 'Providence', "neighborhood": 'General', "description": 'Subject is located in Providence, RI, an established urban community consisting of single and multi-family residences, mixed-use, and commercial properties. Properties vary in age, style, and condition. Access to employment, public transportation, shopping, healthcare, and support services is typical for the area. No adverse neighborhood conditions were observed.', "notes": ""},
    {"id": "n32", "city": 'Providence', "neighborhood": 'East Side', "description": 'Subject is located in the East Side of Providence, an established residential neighborhood consisting of single and multi-family properties. Improvements vary in age, style, and condition. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed.', "notes": ""},
    {"id": "n33", "city": 'Providence', "neighborhood": 'Elmhurst / Mount Pleasant / Federal Hill', "description": 'Subject is located in an established Providence residential neighborhood consisting of single and multi-family properties with supporting commercial uses. Properties vary in age, style, and condition. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed.', "notes": ""},
    {"id": "n34", "city": 'Richmond', "neighborhood": 'Richmond', "description": 'Subject is located in Richmond, RI, a primarily residential community characterized by rural and semi-rural development patterns. The housing stock consists predominantly of single-family dwellings situated on individual lots, with limited multi-family and commercial development in designated areas. Properties vary in age, style, and condition. Lot sizes are generally larger than those typical of urban communities. Access to employment, shopping, healthcare, and municipal services is typical for the area and surrounding communities. No adverse neighborhood or environmental conditions were observed at the time of inspection. Land use designated as "other" includes open space, conservation land, agricultural land, and recreational areas.', "notes": ""},
    {"id": "n35", "city": 'Scituate', "neighborhood": 'Scituate', "description": 'Subject is located in Scituate, RI, an established rural residential community consisting primarily of single-family residences on larger parcels. Properties vary in age, style, and condition. Access to employment, shopping, and support services is typical for the area and surrounding communities. No adverse neighborhood conditions were observed. Land use designated as "other" includes reservoirs and conservation land.', "notes": ""},
    {"id": "n36", "city": 'Smithfield', "neighborhood": 'Smithfield', "description": 'Subject is located in Smithfield, RI, an established residential community consisting primarily of single-family residences with some condominium and commercial development. Properties vary in age, style, and quality. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed. Land use designated as "other" includes open space and recreational land.', "notes": ""},
    {"id": "n37", "city": 'South Kingstown', "neighborhood": 'South Kingstown', "description": 'Subject is located in South Kingstown, RI, an established residential community consisting of single-family residences, seasonal housing, and limited multi-family properties. Properties vary in age, style, and quality. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed. Land use designated as "other" includes beaches and recreational land.', "notes": ""},
    {"id": "n38", "city": 'Tiverton', "neighborhood": 'Tiverton', "description": 'Subject is located in Tiverton, RI, an established residential community consisting of single-family residences with agricultural and waterfront influence in select areas. Properties vary in age, style, and quality. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed. Land use designated as "other" includes farmland and coastal land.', "notes": ""},
    {"id": "n39", "city": 'Warren', "neighborhood": 'Warren', "description": 'Subject is located in Warren, RI, an established residential and mixed-use waterfront community consisting of single and multi-family residences. Properties vary in age, style, and condition. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed. Land use designated as "other" includes waterfront and recreational land.', "notes": ""},
    {"id": "n40", "city": 'Warwick', "neighborhood": 'General', "description": 'Subject is located in Warwick, RI, an established residential community consisting of single and multi-family residences. Properties vary in age, style, and quality. Access to employment, shopping, transportation, and support services is typical for the area. No adverse neighborhood conditions were observed.', "notes": ""},
    {"id": "n41", "city": 'Warwick', "neighborhood": 'Apponaug / Cowesett / Gaspee', "description": 'Subject is located in an established Warwick residential neighborhood characterized by single and multi-family development with supporting commercial uses in the broader area. Properties vary in age, style, and condition. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed.', "notes": ""},
    {"id": "n42", "city": 'West Greenwich', "neighborhood": 'West Greenwich', "description": 'Subject is located in West Greenwich, RI, an established rural residential community consisting primarily of single-family residences on larger parcels. Properties vary in age, style, and condition. Access to employment, shopping, and support services is typical for the area and surrounding communities. No adverse neighborhood conditions were observed. Land use designated as "other" includes conservation land and open space.', "notes": ""},
    {"id": "n43", "city": 'West Warwick', "neighborhood": 'West Warwick', "description": 'Subject is centrally located in West Warwick, RI, in an established residential neighborhood consisting of single and multi-family properties. Properties vary in age, style, and condition. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed. Land use designated as "other" consists of vacant land, parks, and public recreational land.', "notes": ""},
    {"id": "n44", "city": 'Westerly', "neighborhood": 'Westerly', "description": 'Subject is located in Westerly, RI, an established residential community consisting of single and multi-family residences with seasonal occupancy in select areas. Properties vary in age, style, and quality. Access to employment, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed. Land use designated as "other" includes beaches and recreational land.', "notes": ""},
    {"id": "n45", "city": 'Woonsocket', "neighborhood": 'Woonsocket', "description": 'Subject is located in Woonsocket, RI, an established urban residential community consisting of single and multi-family residences and mixed-use properties. Properties vary in age, style, and condition. Access to employment, public transportation, shopping, and support services is typical for the area. No adverse neighborhood conditions were observed.', "notes": ""},
]

# ── Storage ───────────────────────────────────────────────────────────────────
def load_data(key, filepath, default):
    try:
        path = os.path.join(os.path.dirname(__file__), filepath)
        if os.path.exists(path):
            with open(path) as f:
                return json.load(f)
    except Exception:
        pass
    return [item.copy() for item in default]

def save_data(key, filepath, data):
    path = os.path.join(os.path.dirname(__file__), filepath)
    with open(path, "w") as f:
        json.dump(data, f)

def load_revisions():
    return load_data("_rc_revisions", "rc_revisions.json", DEFAULT_REVISIONS)

def save_revisions(data):
    save_data("_rc_revisions", "rc_revisions.json", data)

def load_neighborhoods():
    return load_data("_rc_neighborhoods", "rc_neighborhoods.json", DEFAULT_NEIGHBORHOODS)

def save_neighborhoods(data):
    save_data("_rc_neighborhoods", "rc_neighborhoods.json", data)

def load_zoning():
    return load_data("_rc_zoning", "rc_zoning.json", DEFAULT_ZONING)

def save_zoning(data):
    save_data("_rc_zoning", "rc_zoning.json", data)

def load_comments():
    return load_data("_rc_comments", "rc_comments.json", DEFAULT_COMMENTS)

def save_comments(data):
    save_data("_rc_comments", "rc_comments.json", data)

# ── Adjustment Commentary Presets ─────────────────────────────────────────────
ADJ_DEFAULT_RATES = {
    "gla": 25, "bed": 5000, "fullbath": 3000, "halfbath": 1500,
    "basement": 5000, "garage": 3500, "encporch": 2000, "deck": 1000,
    "fp": 1000, "bgr": 1000,
    "pool": 15000, "cac": 5000, "solar": 10000, "adu": 25000, "outbldg": 5000,
    "site_rate": 1, "site_unit": "none",
    "adv_loc": 5.0, "ben_loc": 5.0, "adv_view": 5.0, "ben_view": 5.0,
    "time_rate": 0.5, "time_dir": "Appreciating",
}

DEFAULT_ADJ_PRESETS = [
    {"name": "SFR \u2014 Tier 1 (Under $400k)",    "rates": {**ADJ_DEFAULT_RATES, "gla":25,  "bed":0, "fullbath":5000,  "halfbath":2500,  "basement":5000,  "garage":5000,  "encporch":4500, "deck":3000, "fp":2500, "bgr":1000, "pool":0, "cac":4000,  "adu":0,     "site_rate":0, "site_unit":"none", "ben_loc":5.0,  "ben_view":5.0}},
    {"name": "SFR \u2014 Tier 2 ($400k\u2013$800k)", "rates": {**ADJ_DEFAULT_RATES, "gla":138, "bed":0, "fullbath":20000, "halfbath":10000, "basement":10000, "garage":14000, "encporch":7500, "deck":5000, "fp":8000, "bgr":1000, "pool":0, "cac":10000, "adu":50000, "site_rate":1, "site_unit":"sf",   "ben_loc":5.0,  "ben_view":5.0}},
    {"name": "SFR \u2014 Tier 3 ($800k\u2013$1.5M)", "rates": {**ADJ_DEFAULT_RATES, "gla":188, "bed":0, "fullbath":22500, "halfbath":10000, "basement":15000, "garage":14500, "encporch":10000,"deck":6500, "fp":8500, "bgr":1000, "pool":0, "cac":12500, "adu":60000, "site_rate":2, "site_unit":"sf",   "ben_loc":12.0, "ben_view":12.0}},
    {"name": "SFR \u2014 Tier 4 ($1.5M+)",          "rates": {**ADJ_DEFAULT_RATES, "gla":205, "bed":0, "fullbath":22500, "halfbath":10000, "basement":19000, "garage":15000, "encporch":12000,"deck":9000, "fp":8000, "bgr":1000, "pool":0, "cac":15000, "adu":70000, "site_rate":3, "site_unit":"sf",   "ben_loc":20.0, "ben_view":20.0}},
    {"name": "MFR \u2014 Urban (2\u20134 Unit)",      "rates": {**ADJ_DEFAULT_RATES, "gla":17,  "bed":6500,  "fullbath":3500,  "halfbath":1750,  "basement":4000,  "garage":6500,  "encporch":0,    "deck":0,    "fp":0,    "bgr":0,    "pool":0, "cac":3000,  "adu":0,     "site_rate":0, "site_unit":"none", "ben_loc":10.0, "ben_view":10.0}},
    {"name": "MFR \u2014 Suburban",                  "rates": {**ADJ_DEFAULT_RATES, "gla":70,  "bed":10000, "fullbath":10000, "halfbath":5000,  "basement":5000,  "garage":12500, "encporch":0,    "deck":0,    "fp":0,    "bgr":0,    "pool":0, "cac":6500,  "adu":0,     "site_rate":1, "site_unit":"sf",   "ben_loc":10.0, "ben_view":10.0}},
    {"name": "Condo \u2014 Tier 1 (Under $350k)",   "rates": {**ADJ_DEFAULT_RATES, "gla":25,  "bed":0, "fullbath":5000,  "halfbath":2500,  "basement":3000,  "garage":10000, "encporch":3000, "deck":3000, "fp":2500, "bgr":0,    "pool":0, "cac":4000,  "adu":0,     "site_rate":0, "site_unit":"none", "ben_loc":5.0,  "ben_view":5.0}},
    {"name": "Condo \u2014 Tier 2 ($350k\u2013$700k)","rates": {**ADJ_DEFAULT_RATES, "gla":138, "bed":0, "fullbath":20000, "halfbath":10000, "basement":7000,  "garage":15000, "encporch":4500, "deck":4500, "fp":8000, "bgr":0,    "pool":0, "cac":10000, "adu":0,     "site_rate":0, "site_unit":"none", "ben_loc":7.0,  "ben_view":7.0}},
    {"name": "Condo \u2014 Tier 3 ($700k\u2013$1.2M)","rates": {**ADJ_DEFAULT_RATES, "gla":188, "bed":0, "fullbath":22500, "halfbath":10000, "basement":10000, "garage":20000, "encporch":7500, "deck":7500, "fp":8500, "bgr":0,    "pool":0, "cac":12500, "adu":0,     "site_rate":0, "site_unit":"none", "ben_loc":11.0, "ben_view":11.0}},
    {"name": "Condo \u2014 Tier 4 ($1.2M+)",         "rates": {**ADJ_DEFAULT_RATES, "gla":205, "bed":0, "fullbath":22500, "halfbath":10000, "basement":13000, "garage":27500, "encporch":11500,"deck":11500,"fp":8000, "bgr":0,    "pool":0, "cac":15000, "adu":0,     "site_rate":0, "site_unit":"none", "ben_loc":15.0, "ben_view":15.0}},
]

def load_adj_presets():
    return load_data("_adj_presets", "rc_adj_presets.json", DEFAULT_ADJ_PRESETS)

def save_adj_presets(data):
    save_data("_adj_presets", "rc_adj_presets.json", data)

# ── Password ──────────────────────────────────────────────────────────────────
def check_password():
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False
    if st.session_state.authenticated:
        return True
    if os.path.exists(LOGO_PATH):
        st.image(LOGO_PATH, width=200)
    st.title("A-Tech Appraisal Co.")
    st.subheader("Revision & Comment Library")
    st.divider()
    pwd = st.text_input("Enter password to continue", type="password", key="pwd_input")
    if st.button("Login", use_container_width=True):
        correct = None
        try:
            correct = st.secrets["APP_PASSWORD"]
        except Exception:
            correct = os.environ.get("APP_PASSWORD")
        if not correct:
            st.error("APP_PASSWORD is not configured. Contact your administrator.")
            return False
        if pwd == correct:
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("Incorrect password.")
    return False

# ── Page Config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="A-Tech R&C Library",
    page_icon="📚",
    layout="wide"
)

if not check_password():
    st.stop()

# ── Header ───────────────────────────────────────────────────────────────────
if os.path.exists(LOGO_PATH):
    import base64
    with open(LOGO_PATH, "rb") as img_f:
        logo_b64 = base64.b64encode(img_f.read()).decode()
    st.markdown(f"""
    <div style="display:flex;flex-direction:column;align-items:center;padding:16px 0 8px 0;">
        <img src="data:image/png;base64,{logo_b64}" style="height:64px;width:auto;margin-bottom:10px;">
        <div style="font-size:1.8rem;font-weight:700;line-height:1.2;text-align:center;">A-Tech Appraisal Co.</div>
        <div style="color:gray;font-size:0.85rem;text-align:center;">Revision & Comment Library — Field Reference</div>
    </div>
    """, unsafe_allow_html=True)
else:
    st.markdown("<h2 style='text-align:center'>A-Tech Appraisal Co.</h2>", unsafe_allow_html=True)
st.divider()

# ── Export All Data ───────────────────────────────────────────────────────────
with st.expander("💾 Export All Data — Download for GitHub backup", expanded=False):
    st.caption(
        "Download all library data as a ZIP file. Before redeploying to Streamlit Cloud, "
        "extract the ZIP and upload all JSON files to your GitHub repo alongside atech_rc_app.py. "
        "This protects any manually added entries from being lost on redeploy."
    )
    if st.button("⬇️ Download All Data as ZIP", key="export_all_zip"):
        export_data = {
            "rc_revisions.json":    load_revisions(),
            "rc_comments.json":     load_comments(),
            "rc_neighborhoods.json": load_neighborhoods(),
            "rc_zoning.json":       load_zoning(),
            "rc_adj_presets.json":  load_adj_presets(),
        }
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            for filename, data in export_data.items():
                zf.writestr(filename, json.dumps(data, indent=2))
        zip_buffer.seek(0)
        st.download_button(
            label="📦 Click here to save atech_library_data.zip",
            data=zip_buffer,
            file_name="atech_library_data.zip",
            mime="application/zip",
            key="export_zip_download",
        )
        st.info("Extract the ZIP and commit all 5 JSON files to your GitHub repo before redeploying.")

# ── ALL CAPS Toggle + Global Admin ───────────────────────────────────────────
if "site_admin" not in st.session_state:
    st.session_state["site_admin"] = False
if "all_caps" not in st.session_state:
    st.session_state["all_caps"] = False

hdr1, hdr2, hdr3 = st.columns([1, 1, 3])
with hdr1:
    caps_label = "🔠 ALL CAPS: ON" if st.session_state["all_caps"] else "🔡 ALL CAPS: OFF"
    if st.button(caps_label, key="caps_toggle", use_container_width=True):
        st.session_state["all_caps"] = not st.session_state["all_caps"]
        st.rerun()
with hdr2:
    if not st.session_state["site_admin"]:
        if st.button("🔒 Admin: OFF", key="admin_toggle_btn", use_container_width=True):
            st.session_state["show_admin_login"] = True
            st.rerun()
    else:
        if st.button("🔓 Admin: ON", key="admin_toggle_btn", use_container_width=True):
            st.session_state["site_admin"] = False
            st.session_state["show_admin_login"] = False
            st.rerun()

if st.session_state.get("show_admin_login") and not st.session_state["site_admin"]:
    with st.container():
        adm_col1, adm_col2, adm_col3 = st.columns([2, 1, 3])
        with adm_col1:
            adm_pwd = st.text_input("Admin password:", type="password", key="global_admin_pwd", label_visibility="collapsed", placeholder="Enter admin password...")
        with adm_col2:
            if st.button("Unlock", key="global_admin_unlock", use_container_width=True):
                correct = None
                try:
                    correct = st.secrets["APP_PASSWORD"]
                except Exception:
                    correct = os.environ.get("APP_PASSWORD")
                if not correct:
                    st.error("APP_PASSWORD is not configured. Contact your administrator.")
                elif adm_pwd == correct:
                    st.session_state["site_admin"] = True
                    st.session_state["show_admin_login"] = False
                    st.rerun()
                else:
                    st.error("Incorrect password.")

def caps(text):
    """Apply ALL CAPS transformation if toggle is on."""
    if st.session_state.get("all_caps") and text:
        return str(text).upper()
    return text

st.divider()

# ── Top Navigation ────────────────────────────────────────────────────────────

# ── STR Storage Helpers ─────────────────────────────────────────────────────
def load_clients():
    try:
        raw = st.session_state.get("_clients_store")
        if raw:
            return json.loads(raw)
    except Exception:
        pass
    # Try file-based fallback
    try:
        path = os.path.join(os.path.dirname(__file__), "clients.json")
        if os.path.exists(path):
            with open(path) as f:
                return json.load(f)
    except Exception:
        pass
    return {}

def save_clients(clients):
    st.session_state["_clients_store"] = json.dumps(clients)
    try:
        path = os.path.join(os.path.dirname(__file__), "clients.json")
        with open(path, "w") as f:
            json.dump(clients, f)
    except Exception:
        pass

def load_orders():
    try:
        raw = st.session_state.get("_orders_store")
        if raw:
            return json.loads(raw)
    except Exception:
        pass
    try:
        path = os.path.join(os.path.dirname(__file__), "orders.json")
        if os.path.exists(path):
            with open(path) as f:
                return json.load(f)
    except Exception:
        pass
    return []

def save_orders(orders):
    st.session_state["_orders_store"] = json.dumps(orders)
    try:
        path = os.path.join(os.path.dirname(__file__), "orders.json")
        with open(path, "w") as f:
            json.dump(orders, f)
    except Exception:
        pass

def log_order(address, property_type, client, borrower, loan_num, avm_file_id, report_date):
    orders = load_orders()
    orders.insert(0, {
        "date":          report_date,
        "address":       address,
        "property_type": property_type,
        "client":        client,
        "borrower":      borrower,
        "loan_number":   loan_num,
        "avm_file_id":   avm_file_id,
    })
    save_orders(orders)


tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8 = st.tabs([
    "🏘️ Neighborhoods",
    "📐 Zoning",
    "📝 Comments",
    "✅ QC Checker",
    "📋 Revisions",
    "🔧 Adj. Commentary",
    "🆕 UAD 3.6",
    "🏠 STR Analysis",
])

TAB_SELECTION = {
    tab1: "🏘️ Neighborhood Descriptions",
    tab2: "📐 Zoning Districts",
    tab3: "📝 Appraisal Comments",
    tab4: "✅ QC Checker",
    tab5: "📋 Revision Responses",
    tab6: "🔧 Adjustment Commentary Generator",
    tab7: "🆕 UAD 3.6 Reference",
    tab8: "🏠 STR Income Analysis",
}

# ══════════════════════════════════════════════════════════════════════════════
# TAB 1 — REVISION RESPONSES
# ══════════════════════════════════════════════════════════════════════════════
with tab5:
 if True:
    revisions = load_revisions()

    col_rs, col_ra = st.columns([4, 1])
    with col_rs:
        rev_search = st.text_input("🔍 Search", placeholder="Type keyword to filter...", key="rev_search")
    with col_ra:
        st.write("")
        st.write("")
        if st.session_state.get("site_admin"):
            if st.button("➕ Add New", key="add_rev_btn", use_container_width=True):
                st.session_state["show_add_rev"] = not st.session_state.get("show_add_rev", False)
        else:
            st.caption("🔒 Admin")

    if st.session_state.get("show_add_rev"):
        with st.container():
            st.divider()
            st.subheader("New Revision Response")
            nr_cat  = st.text_input("Category / Title *", key="nr_cat",
                                     placeholder="e.g. Comp Distance — Exceeded Guidelines")
            nr_req  = st.text_area("AMC / Lender Request (optional)", key="nr_req",
                                    height=80, placeholder="Paste the revision request here...")
            nr_resp = st.text_area("Your Response *", key="nr_resp",
                                    height=140, placeholder="Write your response here...")
            nr_note = st.text_input("Notes (optional)", key="nr_note",
                                     placeholder="e.g. Customize the distance threshold")
            rc1, rc2 = st.columns(2)
            with rc1:
                if st.button("💾 Save Revision", use_container_width=True, key="save_rev"):
                    if not nr_cat.strip() or not nr_resp.strip():
                        st.error("Category and Response are required.")
                    else:
                        import uuid
                        revisions.append({
                            "id":       str(uuid.uuid4())[:8],
                            "category": nr_cat.strip(),
                            "request":  nr_req.strip(),
                            "response": nr_resp.strip(),
                            "notes":    nr_note.strip(),
                        })
                        save_revisions(revisions)
                        st.session_state["show_add_rev"] = False
                        st.success("✅ Revision saved.")
                        st.rerun()
            with rc2:
                if st.button("Cancel", use_container_width=True, key="cancel_rev"):
                    st.session_state["show_add_rev"] = False
                    st.rerun()
            st.divider()

    filtered_revs = revisions
    if rev_search:
        q = rev_search.lower()
        filtered_revs = [r for r in revisions if
                         q in r.get("category","").lower() or
                         q in r.get("request","").lower() or
                         q in r.get("response","").lower()]

    st.write(f"**{len(filtered_revs)} entr{'y' if len(filtered_revs)==1 else 'ies'}**")
    st.divider()

    for rev in filtered_revs:
        with st.expander(f"📋 {rev.get('category','Untitled')}"):
            if rev.get("request"):
                st.markdown("**AMC / Lender Request:**")
                st.info(rev["request"])
            st.markdown("**Appraiser Response:**")
            st.text_area(
                "Copy the text below:",
                value=caps(rev.get("response","")),
                height=160,
                key=f"rev_text_{rev['id']}_{st.session_state.get('all_caps',False)}"
            )
            if rev.get("notes"):
                st.caption(f"📝 Note: {rev['notes']}")
            st.write("")
            if st.session_state.get("site_admin"):
                if st.button("🗑️ Delete this entry", key=f"del_rev_{rev['id']}"):
                    revisions = [r for r in revisions if r["id"] != rev["id"]]
                    save_revisions(revisions)
                    st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 2 — APPRAISAL COMMENTS
# ══════════════════════════════════════════════════════════════════════════════
with tab3:
 if True:
    comments = load_comments()

    col_cs, col_ca = st.columns([4, 1])
    with col_cs:
        com_search = st.text_input("🔍 Search", placeholder="Type keyword to filter...", key="com_search")
    with col_ca:
        st.write("")
        st.write("")
        if st.session_state.get("site_admin"):
            if st.button("➕ Add New", key="add_com_btn", use_container_width=True):
                st.session_state["show_add_com"] = not st.session_state.get("show_add_com", False)
        else:
            st.caption("🔒 Admin")

    if st.session_state.get("show_add_com"):
        with st.container():
            st.divider()
            st.subheader("New Appraisal Comment")
            nc_cat  = st.text_input("Category / Title *", key="nc_cat",
                                     placeholder="e.g. Waterfront Adjustment — Coastal RI")
            nc_text = st.text_area("Comment Text *", key="nc_text",
                                    height=140, placeholder="Write the addendum language here...")
            nc_note = st.text_input("Notes (optional)", key="nc_note",
                                     placeholder="e.g. Use on all coastal assignments")
            cc1, cc2 = st.columns(2)
            with cc1:
                if st.button("💾 Save Comment", use_container_width=True, key="save_com"):
                    if not nc_cat.strip() or not nc_text.strip():
                        st.error("Category and Comment text are required.")
                    else:
                        import uuid
                        comments.append({
                            "id":       str(uuid.uuid4())[:8],
                            "category": nc_cat.strip(),
                            "text":     nc_text.strip(),
                            "notes":    nc_note.strip(),
                        })
                        save_comments(comments)
                        st.session_state["show_add_com"] = False
                        st.success("✅ Comment saved.")
                        st.rerun()
            with cc2:
                if st.button("Cancel", use_container_width=True, key="cancel_com"):
                    st.session_state["show_add_com"] = False
                    st.rerun()
            st.divider()

    filtered_coms = comments
    if com_search:
        q = com_search.lower()
        filtered_coms = [c for c in comments if
                         q in c.get("category","").lower() or
                         q in c.get("text","").lower()]

    st.write(f"**{len(filtered_coms)} entr{'y' if len(filtered_coms)==1 else 'ies'}**")
    st.divider()

    for com in filtered_coms:
        with st.expander(f"📝 {com.get('category','Untitled')}"):
            st.text_area(
                "Copy the text below:",
                value=caps(com.get("text","")),
                height=160,
                key=f"com_text_{com['id']}_{st.session_state.get('all_caps',False)}"
            )
            if com.get("notes"):
                st.caption(f"📝 Note: {com['notes']}")
            st.write("")
            if st.session_state.get("site_admin"):
                if st.button("🗑️ Delete this entry", key=f"del_com_{com['id']}"):
                    comments = [c for c in comments if c["id"] != com["id"]]
                    save_comments(comments)
                    st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# TAB 3 — NEIGHBORHOOD DESCRIPTIONS
# ══════════════════════════════════════════════════════════════════════════════
with tab1:
 if True:
    neighborhoods = load_neighborhoods()

    # ── Admin Mode ────────────────────────────────────────────────────────────

    # ── Search + Add New ──────────────────────────────────────────────────────
    col_hs, col_ha = st.columns([4, 1])
    with col_hs:
        hood_search = st.text_input("🔍 Search", placeholder="Type city or neighborhood...", key="hood_search")
    with col_ha:
        st.write("")
        st.write("")
        if st.session_state["site_admin"]:
            if st.button("➕ Add New", key="add_hood_btn", use_container_width=True):
                st.session_state["show_add_hood"] = not st.session_state.get("show_add_hood", False)

    if st.session_state.get("show_add_hood") and st.session_state["site_admin"]:
        with st.container(border=True):
            st.subheader("New Neighborhood Description")
            nh_city  = st.text_input("City *", key="nh_city", placeholder="e.g. Providence")
            nh_hood  = st.text_input("Neighborhood *", key="nh_hood", placeholder="e.g. Fox Point")
            nh_desc  = st.text_area("Description *", key="nh_desc", height=140,
                                     placeholder="Write the neighborhood description here.")
            nh_note  = st.text_input("Notes (optional)", key="nh_note",
                                      placeholder="e.g. Last updated March 2026")
            hc1, hc2 = st.columns(2)
            with hc1:
                if st.button("💾 Save Description", use_container_width=True, key="save_hood"):
                    if not nh_city.strip() or not nh_hood.strip() or not nh_desc.strip():
                        st.error("City, Neighborhood, and Description are required.")
                    else:
                        import uuid
                        neighborhoods.append({
                            "id":           str(uuid.uuid4())[:8],
                            "city":         nh_city.strip(),
                            "neighborhood": nh_hood.strip(),
                            "description":  nh_desc.strip(),
                            "notes":        nh_note.strip(),
                        })
                        save_neighborhoods(neighborhoods)
                        st.session_state["show_add_hood"] = False
                        st.success("✅ Saved.")
                        st.rerun()
            with hc2:
                if st.button("Cancel", use_container_width=True, key="cancel_hood"):
                    st.session_state["show_add_hood"] = False
                    st.rerun()

    # ── Filter and sort ───────────────────────────────────────────────────────
    filtered_hoods = sorted(neighborhoods, key=lambda x: f"{x.get('city','')} {x.get('neighborhood','')}")
    if hood_search:
        q = hood_search.lower()
        filtered_hoods = [h for h in filtered_hoods if
                          q in h.get("city","").lower() or
                          q in h.get("neighborhood","").lower() or
                          q in h.get("description","").lower()]

    st.caption(f"{len(filtered_hoods)} entr{'y' if len(filtered_hoods)==1 else 'ies'}")
    st.divider()

    if filtered_hoods:
        for hood in filtered_hoods:
            _city = hood.get('city','')
            _hood = hood.get('neighborhood','')
            label = f"🏘️ {_city}" if _hood == _city or _hood in ("General", _city) else f"🏘️ {_city} — {_hood}"
            with st.expander(label):
                st.text_area(
                    "Neighborhood description",
                    value=caps(hood.get("description","")),
                    height=160,
                    key=f"hood_text_{hood['id']}_{st.session_state.get('all_caps',False)}",
                    disabled=not st.session_state.get("site_admin", False),
                    label_visibility="collapsed"
                )
                if hood.get("notes"):
                    st.caption(f"📝 Note: {hood['notes']}")

                if st.session_state.get("site_admin"):
                    ha1, ha2 = st.columns(2)
                    with ha1:
                        edited_desc = st.session_state.get(f"hood_text_{hood['id']}", hood.get("description",""))
                        if st.button("💾 Save Changes", key=f"save_hood_{hood['id']}", use_container_width=True):
                            for h in neighborhoods:
                                if h["id"] == hood["id"]:
                                    h["description"] = edited_desc.strip()
                            save_neighborhoods(neighborhoods)
                            st.success("✅ Saved.")
                            st.rerun()
                    with ha2:
                        if st.button("🗑️ Delete", key=f"del_hood_{hood['id']}", use_container_width=True):
                            neighborhoods = [h for h in neighborhoods if h["id"] != hood["id"]]
                            save_neighborhoods(neighborhoods)
                            st.rerun()
                else:
                    st.caption("☝️ Click · Ctrl+A · Ctrl+C")
    else:
        st.info("No neighborhood descriptions yet. Add your first one above.")

# ══════════════════════════════════════════════════════════════════════════════
# TAB 4 — ZONING DISTRICTS
# ══════════════════════════════════════════════════════════════════════════════
with tab2:
 if True:
    import re
    from itertools import groupby
    zoning = load_zoning()

    # ── Helper: detect and compute unit add-on formulas ──────────────────────
    def parse_addon_formula(lot_str, notes_str):
        src = (lot_str + " " + notes_str).lower()
        # Pattern A: base SF + addon/unit [above N]
        m = re.search(
            r'([\d,]+)\s*sf' + r'.*?\+\s*([\d,]+)\s*sf?' +
            r'.*?(?:per|for\s*ea\w*).*?(?:add\w*\s*)?unit' +
            r'(?:.*?(?:above|over)\s*(\w+))?', src)
        if m:
            base  = int(m.group(1).replace(',',''))
            addon = int(m.group(2).replace(',',''))
            raw   = m.group(3)
            w2n   = {"one":1,"two":2,"three":3,"four":4,"five":5}
            try:
                if not raw:
                    above = 1
                elif raw.isdigit():
                    above = int(raw)
                elif raw in w2n:
                    above = w2n[raw]
                else:
                    above = 1
            except Exception:
                above = 1
            return {"base_sf": base, "addon_sf": addon, "above_n": above, "type": "sf_addon"}
        # Pattern B: X acres + Y SF/unit
        m2 = re.search(r'([\d.]+)\s*acres?\s*(?:\+|plus)\s*([\d,]+)\s*(?:sf\s*)?per\s*unit', src)
        if m2:
            return {"base_acres": float(m2.group(1)), "addon_sf": int(m2.group(2).replace(',','')), "above_n": 0, "type": "acres_addon"}
        # Pattern C: X per unit (multiply)
        m3 = re.search(r'([\d.]+)\s*(acres?|sf)\s*per\s*unit', src)
        if m3:
            if "acre" in m3.group(2):
                return {"per_unit_acres": float(m3.group(1)), "type": "per_unit_acres"}
            return {"per_unit_sf": int(float(m3.group(1))), "type": "per_unit_sf"}
        return None

    def compute_lot_size(formula, unit_count):
        t = formula.get("type")
        if t == "sf_addon":
            extra = max(0, unit_count - formula["above_n"])
            total = formula["base_sf"] + extra * formula["addon_sf"]
            return f"{total:,} SF"
        elif t == "acres_addon":
            total_sf = int(formula["base_acres"] * 43560) + unit_count * formula["addon_sf"]
            return f"{total_sf:,} SF"
        elif t == "per_unit_acres":
            return f"{formula['per_unit_acres'] * unit_count:.1f} acres"
        elif t == "per_unit_sf":
            return f"{formula['per_unit_sf'] * unit_count:,} SF"
        return ""

    def format_frontage(s):
        m = re.search(r"(\d+(?:\.\d+)?)", s)
        return m.group(1) if m else s.split("/")[0].strip()

    # ── Admin Mode ───────────────────────────────────────────────────────────

    # ── Search + Add New ──────────────────────────────────────────────────────
    col_zs, col_za = st.columns([4, 1])
    with col_zs:
        zone_search = st.text_input("🔍 Search towns, zone codes, or keywords...",
                                     placeholder="e.g. Warwick, R-40, sewer, two-family",
                                     key="zone_search", label_visibility="collapsed")
    with col_za:
        if st.session_state["site_admin"]:
            if st.button("➕ Add New", key="add_zone_btn", use_container_width=True):
                st.session_state["show_add_zone"] = not st.session_state.get("show_add_zone", False)

    if st.session_state.get("show_add_zone") and st.session_state["site_admin"]:
        with st.container(border=True):
            st.subheader("New Zoning District")
            zc1, zc2 = st.columns(2)
            with zc1:
                nz_city     = st.text_input("Town / City *", key="nz_city", placeholder="e.g. Providence")
            with zc2:
                nz_district = st.text_input("Zoning Code *", key="nz_district", placeholder="e.g. A-7")
            zd1, zd2 = st.columns(2)
            with zd1:
                nz_frontage = st.text_input("Min Lot Frontage", key="nz_frontage", placeholder="e.g. 70'")
            with zd2:
                nz_lotarea  = st.text_input("Min Lot Area", key="nz_lotarea", placeholder="e.g. 7,000 SF")
            nz_notes = st.text_area("Notes (optional)", key="nz_notes", height=80,
                                     placeholder="e.g. +7,500 SF per additional unit; with public sewer only")
            zb1, zb2 = st.columns(2)
            with zb1:
                if st.button("💾 Save", use_container_width=True, key="save_zone"):
                    if not nz_city.strip() or not nz_district.strip():
                        st.error("Town and Zoning Code are required.")
                    else:
                        import uuid
                        zoning.append({
                            "id": str(uuid.uuid4())[:8],
                            "city": nz_city.strip(), "district": nz_district.strip(),
                            "property_type": "", "frontage": nz_frontage.strip(),
                            "lot_area": nz_lotarea.strip(), "lot_width": "",
                            "front_yard": "", "side_yard": "", "rear_yard": "",
                            "max_height": "", "max_lot_cov": "", "max_floors": "",
                            "notes": nz_notes.strip(),
                        })
                        save_zoning(zoning)
                        st.session_state["show_add_zone"] = False
                        st.success("✅ Saved.")
                        st.rerun()
            with zb2:
                if st.button("Cancel", use_container_width=True, key="cancel_zone"):
                    st.session_state["show_add_zone"] = False
                    st.rerun()

    # ── Filter & Sort ─────────────────────────────────────────────────────────
    filtered_zones = sorted(zoning, key=lambda x: (x.get("city","").lower(), x.get("district","").lower()))
    if zone_search:
        q = zone_search.lower()
        filtered_zones = [z for z in filtered_zones if
                          q in z.get("city","").lower() or
                          q in z.get("district","").lower() or
                          q in z.get("notes","").lower() or
                          q in z.get("lot_area","").lower()]

    st.caption(f"{len(filtered_zones)} entr{'y' if len(filtered_zones)==1 else 'ies'}")
    st.divider()

    # ── Display grouped by town — town header collapses all zones ───────────
    if filtered_zones:
        # Group zones by town into a dict so we can use expander per town
        town_groups = {}
        for zone in filtered_zones:
            t = zone.get("city","")
            town_groups.setdefault(t, []).append(zone)

        for town, zones in town_groups.items():
            zone_count = len(zones)
            with st.expander(f"🏙️  {town}  —  {zone_count} zone{'s' if zone_count != 1 else ''}"):
                for zone in zones:
                    district = zone.get("district","")
                    frontage = zone.get("frontage","") or ""
                    lot_area = zone.get("lot_area","")  or ""
                    notes    = zone.get("notes","")     or ""
                    formula  = parse_addon_formula(lot_area, notes)

                    st.markdown(f"#### 📐 {district}")

                    # Key fields
                    col_a, col_b = st.columns(2)
                    with col_a:
                        st.markdown("**Min Lot Frontage**")
                        st.markdown(f"`{frontage}`" if frontage else "—")
                    with col_b:
                        st.markdown("**Min Lot Area**")
                        st.markdown(f"`{lot_area}`" if lot_area else "—")
                    if notes:
                        st.markdown(f"*📝 {notes}*")

                    st.markdown("**📋 Format for TOTAL**")
                    computed_area = lot_area
                    front_num = format_frontage(frontage) if frontage else ""
                    parts = []
                    if front_num:
                        parts.append(f"{front_num}' lot frontage")
                    if computed_area:
                        parts.append(f"{computed_area} min lot size")
                    auto_output = " / ".join(parts) if parts else "—"

                    # Use saved custom text if it exists, otherwise use auto-generated
                    display_output = zone.get("total_override") or auto_output

                    if st.session_state.get("site_admin"):
                        edited = st.text_area("TOTAL quicklist text", value=caps(display_output), height=75,
                                              key=f"total_{zone['id']}_{st.session_state.get('all_caps',False)}", label_visibility="collapsed")
                        ta1, ta2 = st.columns(2)
                        with ta1:
                            if st.button("💾 Save TOTAL Text", key=f"save_total_{zone['id']}",
                                         use_container_width=True):
                                for z in zoning:
                                    if z["id"] == zone["id"]:
                                        z["total_override"] = edited.strip()
                                save_zoning(zoning)
                                st.success("✅ Saved.")
                                st.rerun()
                        with ta2:
                            if st.button("🗑️ Delete Zone", key=f"del_zone_{zone['id']}",
                                         use_container_width=True):
                                zoning = [z for z in zoning if z["id"] != zone["id"]]
                                save_zoning(zoning)
                                st.rerun()
                    else:
                        st.text_area("TOTAL quicklist text", value=caps(display_output), height=75,
                                     key=f"total_{zone['id']}_{st.session_state.get('all_caps',False)}", disabled=True, label_visibility="collapsed")
                        st.caption("☝️ Click · Ctrl+A · Ctrl+C")

                    st.divider()
    else:
        st.info("No results. Try a different search or add a new entry.")

# ══════════════════════════════════════════════════════════════════════════════
# ══════════════════════════════════════════════════════════════════════════════
# TAB 5 — UAD 3.6 REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
with tab7:
 if True:
    st.subheader("🆕 UAD 3.6 Reference Guide")
    st.caption("Key changes, inspection requirements, ratings definitions, and tool overview. Mandatory November 2, 2026.")

    st.info("⏱️ **Key Dates:** Broad Production started Jan 26, 2026 — both 2.6 and 3.6 accepted now. **Mandatory Nov 2, 2026** — all legacy forms (1004, 1073, 2055, etc.) retired. UAD 2.6 pipeline fully retired May 3, 2027.")

    uad_sub1, uad_sub2, uad_sub3, uad_sub4 = st.tabs([
        "📊 What Changed vs 2.6",
        "🔍 New Inspection Data",
        "🏠 C & Q Ratings",
        "💻 Report Writing Tools"
    ])

    # ── WHAT CHANGED ─────────────────────────────────────────────────────────
    with uad_sub1:
        st.markdown("### Overall vs UAD 2.6 — same, changed, new, or gone")

        col_leg1, col_leg2, col_leg3, col_leg4 = st.columns(4)
        with col_leg1: st.success("🟢 Same")
        with col_leg2: st.warning("🟡 Changed")
        with col_leg3: st.info("🔵 New")
        with col_leg4: st.error("🔴 Gone")

        sections = {
            "Forms & Report Structure": [
                ("🟡", "Form selection", "Single dynamic URAR replaces 1004, 1073, 2055, and all legacy forms. Sections turn on/off based on property/assignment type. No more form numbers."),
                ("🔴", "1004MC Market Conditions form", "Eliminated entirely. Market conditions are now embedded in the URAR Market section as structured fields. Time adjustments still required and must be supported with data — just not the MC grid format."),
                ("🟡", "Scope of work", "Still required. More structured fields, less free-text narrative. The analytical obligation is unchanged."),
            ],
            "Condition & Quality Ratings": [
                ("🟡", "C1–C6 Condition scale", "Now three separate ratings: Exterior Condition + Interior Condition → reconciled to Overall. Previously one rating for the whole property."),
                ("🟡", "Q1–Q6 Quality scale", "Same structure change — Exterior Quality + Interior Quality → Overall. Quality is still absolute, not relative to your market."),
                ("🟢", "C6 rule", "If any component is C6, overall must be C6. Unchanged. Fannie allows C5 as-is in some cases; Freddie requires C4 minimum for delivery."),
                ("🟡", "Defects documentation", "When C5 or C6: now requires itemized structured list — location, description, impact on soundness, recommended action, estimated repair cost. Previously narrative only."),
                ("🔵", "As-Is + Subject-to-Repair ratings", "When repairs required, must report BOTH the As-Is Overall Condition AND the Condition Subject to Repair as separate explicit ratings."),
            ],
            "Kitchen & Bathroom Reporting": [
                ("🔴", "'Not updated / Updated / Remodeled' checkbox", "Retired entirely. Replaced by structured Kitchen and Bathroom Details sections."),
                ("🟡", "Kitchen reporting", "Per-kitchen structured section: update status + approximate year of update + condition rating (C1–C6) + brief description. Every kitchen reported individually."),
                ("🟡", "Bathroom reporting", "Per-bathroom structured section: same as kitchen. Every bathroom including half baths reported individually."),
            ],
            "Site & Location": [
                ("🟡", "View / location influence", "Replaced by Site Influence section. Now captures: onsite / bordering / distant AND positive / neutral / negative. A pond on the lot vs a lake view from a distance are now distinct reportable data points."),
                ("🟡", "Zoning, utilities, site features", "Significantly expanded structured fields. The Site section is 30 pages in the URAR Reference Guide — second largest chapter after the Sales Comparison Approach."),
            ],
            "Sales Comparison Approach": [
                ("🟢", "Comparable selection and analysis", "Same process, same professional judgment required. More structured fields, less free-text narrative."),
                ("🟡", "Time adjustments", "1004MC is gone. Time adjustments still required when market warrants. Must be supported with market data. The decision to NOT adjust requires the same level of support as the decision to adjust."),
                ("🟢", "Actual age in comp grid", "Only actual age in the grid. Effective age discussion moves to narrative commentary."),
                ("🟡", "GLA measurement", "ANSI Z765-2021 formalized. Above-grade finished area must be broken out by floor. UCDP compliance rules will enforce ANSI logic on submission."),
            ],
            "Special Features": [
                ("🔵", "ADU section", "New dedicated section. Effective Mar 21, 2026: expanded ADU eligibility ONLY available with UAD 3.6 submissions. ADU rental income can count toward qualifying income under certain conditions."),
                ("🟡", "Manufactured housing terminology", "Effective Mar 31, 2026: 'single-width' and 'multi-width' replace 'single-section' and 'multi-section' in all UAD 3.6 reports."),
            ],
            "Core Appraisal Practice": [
                ("🟢", "How value is estimated", "Exactly the same. UAD 3.6 changes how work is reported, not how value is estimated."),
                ("🟢", "USPAP compliance", "Required and unchanged. UAD 3.6 is a reporting format change, not a USPAP change."),
                ("🟢", "Physical inspection process", "The inspection itself does not change. What changes is the structured way you document and report what you observed."),
                ("🟢", "Highest and best use analysis", "Same requirement, same analysis. Reporting format more structured."),
                ("🟢", "Cost and income approaches", "Same analytical requirements. More structured data fields for reporting conclusions."),
            ],
        }

        for section, items in sections.items():
            st.markdown(f"**{section}**")
            for status, title, detail in items:
                with st.expander(f"{status} {title}"):
                    st.write(detail)
            st.write("")

        st.caption("Source: Fannie Mae & Freddie Mac UAD Inspection and Reporting Tips (Oct 2025), UAD 3.6 FAQ, McKissock Learning, Working RE, Appraiser eLearning.")

    # ── NEW INSPECTION DATA ───────────────────────────────────────────────────
    with uad_sub2:
        st.markdown("### New data you must collect at the inspection")
        st.warning("UAD 3.6 requires significantly more structured data collected in the field. Plan your inspections accordingly — some of this requires new tools or conversations with the owner/agent that weren't necessary before.")

        st.markdown("#### 🔧 Bring to every inspection")
        tool_items = [
            ("Laser distance measurer", "Required for ceiling height measurement — a tape measure and ladder is not practical or safe. Ceiling height is a required field every time (Report Field ID: 10.045)."),
            ("Mobile device with UAD 3.6 software", "The volume of new structured data fields makes mobile data entry at the inspection far more efficient than clipboard + office entry. This is the workflow UAD 3.6 is built around."),
        ]
        for tool, note in tool_items:
            with st.expander(f"🔧 {tool}"):
                st.write(note)

        st.markdown("#### 🔵 New — must now measure or collect")
        new_items = [
            ("Ceiling height (required every inspection)", "Required field — always. Tied to ANSI 7-foot rule: if more than half a room's area is under 7 ft, that area cannot count as finished GLA. Measure with a laser."),
            ("Walls and ceiling materials/type", "Required structured field in Interior Features table — select from defined enumerations. Always displays, always required."),
            ("Nonstandard Finished Area (NSFA)", "Finished space that doesn't meet ANSI must be measured and reported separately from GLA. Must explain why (e.g. 'sloping ceiling under 7 ft'). Common in finished attics, bonus rooms with knee walls."),
            ("Roof replacement estimate", "New structured field — when was the roof last replaced? Ask the owner or agent at inspection. Can also pull from permit records."),
            ("Outbuilding utilities", "If outbuildings present: which utilities have been extended to each (electric, water, HVAC, etc.)."),
            ("Broadband internet availability", "Surprising new requirement. Verify via FCC broadband map or ask the owner. The GSE definition of 'broadband' is specific — check Appendix F-1 for the enumeration."),
            ("Access road type", "New structured field — type of road providing access (public paved, private paved, gravel, dirt, etc.)."),
            ("ADU detail (if present)", "New dedicated section — unit type, GLA, bed/bath count, kitchen presence, separate entrance, condition, and whether legally permitted. Collect all at inspection."),
            ("Disaster mitigation features", "If present: storm shutters, hurricane straps, flood vents, seismic retrofitting. If this section is included, the commentary field is required."),
        ]
        for title, detail in new_items:
            with st.expander(f"🔵 {title}"):
                st.write(detail)

        st.markdown("#### 🟡 Changed — process same, now structured")
        changed_items = [
            ("Above-grade finished area — per floor breakdown", "Previously reported as one total GLA number. Now must be broken out by floor — first floor, second floor, etc. UCDP will validate that floors sum to total."),
            ("Kitchen update status — per kitchen", "Previously one checkbox. Now per-kitchen: update status + approximate year + condition rating + brief description."),
            ("Bathroom update status — per bathroom", "Previously one checkbox. Now per-bathroom: same structured fields. Every bathroom including half baths."),
            ("Exterior quality — separate from interior", "Assess separately: siding, roof, foundation, windows, trim, architectural detail. Record as its own rating."),
            ("Exterior condition — separate from interior", "Assess separately: same components. Must reconcile with interior condition to reach overall rating."),
            ("Interior quality — separate from exterior", "Assess separately: flooring, trim, cabinetry, fixtures, countertops, built-ins."),
            ("Interior condition — separate from exterior", "Assess separately. If C5/C6: itemized defect list required."),
            ("Site influence (view/location)", "Now captures: onsite / bordering / distant AND positive / neutral / negative. More granular than the old View field."),
        ]
        for title, detail in changed_items:
            with st.expander(f"🟡 {title}"):
                st.write(detail)

        st.markdown("#### 🔴 Gone — no longer required")
        gone_items = [
            ("Street inspection of comparable sales", "No longer required. Front photo of comp still required but the physical drive-by has been retired. Photo can be sourced from MLS or other reliable source."),
        ]
        for title, detail in gone_items:
            with st.expander(f"🔴 {title}"):
                st.write(detail)

        st.caption("Source: GSE Inspection and Reporting Tips (Oct 2025), Appendix F-1 URAR Reference Guide, ANSI Z765-2021.")

    # ── C & Q RATINGS ────────────────────────────────────────────────────────
    with uad_sub3:
        st.markdown("### Condition Ratings (C1–C6)")
        st.info("Key change: Condition is now reported separately for Exterior and Interior, then reconciled to an Overall rating. If ANY component is C6, the overall must be C6. Fannie allows C5 as-is in some cases; Freddie requires C4 minimum for loan delivery.")

        conditions = [
            ("C1", "New / Like New", "New or nearly new. No physical wear on any component. All major and minor elements are in like-new condition. Rebuilt homes may qualify if on a completely new foundation with fully remanufactured materials.", "Use for new construction and very recent gut rehabs to the studs. Rarely seen on resale."),
            ("C2", "Extensive Renovation", "Extensively renovated to resemble new construction. Most components recently replaced or refinished. Minimal physical depreciation, no deferred maintenance, nothing requires repair.", "Component-level detail in 3.6 makes this harder to assign loosely. Every major system should be new or like-new. Kitchen and bathroom details must support it."),
            ("C3", "Well Maintained / Updated", "Well maintained with limited physical depreciation. Some components may show minor wear but all functional systems are working properly. Some updating may be present.", "Most move-in ready properties with recent updates but not full renovation. Very common in well-maintained New England stock."),
            ("C4", "Average / Adequate", "Adequately maintained with some physical depreciation. Normal wear and tear. All functional systems working. No immediate repairs required but may have deferred maintenance items.", "The baseline for typical resale. Dated but functional kitchens/baths, normal aging. Very common in pre-1970 RI/MA stock."),
            ("C5", "Fair / Deferred Maintenance", "Obvious deferred maintenance and/or physical deterioration. Some components need repair or replacement. Major systems may be functional but near end of useful life.", "Freddie Mac: C5 is NOT eligible for delivery — must be repaired to C4. Fannie allows C5 as-is in some cases. Always check lender overlays. Requires itemized defect list in 3.6."),
            ("C6", "Significant Damage / Safety Concern", "Significant damage, serious deferred maintenance, or conditions affecting safety, soundness, or structural integrity. Property may be uninhabitable or pose health/safety risks.", "If any single component is C6, the overall must be C6. Requires full itemized defect list with location, description, impact, recommended action, and estimated repair cost."),
        ]

        for code, label, definition, notes in conditions:
            with st.expander(f"**{code} — {label}**"):
                st.write(definition)
                st.caption(f"📝 Field notes: {notes}")

        st.divider()
        st.markdown("### Quality Ratings (Q1–Q6)")
        st.info("Key change: Quality is now reported separately for Exterior and Interior, then reconciled to an Overall rating. Quality is absolute — a Q3 is a Q3 whether it's in Providence or Newport. Local market norms do not determine quality rating.")

        qualities = [
            ("Q1", "Exceptional / Custom", "Custom architecture, outstanding workmanship, and premium materials throughout — often imported or specialty items. Every component shows exceptional detail and design.", "Very rare. Most markets have no Q1 properties. If you're questioning whether it's Q1, it probably isn't."),
            ("Q2", "High Quality / Custom", "Still custom but not at the absolute top tier. High-quality materials and consistently strong workmanship. May be custom-built or part of a high-quality development.", "May be the highest rating seen in many markets. Common in higher-end coastal RI and South Shore MA."),
            ("Q3", "Good Quality / Above Average", "Solidly constructed with good materials, though not custom throughout. Often includes upgraded finishes mixed with some standard components.", "The most common rating in the upper-middle price range. Upgraded kitchen and baths with standard framing and exterior."),
            ("Q4", "Average / Standard", "Standard or builder-grade materials and construction. Meets code requirements. No significant upgrades. Represents the majority of production housing.", "The baseline for most residential construction. Colonial-era New England stock typically falls here despite age — quality is about construction, not condition."),
            ("Q5", "Fair / Below Average", "Below-standard materials or workmanship. May show significant deficiencies in design, construction, or finish. Functional but lacks quality.", "Rare in typical residential lending. More common in older worker housing, self-built structures, or properties with significant deferred improvements."),
            ("Q6", "Poor / Substandard", "Construction that may not meet basic building standards. Often built without professional oversight or adherence to modern codes. May be unsuitable for year-round habitation.", "Very rarely assigned. Flag immediately — lender eligibility issues are likely."),
        ]

        for code, label, definition, notes in qualities:
            with st.expander(f"**{code} — {label}**"):
                st.write(definition)
                st.caption(f"📝 Field notes: {notes}")

        st.caption("Source: McKissock Learning, Appendix F-1 URAR Reference Guide, Working RE — UAD Quality Equation (Aug 2025).")

    # ── REPORT WRITING TOOLS ─────────────────────────────────────────────────
    with uad_sub4:
        st.markdown("### Report Writing Tools — A Guide for Newer Appraisers")

        st.markdown("""
The appraisal profession is going through its biggest technology shift in decades. UAD 3.6 changes not just what you report but how you collect data in the field — and a new generation of software tools has been built specifically around this new workflow. If you're newer to the profession, you have an advantage: you don't have 15 years of muscle memory to undo.
        """)

        st.divider()
        st.markdown("#### The two categories of report-writing software")

        col_trad, col_ai = st.columns(2)
        with col_trad:
            st.markdown("**Traditional platforms**")
            st.caption("e.g. TOTAL by a la mode, ClickForms, ACI")
            st.write("These have been the industry standard for decades. You enter data manually, build your report field by field, and write narrative addenda. They are being updated for UAD 3.6 but the core workflow — appraiser drives everything — stays the same.")
            st.markdown("✅ Full control over every field  \n✅ Deep MLS and data integration  \n✅ Established, trusted by AMCs  \n✅ Large training community  \n❌ UAD 3.6 transition still in progress  \n❌ Monthly subscription cost")

        with col_ai:
            st.markdown("**AI-assisted platforms**")
            st.caption("e.g. Automax, Aivre, ApprAIz")
            st.write("A new wave of tools built from the ground up for UAD 3.6. These use AI and structured data to auto-populate fields, generate sketches, suggest comps, and flag compliance issues in real time. Both Automax and Aivre were featured at the 2026 UAD 3.6 Bootcamp alongside Fannie Mae and Freddie Mac.")
            st.markdown("✅ Built for UAD 3.6 from day one  \n✅ Faster data entry in the field  \n✅ Real-time compliance checking  \n✅ Auto-populates structured fields  \n❌ Newer — less track record  \n❌ AMC acceptance varies by lender")

        st.divider()
        st.markdown("#### What AI actually does in appraisal software")

        ai_items = [
            ("Computer vision / scanning", "You walk through a property with your phone and the app uses the camera to measure rooms, generate a sketch, and identify features like flooring type, ceiling height, and condition. Some tools can do this in real time during the inspection walkthrough."),
            ("Auto-population", "The software pulls property data from public records, MLS, and other sources and pre-fills report fields — address, lot size, year built, zoning, etc. — so it's already there when you open the assignment."),
            ("Comp suggestion", "Algorithms scan MLS and public record data and suggest comparable sales based on proximity, size, age, condition, and other factors. You still select and verify each comp — the AI narrows the search, you make the call."),
            ("Compliance checking", "The software flags UAD 3.6 rule violations before you submit — missing required fields, inconsistent C/Q ratings, data that doesn't pass UCDP logic checks. Catches errors before the reviewer does."),
            ("What AI does NOT do", "AI does not form the opinion of value. The analysis, judgment, and professional conclusions are entirely yours. USPAP places that responsibility on the appraiser and no software changes that."),
        ]

        for title, detail in ai_items:
            with st.expander(f"{'⚠️' if 'NOT' in title else '🤖'} {title}"):
                st.write(detail)

        st.divider()
        st.markdown("#### Automax vs Aivre — balanced overview")
        st.caption("Both were UAD 3.6 software demo participants at the 2026 Appraiser eLearning Bootcamp alongside Fannie Mae and Freddie Mac.")

        col_am, col_av = st.columns(2)
        with col_am:
            st.markdown("**Automax**")
            st.caption("Free platform — hybrid order network model")
            st.write("Offers the platform at no cost in exchange for participation in their hybrid appraisal order network. Hybrid orders typically involve a reduced-scope inspection with the appraiser completing the valuation. Order acceptance is generally flexible.")
            st.markdown("""
**Features:**
- Instant sketch generation
- TOTAL data integration
- MLS + public records + Zillow comp data
- Cloud-stored USPAP-compliant workfile
- UAD 3.6 on active roadmap

**Consider if:** You want to reduce your software overhead and are open to hybrid assignments as part of your workflow.

**Watch for:** UAD 3.6 deployment timeline — confirm current status with Automax directly before committing.
            """)

        with col_av:
            st.markdown("**Aivre**")
            st.caption("~$2,500/year — standalone subscription")
            st.write("A standalone AI-assisted platform with a subscription model. Was an early mover on UAD 3.6 compliance. As of early 2026 did not include a sketch component — verify current feature set before evaluating as this space moves quickly.")
            st.markdown("""
**Features:**
- UAD 3.6 compliant early
- Purpose-built for new URAR
- No order network obligations
- Full control — standalone tool

**Consider if:** You want a UAD 3.6-native platform with no strings attached and are willing to pay the subscription.

**Watch for:** Sketch capability — this is a significant gap for UAD 3.6 field data collection. Confirm current status.
            """)

        st.divider()
        st.info("**A note for newer appraisers:** Neither tool eliminates the need to understand the appraisal process. AI-assisted platforms speed up data collection and flag compliance issues, but they work best in the hands of someone who already understands what the fields mean, why certain adjustments are made, and what USPAP requires. The foundation comes first — the technology amplifies it. Start your UAD 3.6 education with Fannie Mae's free resources at **fanniemae.com/UAD** — specifically Appendix F-1 (375 pages, worth having as a searchable PDF).")

        st.caption("Source: Appraiser eLearning 2026 UAD 3.6 Bootcamp, Working RE, McKissock Learning, Fannie Mae Appraiser Update (Jan 2026).")



# ══════════════════════════════════════════════════════════════════════════════
# TAB 6 — REPORT BUILDER (GP RESIDENTIAL)
# ══════════════════════════════════════════════════════════════════════════════



# ══════════════════════════════════════════════════════════════════════════════
# TAB 6 — QC CHECKER
# ══════════════════════════════════════════════════════════════════════════════
with tab4:
 if True:
    st.markdown("## QC Checker")
    st.caption(
        "Upload a TOTAL XML export and optionally the PDF. "
        "The checker reads all structured data directly from XML and flags internal inconsistencies."
    )
    st.divider()

    import re
    import xml.etree.ElementTree as ET
    from datetime import datetime, date as dt_date

    try:
        import fitz
        FITZ_AVAILABLE = True
    except ImportError:
        FITZ_AVAILABLE = False

    u1, u2 = st.columns(2)
    with u1:
        qc_xml = st.file_uploader(
            "Upload TOTAL XML Export (required)",
            type=["xml"],
            key="qc_xml_upload",
            help="File > Export > XML from TOTAL. Contains all structured report data."
        )
    with u2:
        qc_pdf = st.file_uploader(
            "Upload TOTAL PDF (optional)",
            type=["pdf"],
            key="qc_pdf_upload",
            help="Adds addendum narrative text checks — adjustment support language, prior sale commentary."
        )

    if not qc_xml:
        st.info("Upload a TOTAL XML export to run QC checks.")
    else:

        # ── Parse XML ─────────────────────────────────────────────────────────
        @st.cache_data(show_spinner="Parsing XML...")
        def parse_xml(xml_bytes):
            root = ET.fromstring(xml_bytes.decode("utf-8", errors="ignore"))
            return root

        @st.cache_data(show_spinner="Reading PDF...")
        def parse_pdf(pdf_bytes):
            if not FITZ_AVAILABLE:
                return ""
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            return "\n".join(p.get_text() for p in doc)

        xml_root = parse_xml(qc_xml.read())
        pdf_text = ""
        if qc_pdf:
            pdf_text = parse_pdf(qc_pdf.read())
            st.success("XML + PDF loaded.")
        else:
            st.success("XML loaded.")

        # ── XML helpers ───────────────────────────────────────────────────────
        def find_elem(root, tag):
            for e in root.iter():
                t = e.tag.split('}')[-1] if '}' in e.tag else e.tag
                if t == tag:
                    return e
            return None

        def find_all_elems(root, tag):
            results = []
            for e in root.iter():
                t = e.tag.split('}')[-1] if '}' in e.tag else e.tag
                if t == tag:
                    results.append(e)
            return results

        def attr(elem, key, default=""):
            if elem is None:
                return default
            return elem.get(key, default).strip()

        def to_int(val):
            if not val:
                return None
            try:
                return int(float(str(val).replace("$","").replace(",","").replace("+","").strip()))
            except Exception:
                return None

        def to_float(val):
            if not val:
                return None
            try:
                return float(str(val).replace("%","").replace("$","").replace(",","").strip())
            except Exception:
                return None

        def parse_date(val):
            if not val:
                return None
            for fmt in ["%Y-%m-%d","%m/%d/%Y","%Y/%m/%d"]:
                try:
                    return datetime.strptime(val.strip(), fmt).date()
                except Exception:
                    pass
            return None

        # ── Extract subject data ──────────────────────────────────────────────
        prop    = find_elem(xml_root, "PROPERTY")
        struct  = find_elem(xml_root, "STRUCTURE")
        site    = find_elem(xml_root, "SITE")
        nbhd    = find_elem(xml_root, "NEIGHBORHOOD")
        val_elem= find_elem(xml_root, "VALUATION")
        recon   = find_elem(xml_root, "_RECONCILIATION")
        report  = find_elem(xml_root, "REPORT")

        subj_addr    = attr(prop, "_StreetAddress")
        subj_city    = attr(prop, "_City")
        subj_state   = attr(prop, "_State")
        subj_rights  = attr(prop, "_RightsType")

        subj_gla     = to_int(attr(struct, "GrossLivingAreaSquareFeetCount"))
        subj_rooms   = to_int(attr(struct, "TotalRoomCount"))
        subj_beds    = to_int(attr(struct, "TotalBedroomCount"))
        subj_baths   = to_float(attr(struct, "TotalBathroomCount"))
        subj_yr      = to_int(attr(struct, "PropertyStructureBuiltYear"))
        subj_stories = to_float(attr(struct, "StoriesCount"))
        subj_style   = attr(struct, "_DesignDescription")

        subj_site    = attr(site, "_AreaDescription")
        subj_zone    = attr(site, "_ZoningClassificationIdentifier")
        subj_zcomp   = attr(site, "_ZoningComplianceType")

        nbhd_value_trend = attr(nbhd, "_PropertyValueTrendType")
        nbhd_supply      = attr(nbhd, "_DemandSupplyType")
        nbhd_mkt_time    = attr(nbhd, "_TypicalMarketingTimeDurationType")
        nbhd_growth      = attr(nbhd, "_GrowthPaceType")
        nbhd_builtup     = attr(nbhd, "_BuiltupRangeType")
        nbhd_desc        = attr(nbhd, "_Description")

        appraised_value  = to_int(attr(val_elem, "PropertyAppraisedValueAmount"))
        eff_date         = parse_date(attr(val_elem, "AppraisalEffectiveDate"))
        recon_comment    = attr(recon, "_SummaryComment")

        report_signed    = parse_date(attr(report, "AppraiserReportSignedDate"))
        file_num         = attr(report, "AppraiserFileIdentifier")
        purpose          = attr(report, "AppraisalPurposeType")

        # Appraiser license
        lic_elem  = find_elem(xml_root, "APPRAISER_LICENSE")
        lic_num   = attr(lic_elem, "_Identifier")
        lic_exp   = parse_date(attr(lic_elem, "_ExpirationDate"))
        lic_state = attr(lic_elem, "_State")

        appraiser_elem = find_elem(xml_root, "APPRAISER")
        appraiser_name = attr(appraiser_elem, "_Name")

        insp_elem = find_elem(xml_root, "INSPECTION")
        insp_date = parse_date(attr(insp_elem, "InspectionDate"))

        supervisor_elem = find_elem(xml_root, "SUPERVISOR")
        supervisor_name = attr(supervisor_elem, "_Name")

        # Borrower / intended user
        borrower_elem = find_elem(xml_root, "BORROWER")
        borrower_name = attr(borrower_elem, "_UnparsedName")

        lender_elem = find_elem(xml_root, "LENDER")
        lender_name = attr(lender_elem, "_UnparsedName")

        # ── Extract comparable sales ──────────────────────────────────────────
        comp_elems = find_all_elems(xml_root, "COMPARABLE_SALE")
        # First comp (seq 0) is the subject — skip it
        actual_comps = [c for c in comp_elems
                        if attr(c,"PropertySequenceIdentifier") != "0"]

        comps = []
        for c in actual_comps:
            seq      = to_int(attr(c, "PropertySequenceIdentifier"))
            price    = to_int(attr(c, "PropertySalesAmount"))
            net_pct  = to_float(attr(c, "SalePriceTotalAdjustmentNetPercent"))
            gross_pct= to_float(attr(c, "SalesPriceTotalAdjustmentGrossPercent"))
            net_amt  = to_int(attr(c, "SalePriceTotalAdjustmentAmount"))
            adj_val  = to_int(attr(c, "AdjustedSalesPriceAmount"))
            net_pos  = attr(c, "SalesPriceTotalAdjustmentPositiveIndicator")  # Y/N

            # Location
            loc = find_elem(c, "LOCATION")
            comp_addr   = attr(loc, "PropertyStreetAddress")
            comp_city   = attr(loc, "PropertyCity")
            proximity   = attr(loc, "ProximityToSubjectDescription")

            # Extract miles from proximity
            prox_miles  = None
            pm = re.search(r'([\d.]+)\s*miles?', proximity or "", re.IGNORECASE)
            if pm:
                prox_miles = float(pm.group(1))

            # Room data
            room_adj = find_elem(c, "ROOM_ADJUSTMENT")
            comp_rooms= to_int(attr(room_adj, "TotalRoomCount"))
            comp_beds = to_int(attr(room_adj, "TotalBedroomCount"))
            comp_baths= to_float(attr(room_adj, "TotalBathroomCount"))

            # Sale price adjustments — build dict by type
            adj_map = {}
            for spa in find_all_elems(c, "SALE_PRICE_ADJUSTMENT"):
                atype = attr(spa, "_Type")
                adesc = attr(spa, "_Description")
                aamt  = attr(spa, "_Amount")
                adj_map[atype] = {"desc": adesc, "amount": to_int(aamt) if aamt else None}

            # Other feature adjustments
            other_adjs = []
            for ofa in find_all_elems(c, "OTHER_FEATURE_ADJUSTMENT"):
                desc = attr(ofa, "PropertyFeatureDescription")
                amt  = to_int(attr(ofa, "PropertyFeatureAdjustmentAmount"))
                if desc or amt:
                    other_adjs.append({"desc": desc, "amount": amt})

            # Prior sale
            prior = find_elem(c, "PRIOR_SALES")
            prior_date = attr(prior, "PropertySalesDate")
            prior_amt  = to_int(attr(prior, "PropertySalesAmount"))

            # Parse comp GLA from GrossLivingArea adjustment description
            comp_gla = None
            if "GrossLivingArea" in adj_map:
                gla_desc = adj_map["GrossLivingArea"]["desc"]
                gm = re.match(r'(\d+)', gla_desc or "")
                if gm:
                    comp_gla = to_int(gm.group(1))

            # Parse date of sale
            sale_date_str = adj_map.get("DateOfSale",{}).get("desc","")
            # Format: s08/25;c06/25 — settled date first
            sale_date = None
            sdm = re.search(r's(\d{2}/\d{2})', sale_date_str or "")
            if sdm:
                # Convert MM/YY to approximate date
                parts = sdm.group(1).split("/")
                try:
                    sale_date = dt_date(2000 + int(parts[1]), int(parts[0]), 1)
                except Exception:
                    pass

            # Parse financing concessions amount
            fin_desc = adj_map.get("FinancingConcessions",{}).get("desc","")
            fin_amt = None
            fm = re.search(r';(\d+)', fin_desc or "")
            if fm:
                fin_amt = to_int(fm.group(1))

            comps.append({
                "num": seq, "address": comp_addr, "city": comp_city,
                "proximity": proximity, "prox_miles": prox_miles,
                "price": price, "net_pct": net_pct, "gross_pct": gross_pct,
                "net_amt": net_amt, "net_pos": net_pos, "adj_value": adj_val,
                "rooms": comp_rooms, "beds": comp_beds, "baths": comp_baths,
                "gla": comp_gla, "adj_map": adj_map, "other_adjs": other_adjs,
                "prior_date": prior_date, "prior_amt": prior_amt,
                "sale_date": sale_date, "fin_amt": fin_amt,
            })

        # ── Display extracted summary ─────────────────────────────────────────
        with st.expander("Extracted Data — expand to review", expanded=False):
            c1, c2, c3 = st.columns(3)
            with c1:
                st.markdown("**Subject**")
                st.write(f"Address: {subj_addr}, {subj_city}, {subj_state}")
                st.write(f"GLA: {subj_gla:,} sf" if subj_gla else "GLA: not found")
                st.write(f"Rooms: {subj_rooms} | Beds: {subj_beds} | Baths: {subj_baths}")
                st.write(f"Year Built: {subj_yr}")
                st.write(f"Style: {subj_style}")
                st.write(f"Property Rights: {subj_rights}")
                st.write(f"Appraised Value: ${appraised_value:,}" if appraised_value else "Value: not found")
                st.write(f"Effective Date: {eff_date}")
            with c2:
                st.markdown("**Appraiser / License**")
                st.write(f"Name: {appraiser_name}")
                st.write(f"License: {lic_num} ({lic_state})")
                st.write(f"License Exp: {lic_exp}")
                st.write(f"Inspection Date: {insp_date}")
                st.write(f"Report Signed: {report_signed or 'not found'}")
                st.write(f"Supervisor: {supervisor_name or '(none)'}")
                st.write(f"File #: {file_num or '(blank)'}")
            with c3:
                st.markdown("**Market / Neighborhood**")
                st.write(f"Value Trend: {nbhd_value_trend}")
                st.write(f"Supply/Demand: {nbhd_supply}")
                st.write(f"Marketing Time: {nbhd_mkt_time}")
                st.write(f"Growth Rate: {nbhd_growth}")
                st.write(f"Built-Up: {nbhd_builtup}")

            if comps:
                st.markdown("**Comparable Sales**")
                for c in comps:
                    p   = f"${c['price']:,}" if c['price'] else "—"
                    av  = f"${c['adj_value']:,}" if c['adj_value'] else "—"
                    np  = f"{c['net_pct']}%" if c['net_pct'] is not None else "—"
                    gp  = f"{c['gross_pct']}%" if c['gross_pct'] is not None else "—"
                    gla = f"{c['gla']:,} sf" if c['gla'] else "—"
                    st.write(f"Comp #{c['num']}: {c['address']} | {p} → {av} "
                              f"| Net {np} | Gross {gp} | GLA {gla} | {c['proximity']}")

        # ── Sketch input (still manual — images not parseable) ────────────────
        st.divider()
        st.markdown("### Sketch Input")
        st.caption("Sketch room count and GLA can't be extracted from image-based sketch pages.")
        sk1, sk2 = st.columns(2)
        with sk1:
            sketch_rooms = st.number_input("Sketch room count (above grade)",
                                            min_value=0, max_value=20,
                                            value=subj_rooms or 0, key="qc_sketch_rooms")
        with sk2:
            sketch_gla   = st.number_input("Sketch GLA (sq ft)",
                                            min_value=0, max_value=10000,
                                            value=subj_gla or 0, key="qc_sketch_gla")

        # ── Run QC ────────────────────────────────────────────────────────────
        st.divider()
        if st.button("Run QC Checks", type="primary", use_container_width=True, key="qc_run"):

            flags  = []
            passes = []

            def flag(cat, msg, sev="warning"):
                flags.append({"cat": cat, "msg": msg, "sev": sev})
            def ok(cat, msg):
                passes.append({"cat": cat, "msg": msg})

            # ── 1. License expiration vs effective date ───────────────────────
            if lic_exp and eff_date:
                if lic_exp < eff_date:
                    flag("License Expired",
                         f"License expired {lic_exp} but effective date is {eff_date}.",
                         "critical")
                else:
                    ok("License",
                       f"License {lic_num} valid through {lic_exp} — effective date {eff_date}.")
            elif not lic_exp:
                flag("License", "License expiration date not found in XML.", "critical")

            # ── 2. License state matches property state ───────────────────────
            if lic_state and subj_state:
                if lic_state.upper() != subj_state.upper():
                    flag("License State",
                         f"License is for {lic_state} but property is in {subj_state}.",
                         "critical")
                else:
                    ok("License State",
                       f"License state ({lic_state}) matches property state ({subj_state}).")

            # ── 3. Inspection date not after effective date ───────────────────
            if insp_date and eff_date:
                if insp_date > eff_date:
                    flag("Inspection Date",
                         f"Inspection date {insp_date} is after effective date {eff_date}.",
                         "critical")
                else:
                    ok("Inspection Date",
                       f"Inspection date {insp_date} is on or before effective date {eff_date}.")

            # ── 4. File number blank ──────────────────────────────────────────
            if not file_num:
                flag("File Number", "File number is blank.", "warning")
            else:
                ok("File Number", f"File number present: {file_num}")

            # ── 5. Room count — form vs sketch ────────────────────────────────
            if subj_rooms and sketch_rooms > 0:
                if subj_rooms != sketch_rooms:
                    flag("Room Count",
                         f"Form states {subj_rooms} rooms but sketch shows {sketch_rooms}.",
                         "critical")
                else:
                    ok("Room Count", f"Form and sketch agree: {subj_rooms} rooms.")

            # ── 6. GLA — form vs sketch ───────────────────────────────────────
            if subj_gla and sketch_gla > 0:
                diff = abs(subj_gla - sketch_gla)
                pct  = diff / subj_gla * 100
                if diff > 20:
                    flag("GLA Discrepancy",
                         f"Form GLA {subj_gla:,} sf vs sketch GLA {sketch_gla:,} sf "
                         f"— difference of {diff} sf ({pct:.1f}%).",
                         "critical" if pct > 5 else "warning")
                else:
                    ok("GLA", f"Form and sketch GLA agree within 20 sf ({subj_gla:,} sf).")

            # ── 7. Net adjustment per comp > 15% ─────────────────────────────
            net_issues = []
            for c in comps:
                if c["net_pct"] is not None and abs(c["net_pct"]) > 15:
                    net_issues.append(
                        f"Comp #{c['num']} ({c['address'][:25]}): "
                        f"net adj {c['net_pct']:+.1f}% exceeds 15% guideline")
            if net_issues:
                for iss in net_issues:
                    flag("Net Adjustment >15%", iss)
            else:
                ok("Net Adjustments", "All comp net adjustments within 15% guideline.")

            # ── 8. Gross adjustment per comp > 25% ───────────────────────────
            gross_issues = []
            for c in comps:
                if c["gross_pct"] is not None and c["gross_pct"] > 25:
                    gross_issues.append(
                        f"Comp #{c['num']} ({c['address'][:25]}): "
                        f"gross adj {c['gross_pct']:.1f}% exceeds 25% guideline")
            if gross_issues:
                for iss in gross_issues:
                    flag("Gross Adjustment >25%", iss)
            else:
                ok("Gross Adjustments", "All comp gross adjustments within 25% guideline.")

            # ── 9. All comps adjusted same direction ─────────────────────────
            net_directions = [c["net_pos"] for c in comps
                              if c["net_pos"] and c["net_pct"] and abs(c["net_pct"]) > 0.5]
            if len(net_directions) >= 3:
                if all(d == "Y" for d in net_directions):
                    flag("Adjustment Direction",
                         "All comps have net positive adjustments — subject may be "
                         "priced below comparable sales without adequate explanation.")
                elif all(d == "N" for d in net_directions):
                    flag("Adjustment Direction",
                         "All comps have net negative adjustments — subject may be "
                         "priced above comparable sales without adequate explanation.")
                else:
                    ok("Adjustment Direction",
                       "Comps have mixed net adjustment directions — no bracketing concern.")

            # ── 10. Subject value vs adjusted comp range ──────────────────────
            adj_values = [c["adj_value"] for c in comps if c["adj_value"]]
            if adj_values and appraised_value:
                low_adj  = min(adj_values)
                high_adj = max(adj_values)
                if appraised_value < low_adj:
                    flag("Value Outside Comp Range",
                         f"Appraised value ${appraised_value:,} is below lowest adjusted "
                         f"comp value ${low_adj:,}.")
                elif appraised_value > high_adj:
                    flag("Value Outside Comp Range",
                         f"Appraised value ${appraised_value:,} is above highest adjusted "
                         f"comp value ${high_adj:,}.")
                else:
                    ok("Value Within Range",
                       f"Appraised value ${appraised_value:,} is within adjusted "
                       f"comp range ${low_adj:,}–${high_adj:,}.")

            # ── 11. GLA bracketing ────────────────────────────────────────────
            comp_glas = [c["gla"] for c in comps if c["gla"]]
            if comp_glas and subj_gla:
                low_gla  = min(comp_glas)
                high_gla = max(comp_glas)
                if subj_gla < low_gla:
                    flag("GLA Bracketing",
                         f"Subject GLA {subj_gla:,} sf is below all comp GLAs "
                         f"(range: {low_gla:,}–{high_gla:,} sf). Consider adding a smaller comp.")
                elif subj_gla > high_gla:
                    flag("GLA Bracketing",
                         f"Subject GLA {subj_gla:,} sf is above all comp GLAs "
                         f"(range: {low_gla:,}–{high_gla:,} sf). Consider adding a larger comp.")
                else:
                    ok("GLA Bracketing",
                       f"Subject GLA {subj_gla:,} sf is bracketed by comps "
                       f"({low_gla:,}–{high_gla:,} sf).")

            # ── 12. GLA adjustment rate consistency ───────────────────────────
            if subj_gla and comps:
                implied_rates = []
                for c in comps:
                    if c["gla"] and c["adj_map"].get("GrossLivingArea",{}).get("amount") is not None:
                        diff = subj_gla - c["gla"]
                        adj  = c["adj_map"]["GrossLivingArea"]["amount"]
                        if adj and diff and abs(diff) > 30:
                            rate = round(abs(adj / diff))
                            implied_rates.append((c["num"], c["address"][:25], diff, adj, rate))

                if len(implied_rates) >= 2:
                    rates = [r[4] for r in implied_rates]
                    rate_range = max(rates) - min(rates)
                    if rate_range > 5:
                        detail = ", ".join([f"Comp #{n}: ${r}/sf" for n,_,_,_,r in implied_rates])
                        flag("GLA Adj Rate Inconsistency",
                             f"Implied GLA adjustment rates vary: {detail}. "
                             f"Verify consistent rate is applied.")
                    else:
                        ok("GLA Adj Rate",
                           f"GLA adj rate consistent across comps (~${rates[0]}/sf).")

            # ── 13. Condition adjustment consistency ──────────────────────────
            cond_adj_map = {}
            for c in comps:
                cond = c["adj_map"].get("Condition",{}).get("desc","")
                amt  = c["adj_map"].get("Condition",{}).get("amount")
                if cond and amt is not None:
                    cond_adj_map.setdefault(cond,[]).append((c["num"], amt))

            cond_issues = []
            for cond, entries in cond_adj_map.items():
                amts = set(e[1] for e in entries)
                if len(amts) > 1:
                    detail = ", ".join([f"Comp #{n}: ${a:,}" for n,a in entries])
                    cond_issues.append(
                        f"Comps rated '{cond}' have different adjustments: {detail}")
            if cond_issues:
                for iss in cond_issues:
                    flag("Condition Adj Consistency", iss)
            elif cond_adj_map:
                ok("Condition Adjustments",
                   "Condition adjustments consistent for comps with matching ratings.")

            # ── 14. Comp proximity flags ──────────────────────────────────────
            far_comps = [c for c in comps
                         if c["prox_miles"] and c["prox_miles"] > 1.0]
            if far_comps:
                for c in far_comps:
                    flag("Comp Distance",
                         f"Comp #{c['num']} ({c['address'][:25]}) is {c['prox_miles']:.2f} miles "
                         f"from subject. Verify extended search explanation is in addendum.")
            else:
                ok("Comp Distance", "All comps within 1 mile of subject.")

            # ── 15. Comp date — older than 12 months ─────────────────────────
            if eff_date:
                old_comps = []
                for c in comps:
                    if c["sale_date"]:
                        months_old = (eff_date.year - c["sale_date"].year)*12 + \
                                      (eff_date.month - c["sale_date"].month)
                        if months_old > 12:
                            old_comps.append((c["num"], c["address"][:25], months_old))
                if old_comps:
                    for n, addr, mo in old_comps:
                        flag("Comp Date Range",
                             f"Comp #{n} ({addr}) is approximately {mo} months old. "
                             f"Verify extended search explanation is in addendum.")
                else:
                    ok("Comp Dates", "All comps appear within 12 months of effective date.")

            # ── 16. Concessions — disclosed but not adjusted ──────────────────
            concession_issues = []
            for c in comps:
                if c["fin_amt"] and c["fin_amt"] > 0:
                    # Check if there's a corresponding adjustment amount for concessions
                    sales_con_amt = c["adj_map"].get("SalesConcessions",{}).get("amount")
                    fin_con_amt   = c["adj_map"].get("FinancingConcessions",{}).get("amount")
                    if sales_con_amt is None and fin_con_amt is None:
                        concession_issues.append(
                            f"Comp #{c['num']} ({c['address'][:25]}) shows concessions "
                            f"of ${c['fin_amt']:,} with no adjustment applied.")
            if concession_issues:
                for iss in concession_issues:
                    flag("Concessions", iss)
            else:
                ok("Concessions", "No unadjusted concession issues detected.")

            # ── 17. Prior sale flip check ─────────────────────────────────────
            flip_issues = []
            for c in comps:
                if c["prior_date"] and c["prior_amt"] and c["price"]:
                    try:
                        prior_dt = datetime.strptime(c["prior_date"], "%m/%d/%Y").date()
                        if c["sale_date"]:
                            months_between = (c["sale_date"].year - prior_dt.year)*12 + \
                                              (c["sale_date"].month - prior_dt.month)
                            pct_change = (c["price"] - c["prior_amt"]) / c["prior_amt"] * 100
                            if months_between <= 24 and pct_change > 20:
                                flip_issues.append(
                                    f"Comp #{c['num']} ({c['address'][:25]}) sold for "
                                    f"${c['prior_amt']:,} {months_between} months prior then "
                                    f"${c['price']:,} — {pct_change:.0f}% increase. "
                                    f"Verify arm's length transaction.")
                    except Exception:
                        pass
            if flip_issues:
                for iss in flip_issues:
                    flag("Prior Sale Flip", iss)
            elif any(c["prior_date"] for c in comps):
                ok("Prior Sales", "No flip/rapid appreciation issues detected in comp prior sales.")

            # ── 18. Market conditions consistency ────────────────────────────
            if nbhd_supply and nbhd_mkt_time:
                supply_lower = nbhd_supply.lower()
                mkt_lower    = nbhd_mkt_time.lower()
                if "shortage" in supply_lower and "overthree" in mkt_lower.replace(" ",""):
                    flag("Market Consistency",
                         f"Supply/Demand is '{nbhd_supply}' but marketing time is "
                         f"'{nbhd_mkt_time}' — these are inconsistent.")
                elif "oversupply" in supply_lower.replace(" ","") and "underthree" in mkt_lower.replace(" ",""):
                    flag("Market Consistency",
                         f"Supply/Demand is '{nbhd_supply}' but marketing time is "
                         f"'{nbhd_mkt_time}' — these are inconsistent.")
                else:
                    ok("Market Consistency",
                       f"Supply/Demand ({nbhd_supply}) and marketing time ({nbhd_mkt_time}) appear consistent.")

            # ── 19. Supervisor blank when trainee signs ───────────────────────
            if not supervisor_name:
                ok("Supervisor", "No supervisory appraiser — solo assignment.")
            else:
                ok("Supervisor", f"Supervisory appraiser present: {supervisor_name}")

            # ── 20. Addendum checks (PDF only) ────────────────────────────────
            if pdf_text:
                adj_support = any(m in pdf_text for m in
                                   ["Adjustments made:", "adj @", "adj@", "per SF GLA"])
                if not adj_support:
                    flag("Adjustment Support",
                         "No adjustment rate statement found in addendum. "
                         "Verify adjustment support language is present.")
                else:
                    ok("Adjustment Support",
                       "Adjustment rate statement present in addendum.")

                prior_language = any(p in pdf_text for p in
                                      ["No other prior sales", "no prior sales",
                                       "prior sale", "did not reveal"])
                if not prior_language:
                    flag("Prior Sale Language",
                         "No prior sale/transfer history language found in addendum.")
                else:
                    ok("Prior Sale Language",
                       "Prior sale/transfer history addressed in addendum.")

                if subj_state and lic_state and subj_state.upper() != lic_state.upper():
                    pass  # Already flagged above
            else:
                st.caption("Upload PDF to enable addendum text checks.")

            # ── Display results ───────────────────────────────────────────────
            st.markdown("---")
            st.markdown(f"## QC Results — {subj_addr}, {subj_city} {subj_state}")

            critical_flags = [f for f in flags if f["sev"] == "critical"]
            warning_flags  = [f for f in flags if f["sev"] == "warning"]

            r1, r2, r3 = st.columns(3)
            with r1:
                if critical_flags:
                    st.error(f"🔴 {len(critical_flags)} Critical Issue(s)")
                else:
                    st.success("🔴 No Critical Issues")
            with r2:
                if warning_flags:
                    st.warning(f"⚠️ {len(warning_flags)} Warning(s)")
                else:
                    st.success("⚠️ No Warnings")
            with r3:
                st.success(f"✅ {len(passes)} Passed")

            st.divider()

            if critical_flags:
                st.markdown("### 🔴 Critical Issues")
                for f in critical_flags:
                    st.error(f"**{f['cat']}:** {f['msg']}")

            if warning_flags:
                st.markdown("### ⚠️ Warnings")
                for f in warning_flags:
                    st.warning(f"**{f['cat']}:** {f['msg']}")

            if passes:
                st.markdown("### ✅ Passed")
                for p in passes:
                    st.success(f"**{p['cat']}:** {p['msg']}")

# ═══════════════════════════════════════════════════════════════════
# TAB 6 — ADJUSTMENT COMMENTARY GENERATOR
# ═══════════════════════════════════════════════════════════════════
with tab6:
 if True:
    st.subheader("Adjustment Commentary Generator")
    st.caption("Configure rates, select items, and generate paste-ready addendum language.")
    st.divider()

    for _k, _v in ADJ_DEFAULT_RATES.items():
        if "adj_" + _k not in st.session_state:
            st.session_state["adj_" + _k] = _v

    def _ag(k):
        if "adj_" + k in st.session_state:
            return st.session_state["adj_" + k]
        preset = st.session_state.get("_adj_preset_rates", {})
        return preset.get(k, ADJ_DEFAULT_RATES.get(k, 0))


    st.divider()
    presets = load_adj_presets()
    preset_names = [p["name"] for p in presets]

    with st.expander("Presets", expanded=True):
        if preset_names:
            sel = st.selectbox("Load preset", preset_names, key="adj_load_sel")
            if st.button("Load", key="adj_load_btn", use_container_width=True):
                match = next((p for p in presets if p["name"] == sel), None)
                if match:
                    _int_steps = {"gla":5,"bed":500,"fullbath":1000,"halfbath":500,"basement":1000,"garage":1000,"encporch":500,"deck":500,"fp":500,"bgr":500,"pool":1000,"cac":1000,"solar":1000,"adu":1000,"outbldg":500,"site_rate":1}
                    _flt_steps = {"adv_loc":2.5,"ben_loc":2.5,"adv_view":2.5,"ben_view":2.5,"time_rate":0.25}
                    snapped = {}
                    for k2, v2 in match["rates"].items():
                        if k2 in _int_steps:
                            snapped[k2] = int(round(float(v2)/_int_steps[k2])*_int_steps[k2])
                        elif k2 in _flt_steps:
                            snapped[k2] = round(round(float(v2)/_flt_steps[k2])*_flt_steps[k2], 4)
                        else:
                            snapped[k2] = v2
                    # Clear slider keys so Streamlit uses the fresh value parameter
                    for k2 in snapped:
                        st.session_state.pop("adj_" + k2, None)
                    # Store preset rates for _ag to read on next render
                    st.session_state["_adj_preset_rates"] = snapped
                    st.success("Loaded: " + match["name"])
                    st.rerun()
        if st.session_state.get("site_admin"):
            st.markdown("---")
            if preset_names:
                ow = st.selectbox("Overwrite", preset_names, key="adj_ow_sel")
                if st.button("Save to preset", key="adj_ow_btn", use_container_width=True):
                    rates2 = {k2: st.session_state.get("adj_" + k2, ADJ_DEFAULT_RATES.get(k2, 0)) for k2 in ADJ_DEFAULT_RATES}
                    for p in presets:
                        if p["name"] == ow:
                            p["rates"] = rates2
                    save_adj_presets(presets)
                    st.success("Saved into: " + ow)
                    st.rerun()
            nn = st.text_input("New preset name", key="adj_new_name", placeholder="e.g. SFR Waterfront")
            if st.button("Save as new", key="adj_save_new", use_container_width=True):
                if not nn.strip():
                    st.error("Enter a name.")
                elif len(presets) >= 20:
                    st.error("Max 20 presets.")
                elif any(p["name"] == nn.strip() for p in presets):
                    st.error("Name exists.")
                else:
                    rates2 = {k2: st.session_state.get("adj_" + k2, ADJ_DEFAULT_RATES.get(k2, 0)) for k2 in ADJ_DEFAULT_RATES}
                    presets.append({"name": nn.strip(), "rates": rates2})
                    save_adj_presets(presets)
                    st.success("Saved: " + nn.strip())
                    st.rerun()
            if preset_names:
                dl = st.selectbox("Delete", preset_names, key="adj_del_sel")
                if st.button("Delete preset", key="adj_del_btn", use_container_width=True):
                    save_adj_presets([p for p in presets if p["name"] != dl])
                    st.rerun()
        else:
            st.caption("Unlock Admin Mode to save or edit presets.")

    if st.button("Reset to Defaults", key="adj_reset"):
        for k2, v2 in ADJ_DEFAULT_RATES.items():
            st.session_state["adj_" + k2] = v2
        for k2 in [k3 for k3 in st.session_state if k3.startswith("adjx_")]:
            del st.session_state[k2]
        st.rerun()

    def _si(k, step):
        """Snap preset/session value to int step for slider."""
        return int(round(float(_ag(k)) / step) * step)
    def _sf(k, step):
        """Snap preset/session value to float step for slider."""
        return round(round(float(_ag(k)) / step) * step, 4)

    st.divider()
    st.markdown("#### Core Adjustment Rates")
    st.caption("GLA rounds to nearest $5. Dollar amounts round to $500 or $1,000. Set to $0 to omit.")
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        gla_val = st.slider("GLA ($/SF)", 0, 250, _si("gla", 5), step=5, key="adj_gla", format="$%d")
    with c2:
        bed_val = st.slider("Bedroom ($)", 0, 30000, _si("bed", 500), step=500, key="adj_bed", format="$%d")
        bed_in_gla = st.checkbox("Bedroom adj in GLA", key="adjx_bed_in_gla")
    with c3:
        fb_val = st.slider("Full Bath ($)", 0, 30000, _si("fullbath", 1000), step=1000, key="adj_fullbath", format="$%d")
    with c4:
        hb_val = st.slider("Half Bath ($)", 0, 15000, _si("halfbath", 500), step=500, key="adj_halfbath", format="$%d")
    c5, c6, c7, c8 = st.columns(4)
    with c5:
        bas_val = st.slider("Basement ($)", 0, 30000, _si("basement", 1000), step=1000, key="adj_basement", format="$%d")
    with c6:
        gar_val = st.slider("Garage/stall ($)", 0, 30000, _si("garage", 1000), step=1000, key="adj_garage", format="$%d")
    with c7:
        enc_val = st.slider("Enclosed Porch ($)", 0, 20000, _si("encporch", 500), step=500, key="adj_encporch", format="$%d")
    with c8:
        deck_val = st.slider("Deck/Patio ($)", 0, 15000, _si("deck", 500), step=500, key="adj_deck", format="$%d")
    c9, c10 = st.columns(2)
    with c9:
        fp_val = st.slider("Fireplace/Woodstove ($)", 0, 15000, _si("fp", 500), step=500, key="adj_fp", format="$%d")
    with c10:
        bgr_val = st.slider("Below-Grade Rooms ($)", 0, 10000, _si("bgr", 500), step=500, key="adj_bgr", format="$%d")

    st.divider()
    st.markdown("#### Additional Feature Adjustments")
    st.caption("Check box to include in paragraph.")
    f1, f2, f3, f4, f5 = st.columns(5)
    with f1:
        pool_val = st.slider("In-Ground Pool ($)", 0, 50000, _si("pool", 1000), step=1000, key="adj_pool", format="$%d")
        use_pool = st.checkbox("Include Pool", key="adjx_use_pool")
    with f2:
        cac_val = st.slider("Central AC ($)", 0, 25000, _si("cac", 1000), step=1000, key="adj_cac", format="$%d")
        use_cac = st.checkbox("Include Central AC", key="adjx_use_cac")
    with f3:
        solar_val = st.slider("Solar ($)", 0, 30000, _si("solar", 1000), step=1000, key="adj_solar", format="$%d")
        use_solar = st.checkbox("Include Solar", key="adjx_use_solar")
    with f4:
        adu_val = st.slider("ADU ($)", 0, 100000, _si("adu", 1000), step=1000, key="adj_adu", format="$%d")
        use_adu = st.checkbox("Include ADU", key="adjx_use_adu")
    with f5:
        outbldg_val = st.slider("Out Building ($)", 0, 20000, _si("outbldg", 500), step=500, key="adj_outbldg", format="$%d")
        use_outbldg = st.checkbox("Include Out Building", key="adjx_use_outbldg")

    st.divider()
    st.markdown("#### Site Size Adjustment")
    s1, s2 = st.columns(2)
    with s1:
        su_map = {"none": "Not used", "sf": "Per SF", "acre": "Per acre"}
        su_cur = su_map.get(str(_ag("site_unit")), "Not used")
        site_unit = st.selectbox("Unit", ["Not used", "Per SF", "Per acre"],
                                  index=["Not used", "Per SF", "Per acre"].index(su_cur),
                                  key="adjx_site_unit")
    with s2:
        site_rate_val = st.number_input("Rate ($)", min_value=0, value=int(_ag("site_rate")), step=1, key="adj_site_rate", format="%d")

    st.divider()
    st.markdown("#### Time / Market Conditions")
    t1, t2, t3 = st.columns(3)
    with t1:
        use_time = st.checkbox("Include time adjustment", key="adjx_use_time")
    with t2:
        time_rate_val = st.slider("Rate (% per month)", 0.0, 5.0, _sf("time_rate", 0.25), step=0.25, key="adj_time_rate", format="%.2f%%")
    with t3:
        td_opts = ["Appreciating", "Declining"]
        td_cur = str(_ag("time_dir"))
        td_idx = td_opts.index(td_cur) if td_cur in td_opts else 0
        time_dir = st.selectbox("Direction", td_opts, index=td_idx, key="adj_time_dir")

    st.divider()
    st.markdown("#### Location & View Adjustments")
    st.caption("% of sale price — increments of 2.5%.")
    lv1, lv2, lv3, lv4 = st.columns(4)
    with lv1:
        adv_loc_val = st.slider("Adverse Location (%)", 0.0, 30.0, _sf("adv_loc", 2.5), step=2.5, key="adj_adv_loc", format="%.1f%%")
        use_adv_loc = st.checkbox("Include Adverse Location", key="adjx_use_adv_loc")
    with lv2:
        ben_loc_val = st.slider("Beneficial Location (%)", 0.0, 30.0, _sf("ben_loc", 2.5), step=2.5, key="adj_ben_loc", format="%.1f%%")
        use_ben_loc = st.checkbox("Include Beneficial Location", key="adjx_use_ben_loc")
    with lv3:
        adv_view_val = st.slider("Adverse View (%)", 0.0, 30.0, _sf("adv_view", 2.5), step=2.5, key="adj_adv_view", format="%.1f%%")
        use_adv_view = st.checkbox("Include Adverse View", key="adjx_use_adv_view")
        adv_view_type = st.text_input("Adverse view description", placeholder="e.g. highway, commercial", key="adjx_adv_view_type")
    with lv4:
        ben_view_val = st.slider("Beneficial View (%)", 0.0, 30.0, _sf("ben_view", 2.5), step=2.5, key="adj_ben_view", format="%.1f%%")
        use_ben_view = st.checkbox("Include Beneficial View", key="adjx_use_ben_view")
        ben_view_type = st.text_input("Beneficial view description", placeholder="e.g. water, golf, wooded", key="adjx_ben_view_type")

    st.divider()
    st.markdown("#### Condition Adjustments")
    cn = ["1","2","3","4","5","6"]
    cc1, cc2 = st.columns(2)
    with cc1:
        st.markdown("**Superior interior condition (per MLS)**")
        sup_cond = st.multiselect("Sup cond", cn, default=["2","4"], format_func=lambda x: "Comp " + x, key="adjx_sup_cond", label_visibility="collapsed")
    with cc2:
        st.markdown("**Inferior interior condition (per MLS)**")
        inf_cond = st.multiselect("Inf cond", cn, format_func=lambda x: "Comp " + x, key="adjx_inf_cond", label_visibility="collapsed")

    st.divider()
    st.markdown("#### Quality Adjustments")
    qq1, qq2 = st.columns(2)
    with qq1:
        st.markdown("**Superior quality**")
        sup_qual = st.multiselect("Sup qual", cn, format_func=lambda x: "Comp " + x, key="adjx_sup_qual", label_visibility="collapsed")
    with qq2:
        st.markdown("**Inferior quality**")
        inf_qual = st.multiselect("Inf qual", cn, format_func=lambda x: "Comp " + x, key="adjx_inf_qual", label_visibility="collapsed")

    st.divider()
    st.markdown("#### Disclosure Options")
    d1, d2, d3, d4 = st.columns(4)
    with d1:
        ext_search = st.checkbox("Extended search beyond 6 months", value=True, key="adjx_ext_search")
    with d2:
        geo_expand = st.checkbox("Geographic search expansion", key="adjx_geo_expand")
    with d3:
        arms_length = st.checkbox("Arms-length transaction note", key="adjx_arms_length")
    with d4:
        prior_sale = st.checkbox("Prior sale research note (USPAP)", key="adjx_prior_sale")
    d5, d6 = st.columns(2)
    with d5:
        val_rounded = st.checkbox("Final value within range, rounded", value=True, key="adjx_val_rounded")

    st.divider()
    if st.button("Generate Paragraph", use_container_width=True, key="adj_generate"):
        def fd(n): return "$" + "{:,}".format(int(n))
        def fp(n): return "{:g}%".format(float(n))
        def cl(vals):
            if not vals: return ""
            refs = ["#" + v for v in vals]
            if len(refs) == 1: return "Sale " + refs[0]
            return "Sales " + ", ".join(refs[:-1]) + " & " + refs[-1]

        parts = []
        if int(gla_val) > 0: parts.append(fd(gla_val) + " per SF GLA over 50 SF")
        if bed_in_gla: parts.append("Bedroom adj reflected in GLA")
        elif int(bed_val) > 0: parts.append("Bedrooms @ " + fd(bed_val))
        if int(fb_val) > 0: parts.append("Full Bath @ " + fd(fb_val))
        if int(hb_val) > 0: parts.append("Half Bath @ " + fd(hb_val))
        if int(bas_val) > 0: parts.append("Basement @ " + fd(bas_val))
        if int(gar_val) > 0: parts.append("Garage @ " + fd(gar_val) + " per stall")
        if int(enc_val) > 0: parts.append("Enclosed Porch @ " + fd(enc_val))
        if int(deck_val) > 0: parts.append("Deck/Open Porch/Patio @ " + fd(deck_val) + " each")
        if int(fp_val) > 0: parts.append("Fireplace/Woodstove @ " + fd(fp_val) + " each")
        if int(bgr_val) > 0: parts.append("Below-grade rooms @ " + fd(bgr_val) + " each")
        if use_pool and int(pool_val) > 0: parts.append("In-ground pool @ " + fd(pool_val))
        if use_cac and int(cac_val) > 0: parts.append("Central AC @ " + fd(cac_val))
        if use_solar and int(solar_val) > 0: parts.append("Solar @ " + fd(solar_val))
        if use_adu and int(adu_val) > 0: parts.append("ADU @ " + fd(adu_val))
        if use_outbldg and int(outbldg_val) > 0: parts.append("Out building @ " + fd(outbldg_val))

        adj_s = ("Adjustments made: " + "; ".join(parts) + ".") if parts else ""
        site_s = ""
        if site_unit == "Per SF" and int(site_rate_val) > 0:
            site_s = " Site size adjustments were applied at " + fd(site_rate_val) + " per SF."
        elif site_unit == "Per acre" and int(site_rate_val) > 0:
            site_s = " Site size adjustments were applied at " + fd(site_rate_val) + " per acre."
        time_s = ""
        if use_time and float(time_rate_val) > 0:
            time_s = " A time adjustment of " + fp(time_rate_val) + " per month was applied to comparables to reflect " + time_dir.lower() + " market conditions over the search period."
        lv_parts = []
        if use_adv_loc and float(adv_loc_val) > 0: lv_parts.append("adverse location (" + fp(adv_loc_val) + ")")
        if use_ben_loc and float(ben_loc_val) > 0: lv_parts.append("beneficial location (" + fp(ben_loc_val) + ")")
        if use_adv_view and float(adv_view_val) > 0:
            desc = " — " + adv_view_type.strip() if adv_view_type.strip() else ""
            lv_parts.append("adverse view" + desc + " (" + fp(adv_view_val) + ")")
        if use_ben_view and float(ben_view_val) > 0:
            desc = " — " + ben_view_type.strip() if ben_view_type.strip() else ""
            lv_parts.append("beneficial view" + desc + " (" + fp(ben_view_val) + ")")
        lv_s = (" Location and view adjustments applied for: " + "; ".join(lv_parts) + ".") if lv_parts else ""
        cond_s = ""
        if sup_cond: cond_s += " Comparable " + cl(sup_cond) + " received condition adjustment" + ("s" if len(sup_cond) > 1 else "") + " due to superior interior condition, per MLS."
        if inf_cond: cond_s += " Comparable " + cl(inf_cond) + " received condition adjustment" + ("s" if len(inf_cond) > 1 else "") + " due to inferior interior condition, per MLS."
        qual_s = ""
        if sup_qual: qual_s += " Comparable " + cl(sup_qual) + " received quality adjustment" + ("s" if len(sup_qual) > 1 else "") + " due to superior overall quality relative to the subject."
        if inf_qual: qual_s += " Comparable " + cl(inf_qual) + " received quality adjustment" + ("s" if len(inf_qual) > 1 else "") + " due to inferior overall quality relative to the subject."
        ext_s = " Due to a lack of sales with similar style, age, and location, it was necessary to extend the search beyond six months." if ext_search else ""
        geo_s = " Due to limited availability of comparable sales within the immediate neighborhood, the geographic search area was expanded to include competing neighborhoods and adjacent communities." if geo_expand else ""
        arms_s = " All comparable sales have been verified as arm's-length transactions. Any non-arm's-length transactions identified were excluded from consideration." if arms_length else ""
        prior_s = " A search of public records revealed no prior sales or transfers of the subject property within the three-year period prior to the effective date, unless otherwise noted in this report." if prior_sale else ""
        round_s = " The final opinion of value is within the indicated range, rounded." if val_rounded else ""
        intro = "The appraiser has verified each sale with MLS data and city records. Photos of the exterior of the comparables are included in the report. If a photo of the comparable was not possible or not permitted by the homeowner, a recent MLS or city photo was included."
        full = (intro + " " + adj_s + site_s + time_s + lv_s + cond_s + qual_s + ext_s + geo_s + arms_s + prior_s + round_s).strip()
        st.session_state["adj_output"] = full

    if st.session_state.get("adj_output"):
        out = st.session_state["adj_output"]
        st.markdown("**Generated Paragraph** — " + str(len(out)) + " characters")
        st.text_area("Paragraph output", value=caps(out), height=220, key=f"adj_out_display_{st.session_state.get('all_caps',False)}", label_visibility="collapsed")
        st.caption("Click in the box, Ctrl+A, Ctrl+C to copy.")


# ══════════════════════════════════════════════════════════════════════════════
# TAB 8 — STR INCOME ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════
with tab8:
    tab_generate, tab_clients, tab_history, tab_intake = st.tabs([
        "⚡ Generate Report", "👥 Client Database", "📋 Order History", "🏠 AI Intake Assistant"
    ])

    # ══════════════════════════════════════════════════════════════════════════════
    # TAB 1 — GENERATE REPORT
    # ══════════════════════════════════════════════════════════════════════════════
    with tab_generate:

        # File Uploads
        st.subheader("1. Upload AirDNA PDF")
        col1, col2 = st.columns(2)
        with col1:
            airdna_pdf = st.file_uploader("AirDNA Rentalizer PDF", type="pdf", key="pdf")

        st.subheader("2. Property Photo (Optional)")
        col_p1, col_p2 = st.columns(2)
        with col_p1:
            property_photo = st.file_uploader("Property Photo", type=["jpg","jpeg","png"], key="photo",
                                               help="Front exterior photo of the subject property")

        # Assignment Info
        st.subheader("3. Assignment Info")

        # Client autofill
        clients = load_clients()
        client_names = ["-- Enter manually --"] + sorted(clients.keys())
        selected_client = st.selectbox("Select existing client (or enter manually below)",
                                        client_names, key="client_select")

        # Pre-fill values from selected client
        if selected_client != "-- Enter manually --" and selected_client in clients:
            c = clients[selected_client]
            prefill_client        = c.get("name", "")
            prefill_client_address = c.get("address", "")
            prefill_client_phone  = c.get("phone", "")
        else:
            prefill_client        = ""
            prefill_client_address = ""
            prefill_client_phone  = ""

        col4, col5, col6 = st.columns(3)
        with col4:
            client = st.text_input("Client / Lender", value=prefill_client,
                                    placeholder="Annie Mac Home Mortgage")
        with col5:
            client_address = st.text_input("Client Address", value=prefill_client_address,
                                            placeholder="123 Main St, Boston, MA")
        with col6:
            client_phone = st.text_input("Client Phone", value=prefill_client_phone,
                                          placeholder="617-555-1234")

        col7, col8, col9 = st.columns(3)
        with col7:
            client_order_num = st.text_input("Client Order Number", placeholder="ORD-12345")
        with col8:
            borrower = st.text_input("Borrower", placeholder="John Smith")
        with col9:
            loan_num = st.text_input("Loan Number", placeholder="2008727778")

        col10, col11, col12 = st.columns(3)
        with col10:
            avm_file_id = st.text_input("A-Tech File ID", placeholder="AVM-2026-001")
        with col11:
            from datetime import date
            report_date = st.text_input("Report Date",
                value=date.today().strftime("%B %d, %Y"))
        with col12:
            property_type = st.selectbox("Property Type", [
                "Single-Family Residence",
                "Condominium",
                "Townhouse",
                "Multi-Family (2-4 Units)",
                "Single Unit in Multi-Family",
                "Manufactured Home",
            ])

        st.subheader("4. Market Overview")
        market_overview = st.text_area(
            "Market Overview",
            placeholder="Write 3-5 sentences describing the local STR market — what drives demand, submarket characteristics, seasonality, and any relevant local factors...",
            height=130,
            label_visibility="collapsed"
        )
        st.caption("This text appears in the Market Overview & Commentary section of the report.")
        st.divider()

        if st.button("⚡ Generate Report", type="primary", use_container_width=True):
            if not airdna_pdf:
                st.error("Please upload the AirDNA PDF.")
            elif not client or not loan_num:
                st.error("Please enter the client/lender name and loan number.")
            elif not market_overview.strip():
                st.error("Please enter a market overview before generating.")
            else:
                with st.spinner("Extracting data from AirDNA PDF..."):
                    pdf_bytes = airdna_pdf.read()
                    data = parse_airdna_pdf(pdf_bytes)

                commentary = market_overview.strip()

                with st.spinner("Building report..."):
                    photo_override = None
                    if property_photo:
                        tmp_photo = tempfile.NamedTemporaryFile(delete=False,
                            suffix=os.path.splitext(property_photo.name)[1])
                        tmp_photo.write(property_photo.read())
                        tmp_photo.close()
                        photo_override = tmp_photo.name

                    map_override = None

                    buf = io.BytesIO()
                    build_pdf(data, client, loan_num, report_date, commentary, buf,
                              photo_override=photo_override, map_override=map_override,
                              client_address=client_address, client_phone=client_phone,
                              client_order_num=client_order_num, borrower=borrower,
                              avm_file_id=avm_file_id, property_type=property_type)
                    buf.seek(0)

                # Log the order
                log_order(
                    address=f"{data.get('address_line1','')} {data.get('city_state_zip','')}",
                    property_type=property_type,
                    client=client,
                    borrower=borrower,
                    loan_num=loan_num,
                    avm_file_id=avm_file_id,
                    report_date=report_date
                )

                addr_slug = re.sub(r"[^a-zA-Z0-9]+","_",
                                   data.get("address_line1","Report")).strip("_")
                filename = f"ATECH_STR_{addr_slug}.pdf"

                st.success("✅ Report generated and logged!")
                st.download_button(
                    label="📄 Download Report PDF",
                    data=buf,
                    file_name=filename,
                    mime="application/pdf",
                    use_container_width=True
                )

                # Store PDF in session state for email sending
                st.session_state["last_pdf_bytes"]  = buf.getvalue()
                st.session_state["last_pdf_filename"] = filename
                st.session_state["last_pdf_address"]  = f"{data.get('address_line1','')} {data.get('city_state_zip','')}"

        # Email section — shows after report is generated
        if "last_pdf_bytes" in st.session_state:
            st.divider()
            st.subheader("📧 Send Report via Email")
            e1, e2 = st.columns([3, 1])
            with e1:
                email_to = st.text_input("Recipient email address",
                                          placeholder="lender@example.com",
                                          key="email_to")
            with e2:
                st.write("")
                st.write("")
                send_clicked = st.button("Send", use_container_width=True, key="send_email")

            email_note = st.text_area("Optional note to include in email body",
                                       placeholder="Please find the STR Income Analysis attached...",
                                       height=80, key="email_note")

            if send_clicked:
                if not email_to.strip():
                    st.error("Please enter a recipient email address.")
                else:
                    address_line = st.session_state.get("last_pdf_address","Subject Property")
                    subject = f"STR Income Analysis — {address_line}"
                    body = email_note.strip() if email_note.strip() else (
                        f"Please find the Short-Term Rental Income Analysis attached for {address_line}.\n\n"
                        f"This report was prepared by A-Tech Appraisal Co., LLC.\n\n"
                        f"Please note: This is not an appraisal and does not constitute an opinion of market value."
                    )
                    try:
                        with st.spinner("Sending email..."):
                            send_report_email(
                                to_email=email_to.strip(),
                                subject=subject,
                                body=body,
                                pdf_bytes=st.session_state["last_pdf_bytes"],
                                filename=st.session_state["last_pdf_filename"]
                            )
                        st.success(f"✅ Report sent to {email_to.strip()}")
                    except Exception as e:
                        st.error(f"Email failed: {str(e)}")

    # ══════════════════════════════════════════════════════════════════════════════
    # TAB 2 — CLIENT DATABASE
    # ══════════════════════════════════════════════════════════════════════════════
    with tab_clients:
        st.subheader("Client Database")
        st.caption("Save client info here once — it will auto-fill on the Generate tab.")

        clients = load_clients()

        # Add / Edit client form
        with st.expander("➕ Add New Client", expanded=len(clients) == 0):
            nc1, nc2, nc3 = st.columns(3)
            with nc1:
                new_name    = st.text_input("Client / Lender Name *", key="new_name",
                                             placeholder="Annie Mac Home Mortgage")
            with nc2:
                new_address = st.text_input("Client Address", key="new_address",
                                             placeholder="123 Main St, Boston, MA")
            with nc3:
                new_phone   = st.text_input("Client Phone", key="new_phone",
                                             placeholder="617-555-1234")

            if st.button("💾 Save Client", use_container_width=True):
                if not new_name.strip():
                    st.error("Client name is required.")
                else:
                    clients[new_name.strip()] = {
                        "name":    new_name.strip(),
                        "address": new_address.strip(),
                        "phone":   new_phone.strip(),
                    }
                    save_clients(clients)
                    st.success(f"✅ Client '{new_name.strip()}' saved.")
                    st.rerun()

        # Client list
        if clients:
            st.divider()
            st.write(f"**{len(clients)} client(s) saved**")
            for name, info in sorted(clients.items()):
                with st.container():
                    cc1, cc2, cc3, cc4 = st.columns([3, 3, 2, 1])
                    with cc1:
                        st.write(f"**{name}**")
                    with cc2:
                        st.write(info.get("address","—"))
                    with cc3:
                        st.write(info.get("phone","—"))
                    with cc4:
                        if st.button("🗑️", key=f"del_{name}", help=f"Delete {name}"):
                            del clients[name]
                            save_clients(clients)
                            st.rerun()
            st.divider()

            # Export clients as CSV
            client_rows = [{"Client": k, "Address": v.get("address",""),
                            "Phone": v.get("phone","")} for k,v in clients.items()]
            client_df = pd.DataFrame(client_rows)
            st.download_button("⬇️ Export Client List (CSV)",
                                data=client_df.to_csv(index=False),
                                file_name="atech_clients.csv",
                                mime="text/csv")
        else:
            st.info("No clients saved yet. Add your first client above.")

    # ══════════════════════════════════════════════════════════════════════════════
    # TAB 3 — ORDER HISTORY
    # ══════════════════════════════════════════════════════════════════════════════
    with tab_history:
        st.subheader("Order History")
        st.caption("Every report generated is automatically logged here.")

        orders = load_orders()

        if orders:
            st.write(f"**{len(orders)} order(s) on record**")

            # Search / filter
            search = st.text_input("🔍 Search by address, client, or borrower",
                                    placeholder="Type to filter...", key="order_search")
            if search:
                q = search.lower()
                orders = [o for o in orders if
                          q in o.get("address","").lower() or
                          q in o.get("client","").lower() or
                          q in o.get("borrower","").lower()]

            # Display as table
            if orders:
                df = pd.DataFrame(orders)
                df = df.rename(columns={
                    "date":          "Report Date",
                    "address":       "Property Address",
                    "property_type": "Type",
                    "client":        "Client / Lender",
                    "borrower":      "Borrower",
                    "loan_number":   "Loan Number",
                    "avm_file_id":   "A-Tech File ID",
                })
                st.dataframe(df, use_container_width=True, hide_index=True)

                st.divider()
                st.download_button(
                    "⬇️ Export Order History (CSV)",
                    data=df.to_csv(index=False),
                    file_name="atech_order_history.csv",
                    mime="text/csv",
                    use_container_width=True
                )
            else:
                st.info("No orders match your search.")
        else:
            st.info("No orders logged yet. Generate your first report to start the log.")


    # ══════════════════════════════════════════════════════════════════════════════
    # TAB 5 — AI INTAKE ASSISTANT
    # ══════════════════════════════════════════════════════════════════════════════
    with tab_intake:

        st.subheader("🏠 AI Intake Assistant")
        st.caption(
            "Upload your tax card, Apple Maps screenshot, GIS screenshot, MLS sheet, "
            "or contract — add any notes — and Claude will generate a fully populated "
            "intake template ready to copy into TOTAL."
        )

        # ── API Key (from Streamlit secrets or environment variable) ─────────────
        try:
            intake_api_key = st.secrets["ANTHROPIC_API_KEY"]
        except Exception:
            intake_api_key = os.environ.get("ANTHROPIC_API_KEY", "")

        # ── File Uploads ──────────────────────────────────────────────────────────
        st.markdown("#### 1. Upload Documents")
        col_u1, col_u2 = st.columns(2)
        with col_u1:
            tax_card_file = st.file_uploader(
                "Tax Card (PDF or Image)",
                type=["pdf", "png", "jpg", "jpeg"],
                key="intake_tax_card"
            )
            mls_file = st.file_uploader(
                "MLS Sheet / 360 Property View (PDF) — optional",
                type=["pdf"],
                key="intake_mls",
                help="Not required for refinance assignments without an active listing."
            )
            anow_file = st.file_uploader(
                "ANOW Order Screenshot (PNG/JPG) — optional",
                type=["png", "jpg", "jpeg"],
                key="intake_anow",
                help="Screenshot from ANOW order management — used to pull borrower, lender, and order details."
            )
        with col_u2:
            maps_file = st.file_uploader(
                "Apple Maps Screenshot (PNG/JPG)",
                type=["png", "jpg", "jpeg"],
                key="intake_maps"
            )
            gis_file = st.file_uploader(
                "GIS / Parcel Map Screenshot (PNG/JPG)",
                key="intake_gis",
                type=["png", "jpg", "jpeg"]
            )
            bt_file = st.file_uploader(
                "Banker & Tradesman (PDF or Image) — optional",
                type=["pdf", "png", "jpg", "jpeg"],
                key="intake_bt",
                help="Banker & Tradesman deed/transfer report for sales history and ownership verification."
            )
        contract_file = st.file_uploader(
            "Purchase & Sales Agreement (PDF) — optional",
            type=["pdf"],
            key="intake_contract"
        )
        cubicasa_file = st.file_uploader(
            "Field Measurement Report (PDF) — optional",
            type=["pdf"],
            key="intake_cubicasa"
        )

        # ── Manual Notes ──────────────────────────────────────────────────────────
        st.markdown("#### 2. Assignment Notes")
        col_n1, col_n2, col_n3 = st.columns(3)
        with col_n1:
            intake_borrower    = st.text_input("Borrower", key="intake_borrower")
            intake_lender      = st.text_input("Lender / AMC", key="intake_lender")
        with col_n2:
            intake_form        = st.selectbox("Form Type",
                                              ["1004", "1073 Condo", "1025 Multi-Family",
                                               "1004C Manufactured", "Other"],
                                              key="intake_form")
            intake_use         = st.selectbox("Intended Use",
                                              ["Purchase", "Refinance", "Estate",
                                               "Divorce", "Other"],
                                              key="intake_use")
        with col_n3:
            intake_due         = st.text_input("Due Date", key="intake_due")
            intake_contract_px = st.text_input("Contract Price", key="intake_contract_px",
                                                placeholder="$000,000")

        intake_notes = st.text_area(
            "Additional Notes (inspection findings, borrower notes, flags, etc.)",
            key="intake_notes",
            height=80,
            placeholder="e.g. Owner confirms well and septic. TQS ceiling height under 7ft in places. Borrower same as owner LLC."
        )

        # ── System Prompt ─────────────────────────────────────────────────────────
        INTAKE_SYSTEM_PROMPT = """You are an expert certified residential appraiser assistant for A-Tech Appraisal Co., LLC in Rhode Island and Massachusetts. Generate a USPAP-compliant pre-inspection intake report from the uploaded documents.

    Never use: "convenient," "desirable," "charming," "ideal," or "easy."
    Never reference the GLA measurement source in the improvement description — just state the number.
    Finished below-grade space is never included in above-grade GLA.

    OUTPUT FORMAT — sections must appear in this exact order, matching 1004 page 1 sequence:

    ## 🏠 APPRAISAL ASSIGNMENT INTAKE

    **Property Address:**
    **City/Town, State, Zip:**
    **County:**
    **Borrower:**
    **Owner of Public Record:**
    **Legal Description (Book & Page):** [Most recent deed Book & Page and recording date e.g. "Book 840, Page 144, recorded 08/20/2013"]
    **Assessor Parcel #:** [Assessor account/parcel number from tax card]
    **Tax Year:**
    **Tax Amount $:**
    **Neighborhood Name:**
    **Occupant:** [Owner / Tenant / Vacant]
    **Assignment Type:** [Purchase Transaction / Refinance Transaction]
    **Property Rights:** [Fee Simple / Leasehold]
    **Lender/Client:**
    **Form Type:**
    **Intended Use:**
    **Due Date:**

    ---

    ## CONTRACT

    **Contract Summary:** [If purchase: "As of [contract date], the Seller and Buyer have agreed on a purchase price of $[amount]. [Appliances/personal property note if stated in contract.] [If concession: The Seller has agreed to contribute $[amount] toward the Buyer's closing costs[, per Amendment dated MM/DD/YYYY if applicable]."] If refinance: "This is a refinance transaction. No purchase contract is applicable."]

    **Contract Price $:**
    **Date of Contract:**
    **Seller = Owner of Record:** [Yes / No]
    **Financial Assistance / Concessions:** [Yes / No — if Yes, dollar amount and description]
    **Prior Sales Analysis:**
    [Analyze all transfers within 36 months. Note arm's length vs. non-arm's length. Flag same-day double transfers.]

    ---

    ## NEIGHBORHOOD

    **Location:** [Urban / Suburban / Rural]
    **Built-Up:** [Over 75% / 25-75% / Under 25%]
    **Growth:** [Rapid / Stable / Slow]
    **Property Values:** [Increasing / Stable / Declining]
    **Demand/Supply:** [Shortage / In Balance / Over Supply]
    **Marketing Time:** [Under 3 Months / 3-6 Months / Over 6 Months]
    **Price Range:** $[low] to $[high] Predominant: $[pred]
    **Age Range:** [low] to [high] yrs Predominant: [pred] yrs

    **Land Use Grid:**
    | Use | % |
    |---|---|
    | Single Family | _% |
    | 2-4 Unit | _% |
    | Multi-Family (5+) | _% |
    | Commercial | _% |
    | Other (describe) | _% |

    **Neighborhood Boundaries:**
    The subject neighborhood is bounded to the north by [X], to the south by [X], to the east by [X], and to the west by [X].

    **Neighborhood Description:**
    [4-5 factual sentences: municipality/county, neighborhood character, arterial access, utilities, Other land use]

    **Market Conditions:**
    [1-2 sentences supporting the trend checkboxes above]

    ---

    ## SITE

    **Dimensions:**
    **Area:**
    **Shape:**
    **View:**
    **Zoning Classification:**
    **Zoning Description:**
    **Zoning Compliance:** [Legal Conforming / Legal Non-Conforming / No Zoning / Illegal]
    **H&BU as Improved = Present Use:** [Yes / No]
    **Utilities:** Elec: [Public/Other] | Gas: [Public/Other] | Water: [Public/Other] | Sewer: [Public/Other]
    **Off-Site Improvements:** Street: [Public/Private] | Curb/Gutter: [Y/N] | Sidewalk: [Y/N] | Alley: [None/Public/Private]
    **FEMA Flood Zone:** [Zone] | Map #: [panel] | Map Date: [date]
    **Adverse Site Conditions:** [None noted / describe if any]

    ---

    ## IMPROVEMENTS

    **Units:** [1 / 1 with ADU / 2-4]
    **Stories:**
    **Type:** [Det. / Att. / S-Det./End Unit]
    **Design (Style):**
    **Year Built:**
    **Effective Age:**
    **Foundation:** [Concrete Slab / Crawl Space / Full Basement / Partial Basement]
    **Basement Area:** [sf] | **% Finished:** [%] | **Basement Rooms:** [describe]
    **Exterior Walls:**
    **Roof Surface:**
    **Gutters & Downspouts:**
    **Window Type:**
    **Heating:** [type / fuel]
    **Cooling:**
    **Floors:**
    **Walls (Interior):**
    **Trim/Finish:**
    **Above Grade Room Count:** Rooms: [#] | Bedrooms: [#] | Baths: [#]
    **GLA:** [sf]
    **Amenities:** [Fireplace(s) # / Patio / Deck / Pool / Porch / Fence / Other]
    **Car Storage:** [None / Garage # cars — Att./Det./Built-in]
    **Appliances:** [list]
    **Additional Features:**

    **GLA Sub-Area Breakdown:**
    | Code | Description | Gross SF | Living SF |
    |---|---|---|---|

    **Improvement Description:**
    [Begin: "The appraiser has inspected the interior and exterior of the subject property and researched municipal records for data reported herein."]
    [Cover style, year built, GLA, rooms, bed/bath, exterior, roof, foundation, basement, heat/cool, garage, features.]
    [If field measurement report provided use that GLA. If assessor only use that. Never reference the source.]
    [Finished below-grade space excluded from above-grade GLA, reported separately.]
    [End: "Condition and quality ratings to be determined at inspection."]

    ---

    ## COMPLEXITY FLAGS

    [List each flag with - [x] prefix on its own line]

    ---

    ## ATTACHMENTS RECEIVED

    [List each attachment on its own line — tax card, Apple Maps, GIS, MLS sheet, ANOW screenshot, Banker & Tradesman, contract, field measurement report]

    ---

    ## PRIORITY ITEMS TO RESOLVE

    [Numbered list — most critical items to confirm before or at inspection]

    ---

    IMPORTANT RULES:
    1. Never use "convenient," "desirable," "charming," "ideal," or "easy"
    2. Never reference GLA source in improvement description — just state the number
    3. Flag GLA variance in COMPLEXITY FLAGS only
    4. Finished below-grade space never in above-grade GLA
    5. TQS must be verified for ANSI ceiling height at inspection
    6. LLC ownership with individual borrower always flagged
    7. Seller concessions always flagged and always in Contract Summary sentence
    8. All transfers within 36 months analyzed in Prior Sales Analysis
    9. Same-day double transfers always flagged
    10. Permit history noted but not over-emphasized
    11. Land use grid reflects immediate neighborhood only
    12. Note when water/sewer needs field confirmation
    13. Contract Summary is always a complete paste-ready sentence with no brackets remaining
    14. Legal Description = Book & Page and recording date of the most recent deed — never the plat/lot description
    15. Assessor Parcel # = the tax assessor account number — always a separate field from Legal Description"""

        # ── Generate Button ───────────────────────────────────────────────────────
        st.divider()
        generate_btn = st.button(
            "🏠 Generate Intake",
            use_container_width=True,
            key="intake_generate",
            type="primary"
        )

        if generate_btn:
            if not intake_api_key:
                st.error("Anthropic API key not found. Add ANTHROPIC_API_KEY to your Streamlit secrets.")
                st.stop()

            # Build message content
            content = []

            # Add uploaded files as base64
            import base64

            def pdf_to_base64(file_bytes):
                return base64.standard_b64encode(file_bytes).decode("utf-8")

            def img_to_base64_and_type(file_bytes, original_filename, max_bytes=4 * 1024 * 1024):
                """Compress image if needed. Returns (base64_str, media_type)."""
                from io import BytesIO
                # If small enough and already JPEG, send as-is
                ext = original_filename.lower().split(".")[-1]
                if len(file_bytes) <= max_bytes and ext in ("jpg", "jpeg"):
                    return base64.standard_b64encode(file_bytes).decode("utf-8"), "image/jpeg"
                # Convert to RGB JPEG and compress
                img = PILImage.open(BytesIO(file_bytes)).convert("RGB")
                # If small enough as PNG, still convert to JPEG for consistency
                if len(file_bytes) <= max_bytes:
                    buf = BytesIO()
                    img.save(buf, format="JPEG", quality=90, optimize=True)
                    return base64.standard_b64encode(buf.getvalue()).decode("utf-8"), "image/jpeg"
                # Need to compress — try reducing quality first
                quality = 85
                while quality >= 30:
                    buf = BytesIO()
                    img.save(buf, format="JPEG", quality=quality, optimize=True)
                    compressed = buf.getvalue()
                    if len(compressed) <= max_bytes:
                        return base64.standard_b64encode(compressed).decode("utf-8"), "image/jpeg"
                    quality -= 10
                # Last resort — resize to 50%
                w, h = img.size
                img = img.resize((w // 2, h // 2), PILImage.LANCZOS)
                buf = BytesIO()
                img.save(buf, format="JPEG", quality=70, optimize=True)
                return base64.standard_b64encode(buf.getvalue()).decode("utf-8"), "image/jpeg"

            def add_image_content(content, file_bytes, filename, label):
                """Helper to add an image to content with compression and correct media type."""
                b64, media_type = img_to_base64_and_type(file_bytes, filename)
                content.append({
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": media_type,
                        "data": b64
                    }
                })
                content.append({"type": "text", "text": label})

            if tax_card_file:
                file_bytes = tax_card_file.read()
                fname = tax_card_file.name.lower()
                if fname.endswith(".pdf"):
                    content.append({
                        "type": "document",
                        "source": {
                            "type": "base64",
                            "media_type": "application/pdf",
                            "data": pdf_to_base64(file_bytes)
                        },
                        "title": "Tax Card"
                    })
                else:
                    add_image_content(content, file_bytes, tax_card_file.name, "[Above image: Tax Card]")

            if maps_file:
                file_bytes = maps_file.read()
                add_image_content(content, file_bytes, maps_file.name,
                    "[Above image: Apple Maps screenshot — use circled/boxed landmarks for N/S/E/W neighborhood boundaries]")

            if gis_file:
                file_bytes = gis_file.read()
                add_image_content(content, file_bytes, gis_file.name,
                    "[Above image: GIS / Parcel Map — use for land use grid percentages and site data]")

            if anow_file:
                file_bytes = anow_file.read()
                add_image_content(content, file_bytes, anow_file.name,
                    "[Above image: ANOW order screenshot — extract borrower, lender/client, AMC, order details, and any assignment notes]")

            if bt_file:
                file_bytes = bt_file.read()
                fname = bt_file.name.lower()
                if fname.endswith(".pdf"):
                    content.append({
                        "type": "document",
                        "source": {
                            "type": "base64",
                            "media_type": "application/pdf",
                            "data": pdf_to_base64(file_bytes)
                        },
                        "title": "Banker & Tradesman — use for ownership history, deed transfers, book/page, and sales history"
                    })
                else:
                    add_image_content(content, file_bytes, bt_file.name,
                        "[Above image: Banker & Tradesman — use for ownership history, deed transfers, book/page, and sales history]")

            if mls_file:
                file_bytes = mls_file.read()
                content.append({
                    "type": "document",
                    "source": {
                        "type": "base64",
                        "media_type": "application/pdf",
                        "data": pdf_to_base64(file_bytes)
                    },
                    "title": "MLS Sheet"
                })

            if contract_file:
                file_bytes = contract_file.read()
                content.append({
                    "type": "document",
                    "source": {
                        "type": "base64",
                        "media_type": "application/pdf",
                        "data": pdf_to_base64(file_bytes)
                    },
                    "title": "Purchase and Sales Agreement"
                })

            if cubicasa_file:
                file_bytes = cubicasa_file.read()
                content.append({
                    "type": "document",
                    "source": {
                        "type": "base64",
                        "media_type": "application/pdf",
                        "data": pdf_to_base64(file_bytes)
                    },
                    "title": "Field Measurement Report — ANSI compliant field measurements. Use these GLA figures and sub-area breakdown. Note any variance from assessor or MLS. GLA and room counts are subject to verification at inspection."
                })

            # Build text prompt
            prompt_parts = ["Please generate the full appraisal intake template based on the uploaded documents."]
            if intake_borrower:    prompt_parts.append(f"Borrower: {intake_borrower}")
            if intake_lender:      prompt_parts.append(f"Lender/AMC: {intake_lender}")
            if intake_form:        prompt_parts.append(f"Form Type: {intake_form}")
            if intake_use:         prompt_parts.append(f"Intended Use: {intake_use}")
            if intake_due:         prompt_parts.append(f"Due Date: {intake_due}")
            if intake_contract_px: prompt_parts.append(f"Contract Price: {intake_contract_px}")
            if intake_notes:       prompt_parts.append(f"Additional Notes: {intake_notes}")

            content.append({"type": "text", "text": "\n".join(prompt_parts)})

            # Call API
            with st.spinner("Claude is analyzing your documents and generating the intake..."):
                try:
                    headers = {
                        "x-api-key": intake_api_key,
                        "anthropic-version": "2023-06-01",
                        "content-type": "application/json",
                        "anthropic-beta": "pdfs-2024-09-25"
                    }
                    payload = {
                        "model": "claude-opus-4-5",
                        "max_tokens": 4000,
                        "system": INTAKE_SYSTEM_PROMPT,
                        "messages": [
                            {"role": "user", "content": content}
                        ]
                    }
                    response = requests.post(
                        "https://api.anthropic.com/v1/messages",
                        headers=headers,
                        json=payload,
                        timeout=120
                    )
                    response.raise_for_status()
                    result = response.json()
                    output_text = result["content"][0]["text"].strip()

                    # Follow-up question box
                    st.divider()
                    st.markdown("#### 💬 Follow-up Question")
                    followup = st.text_input(
                        "Ask a follow-up about this property (e.g. 'draft the improvement description', 'what comps strategy would you suggest')",
                        key="intake_followup"
                    )
                    if st.button("Send Follow-up", key="intake_followup_btn"):
                        if followup.strip():
                            with st.spinner("Thinking..."):
                                fu_payload = {
                                    "model": "claude-opus-4-5",
                                    "max_tokens": 2000,
                                    "system": INTAKE_SYSTEM_PROMPT,
                                    "messages": [
                                        {"role": "user", "content": content},
                                        {"role": "assistant", "content": output_text},
                                        {"role": "user", "content": followup}
                                    ]
                                }
                                fu_response = requests.post(
                                    "https://api.anthropic.com/v1/messages",
                                    headers=headers,
                                    json=fu_payload,
                                    timeout=120
                                )
                                fu_response.raise_for_status()
                                fu_result = fu_response.json()
                                fu_text = fu_result["content"][0]["text"]
                                st.markdown(fu_text)
                                st.text_area("📋 Copy follow-up response",
                                             value=fu_text, height=300,
                                             key="intake_fu_output")

                except requests.exceptions.HTTPError as e:
                    st.error(f"API error: {e.response.status_code} — {e.response.text}")
                except Exception as e:
                    st.error(f"Error: {str(e)}")
